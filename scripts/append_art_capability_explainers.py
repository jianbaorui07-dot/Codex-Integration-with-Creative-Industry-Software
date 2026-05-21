from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT_DIR = Path(r"C:\Users\jian\Documents\New project\output")
DOCX_PATH = OUTPUT_DIR / "艺术行业效率插件建议.docx"
BLUE = "204370"
LIGHT_GRAY = "F7FAFD"
LIGHT_BLUE = "EEF5FF"
AMBER = "FFF6E6"
BORDER = "CBD5E1"


def set_font(run, size=None, bold=None, color=None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, cm):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn("w:" + edge))
        if node is None:
            node = OxmlElement("w:" + edge)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def cell_text(cell, text, size=7.4, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(text)
    set_font(r, size=16 if level == 1 else 11.5, bold=True, color=BLUE if level == 1 else "435B7B")


def add_para(doc, text, size=9):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_font(r, size=size)


def add_note(doc, title, body, fill=AMBER):
    table = doc.add_table(rows=1, cols=1)
    set_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    set_font(r, size=10, bold=True, color=BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.line_spacing = 1.12
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_font(r2, size=8.5)


def capabilities():
    return [
        ("创意编程\nCreative Coding", "用代码画画、做动画、做互动。", "像给画笔写规则：不是一笔一笔画，而是让电脑按你的规则画。", "生成海报、互动网页、课堂练习、视觉实验。", "p5.js / Processing", "先做一个随机线条海报。"),
        ("生成艺术\nGenerative Art", "让规则、随机数和参数一起生成作品。", "像种一片花园：你设计种子和环境，具体长相每次不同。", "系列海报、纹样、动态屏保、数字版画。", "p5.js / GLSL / Python", "先做粒子流动或分形树。"),
        ("数据可视化\nData Visualization", "把表格和数字变成图。", "像把账本变成地图和图表，一眼看懂变化。", "销售趋势、观众热力图、项目预算、艺术家关系图。", "D3.js / Observable / Python", "先做作品销售时间线。"),
        ("AI 图像生成\nAI Image Generation", "用文字或参考图生成图像。", "像一个很快的视觉助理，先给你很多草稿方向。", "海报初稿、空间草图、材质参考、情绪板。", "Stable Diffusion / InvokeAI / ComfyUI", "先做风格参考板。"),
        ("提示词\nPrompting", "把想法写成 AI 更听得懂的描述。", "像给摄影师下 brief，说清主题、风格、镜头、材质和禁忌。", "稳定生成风格、减少返工、整理创作语言。", "中英提示词库 / Prompt Library", "先做自己的艺术风格词库。"),
        ("节点工作流\nNode Workflow", "把复杂流程拆成一个个节点，用线连起来。", "像搭积木：每块负责一件事，连起来就是完整机器。", "AI 图像流程、视频流程、实时视觉、投影控制。", "ComfyUI / TouchDesigner", "先学会加载模型、输入提示词、保存图片。"),
        ("网页基础\nHTML CSS JS", "网页的结构、样式和动作。", "HTML 像骨架，CSS 像衣服，JavaScript 像动作。", "作品集网站、展览页面、互动作品、客户表单。", "HTML / CSS / JavaScript", "先做一个作品详情页。"),
        ("组件化界面\nReact Components", "把页面拆成可复用的小组件。", "像展墙模块：标题、图片卡、按钮、表单都能重复用。", "作品集、后台管理、客户预览页、课程平台。", "React", "先做可复用作品卡片。"),
        ("全栈网站\nFull Stack Website", "前台页面、后台数据、登录和接口一起做。", "像开一个小店：前面给客户看，后面管库存和订单。", "艺术电商、客户 CRM、报名系统、私密预览。", "Next.js / Database / Auth", "先做作品询价表单。"),
        ("网页 3D\nWeb 3D", "在浏览器里显示和互动 3D 模型。", "像把小展厅放进网页，观众不用安装软件也能看。", "3D 展柜、虚拟展厅、产品配置器、网页雕塑。", "Three.js / WebGL / GLTF", "先做 GLTF 作品查看器。"),
        ("3D 建模\n3D Modeling", "做立体模型、材质、灯光和渲染。", "像在电脑里捏泥巴、搭空间、打灯拍照。", "雕塑草图、展览空间、动画、数字制造文件。", "Blender", "先建一个小展厅模型。"),
        ("实时 3D\nReal-time 3D", "画面能实时运行和交互。", "像游戏一样，观众动一下，画面马上反应。", "VR 展厅、互动培训、AR 预览、沉浸体验。", "Unity / Unreal / WebXR", "先做一个可走动的虚拟展厅。"),
        ("AR 增强现实\nAugmented Reality", "把虚拟内容叠到真实世界。", "像手机给现实加一层说明和动画。", "手机看画、AR 展签、雕塑摆放、城市艺术路线。", "WebXR / ARKit / ARCore", "先做 AR 作品尺寸预览。"),
        ("VR 虚拟现实\nVirtual Reality", "进入一个完全虚拟的空间。", "像戴上头显走进不存在的展厅。", "虚拟展览、空间绘画、声音空间、策展预演。", "Unity / WebXR", "先做 VR 小展厅。"),
        ("投影映射\nProjection Mapping", "把影像精准投到墙面、建筑或物体上。", "像给真实物体穿上一层会动的皮肤。", "建筑立面、雕塑投影、舞台视觉、沉浸房间。", "TouchDesigner / MadMapper", "先做小雕塑投影映射。"),
        ("实时视觉\nRealtime Visuals", "现场画面实时生成、切换和响应。", "像 DJ 打碟，只是你打的是画面。", "VJ、演出视觉、互动屏、声音反应投影。", "TouchDesigner / Resolume", "先做声音反应粒子。"),
        ("声音分析\nAudio Analysis", "把声音拆成音量、频率、节拍等数据。", "像让电脑听懂音乐哪里强、哪里快、哪里低沉。", "音乐海报、声音反应灯光、节拍剪辑。", "Web Audio / Tone.js / TouchDesigner", "先做音乐频谱海报。"),
        ("传感器\nSensors", "让作品感知距离、触摸、温度、动作。", "像给作品装眼睛、耳朵和皮肤。", "互动装置、保护监测、观众反馈、环境数据艺术。", "Arduino / ESP32 / Raspberry Pi", "先做距离触发灯光。"),
        ("单片机\nMicrocontroller", "控制灯、电机、按钮和传感器的小电脑。", "像装置里的小管家，负责接收信号和发命令。", "灯光装置、按钮互动、机械运动、电子展签。", "Arduino / ESP32", "先让按钮控制一颗灯。"),
        ("串口通信\nSerial Communication", "让电脑和硬件互相传消息。", "像电脑和装置之间打电话。", "电脑视觉控制灯带、传感器控制投影、现场控制台。", "Arduino Serial / Python / TouchDesigner", "先把传感器数值显示到电脑。"),
        ("计算机视觉\nComputer Vision", "让电脑看懂图片和视频。", "像让电脑学会看边缘、人、手势、颜色和运动。", "自动裁图、人流统计、手势控制、作品相似搜索。", "OpenCV / MediaPipe", "先做作品自动裁边工具。"),
        ("图像处理\nImage Processing", "对图片做裁切、校色、降噪、识别和批处理。", "像一个不会累的修图助理。", "作品拍摄校正、水印、压缩、颜色提取。", "Python / OpenCV / PIL", "先做作品色彩拆解器。"),
        ("数字制造\nDigital Fabrication", "把数字文件变成实体物。", "像从电脑图纸走到激光切割、3D 打印和 CNC。", "展架、灯箱、首饰、雕塑、包装盒。", "Blender / Rhino / Grasshopper / CNC", "先做参数化展架。"),
        ("参数化设计\nParametric Design", "用参数控制形状和结构。", "像做一个可调模具，改尺寸就自动变化。", "首饰、灯罩、展架、浮雕、墙面模块。", "Grasshopper / Processing", "先做可变尺寸画框。"),
        ("自动化表格\nSpreadsheet Automation", "让表格自动算钱、排期和状态。", "像把脑子里的杂事交给一个清楚的账本。", "报价、库存、排期、预算、客户跟进。", "Excel / Google Sheets / Python", "先做作品+客户+报价总表。"),
        ("CRM 客户管理\nClient CRM", "记录客户、偏好、预算和跟进。", "像一本不会忘事的客户关系本。", "收藏客户、品牌合作、课程报名、复购提醒。", "Airtable / Notion / Sheets", "先做艺术客户 CRM。"),
        ("内容自动化\nContent Automation", "自动生成、改写、裁切和排期内容。", "像一个社媒助理，帮你做重复发布工作。", "海报适配、短视频字幕、文案改写、发布日历。", "Canva / Python / APIs", "先做社媒发布日历。"),
        ("视频合成\nVideo Assembly", "把图片、文字、音乐和字幕合成视频。", "像自动剪一个作品介绍短片。", "展览预告、过程视频、动态海报、课程片头。", "HyperFrames / FFmpeg / Premiere", "先做作品过程视频合成。"),
        ("数据库\nDatabase", "把作品、客户、文件和状态结构化存起来。", "像一个可以搜索、筛选、关联的资料柜。", "作品档案、库存、电商、展览管理、知识库。", "SQLite / Supabase / Airtable", "先做个人作品档案库。"),
        ("搜索与标签\nSearch and Tags", "给资料加标签并快速找到。", "像给每张图、每份文件贴清楚标签。", "参考图检索、作品搜索、文献索引、代码片段库。", "Full-text Search / Vector Search", "先做视觉参考检索器。"),
        ("AI 知识库\nRAG Assistant", "让 AI 根据你的资料回答问题。", "像把你的简历、作品说明、访谈喂给一个助手。", "艺术家问答、申请材料、客户回复、策展研究。", "LLM / Embeddings / Vector DB", "先做艺术家知识库助手。"),
        ("中英双语\nBilingual Workflow", "中文创作和英文发布之间稳定转换。", "像给你配一个懂艺术术语的翻译编辑。", "作品说明、网站、申请、提示词、海外客户沟通。", "Glossary / Translation QA", "先做中英艺术术语库。"),
        ("版权与授权\nCopyright Licensing", "管理作品能不能用、怎么用、用多久。", "像给作品加一份使用说明书和边界。", "授权合同、水印、版权声明、侵权监测。", "Contracts / Metadata / Watermark", "先做作品授权检查表。"),
        ("版本管理\nVersion Control", "记录文件和代码每次改了什么。", "像创作过程的时间机器，可以回到旧版本。", "网站代码、文案、合同、技术作品、团队协作。", "Git / GitHub", "先把作品集网站放到 GitHub。"),
        ("自动部署\nDeployment", "把本地作品发布到网上。", "像把做好的展览挂到线上空间。", "作品集、互动网页、项目文档、客户预览。", "Vercel / GitHub Actions", "先做自动部署作品集。"),
        ("测试与检查\nTesting QA", "上线前自动检查有没有坏。", "像展览开幕前巡场，看看灯、屏幕、链接都对不对。", "网页截图检查、表单检查、链接检查、视觉回归。", "Playwright / GitHub Actions", "先做项目截图自动生成。"),
        ("文件归档\nArchive System", "把作品、合同、照片和资料长期整理好。", "像一个有目录、有编号、有备份的工作室档案室。", "作品档案、展览归档、证书整理、文献检索。", "Drive / Database / OCR", "先做艺术项目档案包。"),
        ("隐私与安全\nPrivacy Security", "保护客户资料、高清图和未公开作品。", "像给资料加门锁、权限和到期时间。", "客户下载页、高清图权限、敏感信息打码。", "Access Control / Encryption", "先做客户文件下载页。"),
        ("区块链证书\nBlockchain Certificate", "用链上记录证明数字作品信息。", "像给数字作品做可查的编号和历史记录。", "数字证书、限量版、出处记录、会员权益。", "Ethereum / IPFS / Smart Contract", "先做数字作品证书。"),
        ("工程文档\nTechnical Documentation", "把项目怎么做、怎么运行、怎么维护写清楚。", "像给未来的自己留说明书。", "README、插件手册、安装说明、项目 SOP。", "Markdown / GitHub Wiki", "先给每个项目写 README。"),
    ]


def source_rows():
    return [
        ("p5.js", "面向艺术家、设计师、教育者和初学者的创意编程入门工具。", "https://p5js.org/"),
        ("Three.js", "把浏览器 3D 里的场景、相机、灯光、材质等封装起来，适合网页 3D。", "https://threejs.org/manual/en/fundamentals.html"),
        ("D3.js", "用于定制数据可视化的 JavaScript 库。", "https://d3js.org/"),
        ("Blender", "免费开源 3D 创作套件，可做建模、动画、渲染、视频、模拟等。", "https://docs.blender.org/manual/en/4.3/getting_started/about/index.html"),
        ("OpenCV", "大型计算机视觉库，覆盖图像/视频处理、检测和深度学习模块。", "https://opencv.org/"),
        ("Next.js", "用于构建全栈 Web 应用的 React 框架。", "https://nextjs.org/docs"),
        ("React", "用于渲染用户界面的 JavaScript 库，页面可拆成组件。", "https://react.dev/learn/describing-the-ui"),
        ("Unity", "实时 3D 开发平台，可做 2D、3D、VR、AR、模拟等应用。", "https://unity.com/"),
        ("Arduino", "开源电子原型平台，适合连接传感器、灯、电机和按钮。", "https://www.arduino.cc/"),
        ("TouchDesigner", "常用于实时互动媒体、音画、投影、装置和演出视觉。", "https://derivative.ca/"),
    ]


def add_capability_tables(doc):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.1)
    section.right_margin = Cm(1.1)

    add_heading(doc, "九、技术能力白话解释", 1)
    add_note(
        doc,
        "先看能力，再选项目",
        "前面的 345 个项目很多。真正学习时不要按项目编号硬做，而是先看自己缺哪种能力：想做网页，就看网页基础、React、Next.js；想做互动装置，就看传感器、单片机、串口、TouchDesigner；想做作品包装，就看内容自动化、视频合成、AI 图像和版权发布。",
        fill=AMBER,
    )

    rows = capabilities()
    widths = [3.1, 4.0, 5.0, 4.4, 3.9, 4.9]
    headers = ["能力 / Ability", "一句话解释", "小白比喻", "能帮你做什么", "入门工具", "先做一个"]
    for start in range(0, len(rows), 10):
        if start:
            doc.add_page_break()
        add_heading(doc, f"能力组 {start // 10 + 1}", 2)
        table = doc.add_table(rows=1, cols=6)
        table.autofit = False
        set_table_borders(table)
        for idx, header in enumerate(headers):
            cell = table.rows[0].cells[idx]
            set_cell_shading(cell, BLUE)
            cell_text(cell, header, size=7.6, bold=True, color="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_width(cell, widths[idx])
        repeat_table_header(table.rows[0])
        for offset, row in enumerate(rows[start:start + 10]):
            cells = table.add_row().cells
            for col, value in enumerate(row):
                set_cell_shading(cells[col], "FFFFFF" if offset % 2 == 0 else LIGHT_GRAY)
                cell_text(cells[col], value, size=7.25, bold=(col == 0))
                set_cell_width(cells[col], widths[col])


def add_sources(doc):
    doc.add_page_break()
    add_heading(doc, "十、继续搜索补充来源", 1)
    add_para(doc, "这些来源用于补充“能力白话解释”的定义和边界。文档里的解释已经改写成艺术行业能听懂的说法。")
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    set_table_borders(table)
    widths = [4.0, 13.0, 9.0]
    for idx, header in enumerate(("来源", "我提取的重点", "链接")):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, BLUE)
        cell_text(cell, header, size=7.8, bold=True, color="FFFFFF")
        set_cell_width(cell, widths[idx])
    repeat_table_header(table.rows[0])
    for idx, row in enumerate(source_rows()):
        cells = table.add_row().cells
        for col, value in enumerate(row):
            set_cell_shading(cells[col], "FFFFFF" if idx % 2 == 0 else LIGHT_GRAY)
            cell_text(cells[col], value, size=7.4, bold=(col == 0))
            set_cell_width(cells[col], widths[col])


def update_footer(doc):
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.text = ""
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = footer.add_run("艺术行业效率插件建议 | 白话版 + 300+ 项目库 + 能力解释")
        set_font(r, size=8, color="78808C")


def main():
    global DOCX_PATH
    if not DOCX_PATH.exists():
        matches = list(OUTPUT_DIR.glob("*.docx"))
        if len(matches) == 1:
            DOCX_PATH = matches[0]
    doc = Document(DOCX_PATH)
    add_capability_tables(doc)
    add_sources(doc)
    update_footer(doc)
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
