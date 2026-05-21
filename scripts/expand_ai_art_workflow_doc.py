from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


SRC_PATH = r"C:\Users\jian\Desktop\AI艺术出图项目体系白话总结_5人工作组.docx"
OUT_PATH = r"C:\Users\jian\Documents\New project\output\AI艺术出图项目体系白话总结_5人工作组_扩容版.docx"

ACCENT = "1F4E5F"
ACCENT_DARK = "173A46"
ACCENT_LIGHT = "EAF3F5"
MUTED = "6B7280"
GRID = "CBD5E1"
TEXT = "111827"
SOFT = "F8FAFC"


def font_run(run, size=10, color=TEXT, bold=False):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=GRID, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=115, start=105, bottom=115, end=105):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_pct=100):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_pct * 50))
    tbl_w.set(qn("w:type"), "pct")


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_cell(cell, bold=False, color=TEXT, size=8.6, align=None):
    for paragraph in cell.paragraphs:
        if align is not None:
            paragraph.alignment = align
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.12
        for run in paragraph.runs:
            font_run(run, size=size, color=color, bold=bold)


def add_hyperlink(paragraph, text, url):
    r_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    for tag, val in (("w:color", "0563C1"),):
        node = OxmlElement(tag)
        node.set(qn("w:val"), val)
        r_pr.append(node)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Microsoft YaHei")
    fonts.set(qn("w:hAnsi"), "Microsoft YaHei")
    fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r_pr.append(fonts)
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        font_run(run, size=17 if level == 1 else 13 if level == 2 else 11.5, color=ACCENT_DARK if level != 2 else ACCENT, bold=True)
    p.paragraph_format.keep_with_next = True
    return p


def add_para(doc, text="", bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        font_run(r, size=10.3, bold=True)
        r = p.add_run(text[len(bold_prefix):])
        font_run(r, size=10.3)
    else:
        r = p.add_run(text)
        font_run(r, size=10.3)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.55)
        p.paragraph_format.first_line_indent = Cm(-0.22)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run("• " + item)
        font_run(r, size=9.8)


def add_numbered(doc, items):
    for idx, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.65)
        p.paragraph_format.first_line_indent = Cm(-0.28)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f"{idx}. {item}")
        font_run(r, size=9.8)


def add_callout(doc, title, body, fill=ACCENT_LIGHT):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    set_table_width(table, 100)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, color="B7CDD4")
    set_cell_margins(cell, top=150, bottom=150, start=180, end=180)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    font_run(r, size=10.7, color=ACCENT_DARK, bold=True)
    p = cell.add_paragraph()
    r = p.add_run(body)
    font_run(r, size=9.8)
    p.paragraph_format.space_after = Pt(0)


def add_table(doc, headers, rows, font_size=8.3):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    set_table_width(table, 100)
    hdr = table.rows[0]
    set_repeat_header(hdr)
    for i, header in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = header
        set_cell_shading(cell, ACCENT)
        set_cell_border(cell, color=ACCENT_DARK)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        style_cell(cell, bold=True, color="FFFFFF", size=8.8, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cell = cells[i]
            cell.text = text
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            style_cell(cell, size=font_size, align=WD_ALIGN_PARAGRAPH.LEFT if len(str(text)) > 4 else WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_link_table(doc, rows):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.autofit = True
    set_table_width(table, 100)
    for i, header in enumerate(["类别", "项目/用途", "官方或常用链接"]):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, ACCENT)
        set_cell_border(cell, color=ACCENT_DARK)
        set_cell_margins(cell)
        style_cell(cell, bold=True, color="FFFFFF", size=8.8, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_repeat_header(table.rows[0])
    for category, name, url in rows:
        cells = table.add_row().cells
        cells[0].text = category
        cells[1].text = name
        p = cells[2].paragraphs[0]
        add_hyperlink(p, url, url)
        for i, cell in enumerate(cells):
            set_cell_border(cell)
            set_cell_margins(cell, top=110, bottom=110, start=90, end=90)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            style_cell(cell, size=7.8 if i == 2 else 8.3, align=WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT)


def expand():
    doc = Document(SRC_PATH)
    for section in doc.sections:
        section.left_margin = Cm(1.75)
        section.right_margin = Cm(1.75)
        section.top_margin = Cm(1.65)
        section.bottom_margin = Cm(1.6)

    doc.add_page_break()
    add_heading(doc, "13. 搜索扩容说明：这次新增了什么", 1)
    add_para(
        doc,
        "本次扩容基于 2026-05-16 的公开资料搜索，重点补充了新模型、新界面、云端工具、搜索方法、硬件部署、合规风险和 5 人小组的日常管理。前 12 节适合快速入门；后续章节更适合拿来做团队培训、项目 SOP 和内部知识库。",
    )
    add_callout(
        doc,
        "扩容后的核心判断",
        "出图团队不要只围着某一个软件转。真正稳定的体系应该同时有：本地可控生产线、云端快速试稿能力、模型资产库、工作流模板库、版权合规台账和质检交付标准。",
    )

    add_heading(doc, "14. 新增生态总览：本地、云端、混合三条线", 1)
    rows = [
        ("本地开源线", "ComfyUI、WebUI、Forge、SwarmUI、InvokeAI、Stability Matrix", "可控、可复现、成本可控，适合团队沉淀流程。", "需要显卡、安装维护、模型许可管理。"),
        ("云端商业线", "Midjourney、Adobe Firefly、Ideogram、Runway", "上手快、效果强、适合快速提案和视觉方向探索。", "成本按账号/额度走，参数可控性和批量复现弱一些。"),
        ("混合生产线", "云端定方向 + 本地复现 + 后期交付", "既能快速找风格，又能把正式资产沉淀到团队内部。", "需要统一命名、记录来源、区分试稿和正式稿。"),
    ]
    add_table(doc, ["路线", "代表工具", "优势", "注意点"], rows)

    add_heading(doc, "15. 模型家族扩展：不只 SD", 1)
    rows = [
        ("SD 1.5", "老牌、轻量、资料最多", "低配置练习、二次元、插件生态、老 LoRA。", "画面上限和文字能力弱，商业新项目不建议只靠它。"),
        ("SDXL", "比 SD 1.5 更强的大模型", "海报、插画、真人、产品概念图。", "需要更大显存，LoRA 和 ControlNet 要选 SDXL 版本。"),
        ("Stable Diffusion 3.5", "Stability AI 后续开放模型系列", "需要更强文字理解和复杂提示词时考虑。", "先看社区支持度、许可证和当前工具兼容性。"),
        ("FLUX.1", "Black Forest Labs 的强力图像模型", "提示词理解、质感、复杂场景表现好。", "dev/schnell/pro 许可证不同，商业前必须核对。"),
        ("Qwen-Image", "偏中文和复杂文字渲染的图像模型", "中文海报、带字设计、编辑类任务。", "生态还在快速变化，先做小规模测试。"),
        ("HunyuanImage", "腾讯混元图像模型线", "中文语义、图像理解、复杂任务。", "硬件和部署要求较高，适合技术组评估。"),
        ("Wan / HunyuanVideo / Runway", "视频和动态图方向", "图生视频、短片、动态广告、分镜预览。", "视频稳定性、成本和审核时间要单独排期。"),
    ]
    add_table(doc, ["模型/方向", "白话定位", "适合场景", "团队注意点"], rows, font_size=8.1)

    doc.add_page_break()
    add_heading(doc, "16. 工具扩展表：新增值得关注的项目", 1)
    rows = [
        ("WebUI Forge", "A1111 的增强/分支版", "低显存优化、更多新模型尝试、保留 WebUI 操作习惯。", "进阶用户、旧 WebUI 用户迁移"),
        ("Stability Matrix", "SD 工具安装管理器", "统一安装 ComfyUI/WebUI/InvokeAI 等，减少环境折腾。", "团队装机、培训机器"),
        ("SwarmUI", "更现代的模块化 WebUI", "把强工具做得更易用，也能接 ComfyUI 后端。", "批量出图、多用户体验"),
        ("ComfyUI Manager", "ComfyUI 节点管理器", "安装、更新、查找缺失节点，团队必须会用。", "流程工程师"),
        ("Krita AI Diffusion", "Krita 里的 AI 绘画插件", "适合画师边画边局部重绘，不离开绘画软件。", "插画师、原画师"),
        ("Adobe Firefly", "Adobe 生态生成工具", "和 Photoshop/Illustrator/Express 配合，商业设计更稳。", "设计组、品牌组"),
        ("Midjourney", "云端高质量出图工具", "快速找视觉方向、风格探索、提案图。", "创意前期、灵感探索"),
        ("Ideogram", "擅长文字和设计的云端工具", "Logo、海报字、英文/中文文字排版试稿。", "营销设计"),
        ("Runway", "图像和视频生成平台", "图生视频、广告短片、分镜动态化。", "短视频组、动效组"),
    ]
    add_table(doc, ["工具", "它是什么", "能解决什么", "适合谁"], rows, font_size=8.0)

    add_heading(doc, "17. 团队搜索方法：怎么找资料不踩坑", 1)
    add_callout(
        doc,
        "搜索原则",
        "优先看官方文档、GitHub、Hugging Face 模型卡和项目 README；其次看社区教程；最后才看搬运文章。每次下载模型或节点前，都要先看许可证、更新时间、issue 状态和评论反馈。",
        fill=SOFT,
    )
    rows = [
        ("找官方项目", "项目名 + official / GitHub / docs", "ComfyUI official docs、FLUX.1 GitHub、Qwen-Image Hugging Face", "优先保存官网和 GitHub 链接"),
        ("找工作流", "任务 + ComfyUI workflow + model", "product photo ComfyUI workflow、ControlNet pose workflow", "先在测试环境跑，不直接放生产环境"),
        ("找模型", "模型名 + license + Hugging Face / Civitai", "FLUX.1 dev license、SDXL commercial use", "商业项目必须记录许可证截图或链接"),
        ("找报错", "报错原文 + 项目名 + GitHub issue", "ComfyUI missing node import failed", "优先看已关闭 issue 和官方讨论"),
        ("找教程", "工具名 + beginner guide / 中文教程 / best practices", "ComfyUI beginner guide、Kohya LoRA training", "教程只能参考，最终以官方文档为准"),
    ]
    add_table(doc, ["搜索目标", "推荐搜索词", "例子", "记录要求"], rows, font_size=8.0)

    add_heading(doc, "18. 模型和插件入库标准", 1)
    add_para(doc, "团队不要看到模型就下载、看到节点就安装。入库前做一个 10 分钟检查，可以减少 80% 的后续混乱。")
    add_numbered(
        doc,
        [
            "确认来源：GitHub、Hugging Face、Civitai、官方站点或可信作者。",
            "确认许可证：是否允许商业、是否限制再分发、是否只能研究使用。",
            "确认兼容性：SD1.5、SDXL、SD3.5、FLUX、Qwen 等模型体系不能混用。",
            "确认安全性：自定义节点会执行 Python 代码，不要在生产机随便安装陌生节点。",
            "确认效果：用固定测试提示词跑 6 张图，记录优点和缺陷。",
            "确认版本：记录下载日期、版本号、来源链接、文件名和哈希值。",
        ],
    )
    rows = [
        ("模型名", "写清楚模型/LoRA/节点的显示名。", "FLUX.1-schnell、brand_style_v01"),
        ("真实文件名", "保留下载后的完整文件名，方便排查版本。", "flux1-schnell.safetensors"),
        ("来源链接", "写 GitHub、Hugging Face、Civitai 或内部资产库链接。", "https://huggingface.co/..."),
        ("许可证", "写 Apache-2.0、MIT、非商用、内部自有等。", "Apache-2.0 / internal-only"),
        ("适用工具", "说明能在哪套流程里用，避免模型体系混用。", "SDXL + ComfyUI / FLUX + SwarmUI"),
        ("负责人", "谁入库，后续谁负责更新和下架。", "模型素材管理员"),
        ("状态", "测试中、生产可用、弃用、仅学习。", "生产可用"),
    ]
    add_table(doc, ["字段", "填写说明", "示例"], rows, font_size=8.2)

    add_heading(doc, "19. 5 人工作组扩展：日常节奏和交接规则", 1)
    rows = [
        ("每天 10 分钟", "同步今天要出什么图、谁卡在哪、是否需要新模型/节点。", "项目负责人", "当日任务清单"),
        ("每个项目开始前", "确认用途、尺寸、风格、禁区、参考图、版权要求。", "项目负责人 + 风格设计师", "需求单"),
        ("每次批量生成后", "保留参数、seed、模型版本、workflow、候选图。", "流程工程师", "过程包"),
        ("每周 1 次", "沉淀好用 prompt、workflow、模型、失败案例。", "全员", "周复盘文档"),
        ("每月 1 次", "清理废弃模型、更新工具、检查许可证和备份。", "模型管理员", "资产库报告"),
    ]
    add_table(doc, ["节奏", "做什么", "负责人", "产出"], rows, font_size=8.1)

    add_heading(doc, "20. RACI 分工表：谁负责、谁拍板、谁配合", 1)
    add_para(doc, "R=负责执行，A=最终拍板，C=需要咨询，I=需要同步。这个表能防止 5 个人都在改图、但没人对最终结果负责。")
    rows = [
        ("需求确认", "A/R", "C", "I", "C", "I"),
        ("风格探索", "A", "R", "C", "C", "I"),
        ("工作流搭建", "C", "C", "A/R", "C", "I"),
        ("模型/LoRA 入库", "I", "C", "C", "A/R", "I"),
        ("批量生成", "I", "C", "A/R", "C", "I"),
        ("后期修图", "C", "C", "I", "I", "A/R"),
        ("最终验收", "A/R", "C", "I", "C", "R"),
        ("版权合规记录", "A", "I", "I", "R", "C"),
    ]
    add_table(doc, ["任务", "负责人/美术总监", "提示词与风格", "流程工程师", "模型素材管理员", "后期质检交付"], rows, font_size=7.8)

    add_heading(doc, "21. 硬件与部署建议", 1)
    rows = [
        ("入门练习机", "8GB-12GB 显存", "SD1.5、部分 SDXL、轻量 LoRA 测试", "适合新人学习，不建议承担正式批量任务。"),
        ("主力出图机", "16GB-24GB 显存", "SDXL、ComfyUI、ControlNet、局部重绘、放大", "5 人小组建议至少 1 台主力机器。"),
        ("高配/服务器", "24GB-48GB+ 显存", "FLUX、视频、批量生产、大模型实验", "适合流程工程师和模型管理员共用。"),
        ("云 GPU", "按小时租用", "临时大任务、训练 LoRA、视频生成", "用完关机，成本要写进项目预算。"),
        ("云端商业工具", "账号/额度", "快速试方向、提案图、视频初稿", "不要把客户敏感素材随便上传。"),
    ]
    add_table(doc, ["配置路线", "大致规格", "适合任务", "团队建议"], rows, font_size=8.2)

    add_heading(doc, "22. 质量验收清单：出图不是看着好看就完事", 1)
    rows = [
        ("画面", "构图、主体、光影、色彩、风格是否符合需求。", "负责人/美术总监"),
        ("人物", "脸、手、眼睛、牙齿、肢体、服装结构是否异常。", "后期质检"),
        ("文字", "海报字、Logo、商品标签是否拼错或变形。", "后期质检 + 风格设计师"),
        ("品牌", "颜色、字体、气质、产品形态是否偏离品牌。", "项目负责人"),
        ("技术", "分辨率、比例、边缘、噪点、压缩质量、透明背景。", "后期质检"),
        ("可追溯", "是否保存 prompt、seed、模型、LoRA、工作流和版本号。", "流程工程师"),
        ("合规", "模型许可、人物肖像、商标、敏感内容、客户素材授权。", "模型素材管理员"),
    ]
    add_table(doc, ["检查项", "检查内容", "主要负责人"], rows, font_size=8.3)

    doc.add_page_break()
    add_heading(doc, "23. Prompt 模板库：给新人直接套用", 1)
    add_callout(
        doc,
        "通用 Prompt 结构",
        "主体 + 场景 + 风格 + 构图 + 光线 + 镜头/材质 + 质量要求 + 禁止项。中文写需求，英文补风格和摄影词，都可以；关键是结构清楚、不要一整段乱堆词。",
    )
    rows = [
        ("产品图", "主体：一款白色无线耳机；场景：干净桌面；风格：高端电商摄影；光线：柔和棚拍；要求：清晰、真实、无多余文字。"),
        ("角色图", "主体：年轻女性探险家；服装：轻便户外装备；场景：雨后森林；风格：电影感写实；要求：人物一致、脸部自然。"),
        ("海报图", "主体：夏季新品饮料；构图：产品居中、周围水果飞溅；风格：清爽商业广告；文字区域：上方留白。"),
        ("二次元", "主体：原创魔法少女；风格：精致日系插画；构图：半身像；光线：柔和逆光；要求：线条干净、颜色明亮。"),
        ("负面词", "低清晰度、变形手、错误文字、多余肢体、脸部崩坏、商标水印、模糊、压缩噪点。"),
    ]
    add_table(doc, ["类型", "模板示例"], rows, font_size=8.2)

    add_heading(doc, "24. 合规和版权：小团队最容易忽略的地方", 1)
    add_bullets(
        doc,
        [
            "商业项目不要默认所有模型都能商用。不同 Checkpoint、LoRA、平台输出都有自己的规则。",
            "Civitai、Hugging Face 上的模型要看模型卡和许可证，不要只看下载量。",
            "涉及真实人物、明星、员工、客户、未成年人、医疗金融等场景，要单独走授权和审核。",
            "客户素材、产品图、Logo、内部资料不要随便上传到云端平台，除非合同和平台条款允许。",
            "生成图用于广告、包装、游戏、出版前，建议保留模型来源、提示词、版本和人工修改记录。",
            "训练 LoRA 时，数据集来源要干净；能用自有图、授权图、客户提供图，就不要用网上乱抓的图。",
        ],
    )

    add_heading(doc, "25. 30/60/90 天落地路线图", 1)
    rows = [
        ("0-30 天", "完成工具安装、基础培训、模型库初版、5 人分工、命名规范。", "能完成小型项目交付。"),
        ("31-60 天", "沉淀 5-10 套 ComfyUI 工作流、建立 prompt 模板库、做质检表。", "同类任务能稳定复用。"),
        ("61-90 天", "训练 1-3 个自有 LoRA、建立合规台账、完成批量生产流程。", "形成团队自己的风格资产。"),
        ("90 天后", "根据业务选择云 GPU、内部服务器、自动化脚本、视频工作流。", "从“出图小组”升级为“视觉资产生产线”。"),
    ]
    add_table(doc, ["阶段", "重点动作", "目标"], rows, font_size=8.3)

    add_heading(doc, "26. 推荐知识库结构", 1)
    rows = [
        ("A_工具安装", "安装包、版本、启动方式、常见报错。", "流程工程师维护"),
        ("B_模型资产", "Checkpoint、LoRA、VAE、ControlNet、许可证。", "模型素材管理员维护"),
        ("C_工作流库", "ComfyUI JSON、节点截图、示例输入输出。", "流程工程师维护"),
        ("D_Prompt库", "不同项目类型的提示词模板和反例。", "提示词与风格设计师维护"),
        ("E_项目交付", "需求单、候选图、最终图、验收表。", "项目负责人和交付人员维护"),
        ("F_复盘案例", "好案例、坏案例、参数对比、失败原因。", "全员每周更新"),
    ]
    add_table(doc, ["目录", "放什么", "负责人"], rows, font_size=8.3)

    doc.add_page_break()
    add_heading(doc, "27. 搜索补充链接清单", 1)
    link_rows = [
        ("ComfyUI", "官方文档：节点式生成式 AI 应用和推理引擎", "https://docs.comfy.org/index"),
        ("ComfyUI", "官方 GitHub", "https://github.com/Comfy-Org/ComfyUI"),
        ("ComfyUI", "ComfyUI Manager 文档", "https://docs.comfy.org/manager/pack-management"),
        ("ComfyUI", "ComfyUI Manager GitHub", "https://github.com/Comfy-Org/ComfyUI-Manager"),
        ("WebUI", "AUTOMATIC1111 Stable Diffusion WebUI", "https://github.com/AUTOMATIC1111/stable-diffusion-webui"),
        ("WebUI", "Stable Diffusion WebUI Forge", "https://github.com/lllyasviel/stable-diffusion-webui-forge"),
        ("管理器", "Stability Matrix", "https://github.com/LykosAI/StabilityMatrix"),
        ("管理器", "SwarmUI", "https://github.com/mcmonkeyprojects/SwarmUI"),
        ("绘画插件", "Krita AI Diffusion", "https://github.com/Acly/krita-ai-diffusion"),
        ("模型", "Stable Diffusion 3.5 介绍", "https://stability.ai/news/introducing-stable-diffusion-3-5"),
        ("模型", "Stability AI SD3.5 GitHub", "https://github.com/Stability-AI/sd3.5"),
        ("模型", "FLUX.1 官方推理库", "https://github.com/black-forest-labs/flux"),
        ("模型", "Qwen-Image GitHub", "https://github.com/QwenLM/Qwen-Image"),
        ("模型", "HunyuanImage 3.0 GitHub", "https://github.com/Tencent-Hunyuan/HunyuanImage-3.0"),
        ("视频", "Wan2.1 GitHub", "https://github.com/Wan-Video/Wan2.1"),
        ("平台", "Hugging Face Model Cards 文档", "https://huggingface.co/docs/hub/model-cards"),
        ("平台", "Hugging Face Licenses 文档", "https://huggingface.co/docs/hub/repositories-licenses"),
        ("平台", "Civitai", "https://civitai.com/"),
        ("商业云", "Adobe Firefly", "https://www.adobe.com/products/firefly.html"),
        ("商业云", "Midjourney Terms of Service", "https://docs.midjourney.com/hc/en-us/articles/32083055291277-Terms-of-Service"),
        ("商业云", "Ideogram", "https://ideogram.ai/features/"),
        ("视频云", "Runway Gen-4 帮助文档", "https://help.runwayml.com/hc/en-us/articles/37327109429011-Creating-with-Gen-4-Video"),
    ]
    add_heading(doc, "27.1 开源与本地工具", 2)
    add_link_table(doc, link_rows[:15])
    doc.add_page_break()
    add_heading(doc, "27.2 平台、商业云与视频", 2)
    add_link_table(doc, link_rows[15:])

    add_heading(doc, "28. 扩容版最后建议", 1)
    add_callout(
        doc,
        "一句话落地",
        "五人小组最适合采用“云端找方向，本地做生产，后期做交付，知识库做沉淀”的混合体系。先把流程跑顺，再逐步增加模型、节点、训练和自动化，不要一开始就把团队拖进无止境的安装和调参里。",
    )

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    expand()
