from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


DOCX_PATH = Path(r"C:\Users\jian\Documents\New project\output\艺术行业效率插件建议.docx")


BLUE = "204370"
LIGHT_BLUE = "EEF5FF"
LIGHT_GRAY = "F7FAFD"
AMBER = "FFF6E6"
BORDER = "CBD5E1"


def set_east_asian_font(run, size=None, bold=None, color=None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, cm):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn("w:" + edge))
        if node is None:
            node = OxmlElement("w:" + edge)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def cell_text(cell, text, size=7.2, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(text)
    set_east_asian_font(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    set_east_asian_font(run, size=16 if level == 1 else 12, bold=True, color=BLUE if level == 1 else "435B7B")
    return p


def add_para(doc, text, size=9.2):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_east_asian_font(run, size=size)
    return p


def add_note(doc, title, body, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    set_east_asian_font(r, size=10, bold=True, color=BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.15
    r2 = p2.add_run(body)
    set_east_asian_font(r2, size=8.8)
    doc.add_paragraph()


def build_projects():
    categories = [
        ("创意编程基础", "Creative Coding Basics", "p5.js / Processing / JavaScript", [
            ("随机点线海报生成器", "Random Line Poster Generator", "训练随机数、循环、颜色和导出图片", "random poster generator p5.js generative art", "让代码按规则画线、点和色块，像做一台自动海报机器。"),
            ("参数化纹样画布", "Parametric Pattern Canvas", "训练变量、滑块控制和图案复用", "parametric pattern p5.js sliders", "把图案拆成可调参数，改一个数字就能换一批视觉方案。"),
            ("艺术家签名动画", "Animated Artist Signature", "训练路径、补间动画和 SVG 绘制", "animated signature svg path javascript", "把签名或笔触变成按时间出现的动画。"),
            ("互动色卡生成器", "Interactive Palette Generator", "训练色彩算法和 UI 控件", "color palette generator javascript", "输入作品主色，自动生成适合海报、网页和作品集的配色。"),
            ("作品图像像素化工具", "Artwork Pixelation Tool", "训练图像采样和像素处理", "image pixelation canvas javascript", "把作品拆成像素块，适合做复古视觉或材料分析。"),
            ("网格排版实验室", "Grid Layout Lab", "训练响应式网格和版式规则", "generative grid layout canvas", "用网格控制作品、文字、留白，快速试作品集版式。"),
            ("渐变噪声背景器", "Noise Gradient Background Maker", "训练噪声函数和渐变算法", "perlin noise gradient p5.js", "用噪声生成自然的云雾、纸张、光影背景。"),
            ("鼠标笔刷画板", "Mouse Brush Sketchpad", "训练事件监听和笔刷逻辑", "creative coding brush canvas", "鼠标或触控变成自定义笔刷，可做线上互动涂鸦。"),
            ("字体变形海报", "Deformed Type Poster", "训练文字绘制和几何变形", "typography deformation canvas", "把字体按声音、鼠标或随机值扭曲成实验海报。"),
            ("时间驱动图案钟", "Time Driven Pattern Clock", "训练时间数据和动态视觉", "generative clock p5.js", "让当前时间控制图形变化，适合展览入口屏。"),
            ("作品色彩拆解器", "Artwork Color Extractor", "训练图像读取和聚类思路", "extract colors from image javascript", "从作品里提取主色，做成展览视觉系统。"),
            ("可打印明信片生成器", "Printable Postcard Generator", "训练版面尺寸和导出流程", "printable postcard generator canvas", "一键生成适合印刷尺寸的作品明信片。"),
            ("键盘声音画布", "Keyboard Sound Canvas", "训练键盘事件和音画反馈", "keyboard visualizer p5 sound", "按键会触发图形和声音，适合课堂演示。"),
            ("循环 GIF 草图机", "Looping GIF Sketch Machine", "训练帧动画和导出", "p5.js export gif animation", "生成可以发社媒的无缝循环动态图。"),
            ("作品拼贴生成器", "Artwork Collage Generator", "训练图片裁切、层级和混合模式", "canvas collage generator", "把多张作品自动拼成封面、海报或社媒图。"),
        ]),
        ("生成艺术", "Generative Art", "p5.js / GLSL / SVG / Python", [
            ("粒子流动花园", "Particle Flow Garden", "训练粒子系统和向量场", "particle flow field generative art", "很多小点按看不见的风场移动，形成有生命感的线条。"),
            ("分形树生长器", "Fractal Tree Grower", "训练递归和自然形态", "fractal tree generative art", "用重复规则长出树枝，理解复杂图形怎样由简单规则生成。"),
            ("噪声地形版画", "Noise Terrain Print", "训练高度图和等高线", "perlin noise terrain contour art", "用噪声生成山脉、海浪或地形纹理。"),
            ("规则瓷砖图案", "Rule Based Tile Patterns", "训练平铺、旋转和对称", "wang tiles generative pattern", "把小图块按规则拼成不会重复的装饰图案。"),
            ("生成式水墨", "Generative Ink Wash", "训练透明度、扩散和图层", "generative ink wash p5.js", "模拟墨迹晕染，适合数字水墨方向。"),
            ("算法花朵图谱", "Algorithmic Flower Atlas", "训练极坐标和参数化形态", "algorithmic flowers generative art", "用数学控制花瓣数量、弯曲和颜色。"),
            ("随机行走线稿", "Random Walk Line Drawing", "训练随机游走和路径记录", "random walk line art", "让一条线自己走路，形成抽象线稿。"),
            ("生成式纹身草图", "Generative Tattoo Sketches", "训练 SVG 曲线和对称设计", "generative tattoo svg", "自动生成可再手工修改的纹身图案草稿。"),
            ("城市纹理生成器", "Urban Texture Generator", "训练网格噪声和材质合成", "procedural urban texture", "生成砖墙、路面、霓虹等城市视觉纹理。"),
            ("生成式封面系统", "Generative Cover System", "训练模板和批量输出", "generative book cover system", "同一套规则批量生成展册封面或系列海报。"),
            ("纸张纤维模拟", "Paper Fiber Simulator", "训练线段、透明度和材质感", "procedural paper texture", "用代码做出纸张纤维，适合作品展示背景。"),
            ("算法动物纹样", "Algorithmic Animal Patterns", "训练噪声、遮罩和图案映射", "procedural animal pattern", "生成斑马纹、豹纹、羽毛等图案，理解自然纹理规则。"),
            ("生成式书法骨架", "Generative Calligraphy Skeleton", "训练贝塞尔曲线和笔触宽度", "generative calligraphy bezier", "用曲线骨架控制笔触粗细，做实验书法。"),
            ("参数化首饰草图", "Parametric Jewelry Sketches", "训练曲线、对称和导出 SVG", "parametric jewelry design svg", "自动生成耳环、胸针、吊坠的二维草图。"),
            ("生成式系列编号证书", "Generative Edition Certificate", "训练模板、编号和批量生成", "generative edition certificate art", "为限量作品生成带编号的视觉证书。"),
        ]),
        ("数据可视化", "Data Visualization", "D3.js / Python / Observable / SVG", [
            ("作品销售时间线", "Artwork Sales Timeline", "训练时间序列和图表", "art sales timeline d3", "把作品销售、展览、咨询按时间排出来，看趋势。"),
            ("展览观众热力图", "Exhibition Visitor Heatmap", "训练热力图和空间数据", "visitor heatmap visualization", "把观众停留位置转成热力图，优化展陈动线。"),
            ("社媒互动仪表盘", "Social Engagement Dashboard", "训练 API、图表和指标", "social media analytics dashboard artist", "把点赞、收藏、评论和转化放在一个页面看。"),
            ("作品价格散点图", "Artwork Price Scatterplot", "训练散点图和筛选", "art price scatter plot", "看尺寸、年份、媒介和价格之间有没有规律。"),
            ("艺术家关系网络", "Artist Network Graph", "训练网络图和节点关系", "artist network graph d3", "把艺术家、机构、策展人、展览的关系画出来。"),
            ("材料成本瀑布图", "Material Cost Waterfall", "训练成本拆解图", "cost waterfall chart", "把材料、运输、装裱、场地等成本拆清楚。"),
            ("作品库存状态图", "Inventory Status Map", "训练状态分类和看板", "art inventory dashboard", "用颜色标记已售、可售、借展、待拍摄。"),
            ("展览预算树图", "Exhibition Budget Treemap", "训练树图和层级数据", "budget treemap d3", "把预算按大类小类显示，快速看钱花在哪里。"),
            ("作品主题词云", "Artwork Theme Word Cloud", "训练文本统计和可视化", "word cloud artwork statement", "从作品说明里找高频词，看个人风格叙事。"),
            ("收藏客户分层图", "Collector Segmentation Chart", "训练分群和图表", "customer segmentation visualization", "把客户按预算、兴趣、购买频率分层。"),
            ("展览路线桑基图", "Exhibition Journey Sankey", "训练流向图和路径数据", "visitor journey sankey", "看观众从入口到各展区的流动路径。"),
            ("艺术项目甘特图", "Art Project Gantt Chart", "训练排期和依赖关系", "gantt chart exhibition planning", "把制作、运输、布展、开幕、撤展排成一张时间表。"),
            ("城市艺术地图", "City Art Map", "训练地图坐标和筛选", "art gallery map visualization", "把画廊、艺术空间、客户位置放到地图上。"),
            ("作品颜色统计图", "Artwork Color Analytics", "训练图像色彩统计", "image color analysis visualization", "分析一批作品的主色、亮度和色彩倾向。"),
            ("艺术市场新闻雷达", "Art Market News Radar", "训练信息聚合和分类", "art market news dashboard", "把新闻、拍卖、展览信息按主题归类跟踪。"),
        ]),
        ("AI 图像与提示词", "AI Image and Prompt Tools", "Stable Diffusion / ComfyUI / Python / APIs", [
            ("艺术风格提示词库", "Art Style Prompt Library", "训练结构化提示词和标签", "stable diffusion prompt library art style", "把常用风格、材质、构图词整理成可复用数据库。"),
            ("作品风格参考板", "Style Reference Board", "训练图像检索和相似度", "image similarity style board", "上传参考图后自动整理成风格板。"),
            ("提示词 A/B 测试器", "Prompt A/B Tester", "训练批量生成和对比", "prompt a/b testing stable diffusion", "同一主题换不同提示词，比较哪组更接近目标。"),
            ("AI 海报初稿机", "AI Poster Draft Machine", "训练图像生成和版式筛选", "ai poster generator workflow", "先生成多个海报方向，再人工挑选和精修。"),
            ("局部重绘修图助手", "Inpainting Retouch Assistant", "训练蒙版和图像修补", "stable diffusion inpainting workflow", "只改画面局部，比如背景、道具、缺口，不重做整张图。"),
            ("作品高清放大流程", "Artwork Upscale Workflow", "训练超分辨率和批处理", "ai image upscaling workflow", "把低清图放大到适合展示或打印的尺寸。"),
            ("系列图像一致性工具", "Consistent Series Image Tool", "训练种子、LoRA 和风格一致", "consistent character style diffusion", "让一组图保持同一视觉语言。"),
            ("材质参考生成器", "Material Reference Generator", "训练纹理提示词和图像分类", "ai material texture generator", "快速生成陶瓷、金属、布料、纸张等材质参考。"),
            ("策展视觉情绪板", "Curatorial Moodboard Generator", "训练多图生成和标签管理", "ai moodboard generator exhibition", "把策展主题变成可讨论的视觉方向板。"),
            ("作品说明配图机", "Statement Illustration Generator", "训练文本转图和文案配合", "text to image illustration statement", "根据作品说明生成辅助插图，帮客户理解概念。"),
            ("风格迁移实验室", "Style Transfer Lab", "训练神经风格迁移", "neural style transfer art", "把一张图的风格迁移到另一张图，探索视觉语言。"),
            ("AI 展览空间草图", "AI Exhibition Space Sketcher", "训练空间提示词和透视", "ai exhibition space concept art", "快速生成展厅、墙面、灯光和作品摆放草图。"),
            ("提示词翻译校对器", "Prompt Translation Checker", "训练中英术语翻译", "chinese english prompt translation", "把中文创作想法翻成更稳定的英文提示词。"),
            ("参考图版权标注器", "Reference Rights Annotator", "训练元数据和备注系统", "image metadata copyright notes", "给参考图记录来源、用途和版权风险。"),
            ("AI 作品筛选看板", "AI Output Review Board", "训练评分、标签和看板", "ai image review board", "把生成结果按构图、色彩、可用性打分筛选。"),
        ]),
        ("作品集与网站", "Portfolio and Website", "HTML / CSS / React / Next.js / CMS", [
            ("艺术家个人首页", "Artist Homepage", "训练响应式网页和内容结构", "artist portfolio website react", "把作品、简介、展览、联系方式做成一个清楚的网站。"),
            ("可筛选作品集", "Filterable Portfolio", "训练筛选、分类和路由", "filterable portfolio gallery javascript", "观众可以按年份、媒介、系列筛作品。"),
            ("作品详情页模板", "Artwork Detail Page Template", "训练动态页面和数据字段", "artwork detail page cms", "每件作品都有尺寸、材质、价格、说明和高清图。"),
            ("线上展厅页面", "Online Exhibition Room", "训练长页面叙事和图片优化", "online exhibition website", "把展览做成可滚动浏览的线上版本。"),
            ("艺术家简历生成页", "Artist CV Generator Page", "训练表单和静态导出", "artist cv generator web app", "把展览、奖项、收藏录入后自动生成 CV。"),
            ("作品询价表单", "Artwork Inquiry Form", "训练表单提交和邮件提醒", "art inquiry form website", "客户点作品后直接留下购买或合作意向。"),
            ("双语作品集网站", "Bilingual Portfolio Site", "训练 i18n 和内容管理", "bilingual artist portfolio i18n", "中文英文切换，方便国际客户看懂。"),
            ("作品图片懒加载", "Artwork Lazy Loading Gallery", "训练性能优化", "image lazy loading gallery", "图片多的网站也能打开快，不拖慢用户体验。"),
            ("展览新闻博客", "Exhibition News Blog", "训练 Markdown 和 CMS", "artist blog nextjs mdx", "持续发布展览、创作过程和新闻。"),
            ("沉浸式首页动效", "Immersive Landing Animation", "训练滚动动画和 WebGL", "scroll animation artist website", "让首页有动态视觉，但不影响作品阅读。"),
            ("作品搜索页", "Artwork Search Page", "训练全文搜索和标签", "portfolio search tags", "按标题、主题、年份、材料搜索作品。"),
            ("作品集 PDF 导出", "Portfolio PDF Export", "训练网页转 PDF", "portfolio pdf export react", "从网站数据一键导出客户版 PDF。"),
            ("客户私密预览页", "Private Collector Preview", "训练登录链接和权限", "private gallery preview website", "给特定客户看未公开作品或报价。"),
            ("艺术电商小店", "Art Shop Mini Store", "训练购物车和支付流程", "artist ecommerce store", "卖版画、周边、课程或数字下载。"),
            ("网站访问分析面板", "Portfolio Analytics Panel", "训练埋点和数据面板", "portfolio analytics dashboard", "看哪些作品被点得最多，辅助商业判断。"),
        ]),
        ("互动装置", "Interactive Installation", "TouchDesigner / Arduino / Sensors / Python", [
            ("红外感应光墙", "Infrared Reactive Light Wall", "训练传感器输入和灯光控制", "arduino infrared led wall installation", "人靠近时灯光变化，适合入口互动装置。"),
            ("身体动作投影", "Body Motion Projection", "训练姿态识别和实时投影", "mediapipe projection installation", "摄像头识别人动作，投影跟着身体变化。"),
            ("触摸声音雕塑", "Touch Sound Sculpture", "训练触摸传感和音频触发", "capacitive touch sound sculpture", "观众触摸不同位置，触发不同声音。"),
            ("观众距离视觉反馈", "Distance Reactive Visuals", "训练距离传感器和视觉映射", "ultrasonic sensor generative visuals", "离作品越近，画面越细或越亮。"),
            ("多人互动粒子墙", "Multi User Particle Wall", "训练多目标跟踪", "multi user interactive particle wall", "多个观众同时影响墙面粒子运动。"),
            ("呼吸灯装置", "Breathing Light Installation", "训练节奏、PWM 和灯光渐变", "breathing light arduino installation", "灯光像呼吸一样变化，营造空间情绪。"),
            ("声音触发投影花", "Sound Reactive Projection Bloom", "训练音频分析和投影", "sound reactive projection mapping", "声音越强，投影花朵越开放。"),
            ("NFC 作品讲解台", "NFC Artwork Info Stand", "训练 NFC 和内容触发", "nfc artwork information installation", "手机靠近标签就打开作品讲解页面。"),
            ("观众投票灯箱", "Audience Voting Lightbox", "训练按钮输入和状态显示", "interactive voting lightbox arduino", "观众按按钮选择情绪，灯箱统计并显示结果。"),
            ("环境数据装置", "Environmental Data Installation", "训练传感器和数据可视化", "environmental sensor art installation", "温度、湿度、空气数据变成实时艺术画面。"),
            ("触控水波屏", "Touch Ripple Screen", "训练触控坐标和波纹算法", "touch ripple interactive screen", "手指触碰屏幕产生水波视觉。"),
            ("脚步触发地面投影", "Step Triggered Floor Projection", "训练压力传感和投影互动", "pressure sensor floor projection", "观众踩到地面某处，投影出现反应。"),
            ("手势控制灯带", "Gesture Controlled LED Strip", "训练手势识别和串口通信", "gesture controlled led strip", "挥手就能改变灯光颜色和节奏。"),
            ("互动展签屏", "Interactive Label Screen", "训练小屏 UI 和内容管理", "interactive museum label screen", "展签不再只写一段字，可以切换图片、视频、语音。"),
            ("观众影子剧场", "Audience Shadow Theater", "训练轮廓识别和投影合成", "interactive shadow projection", "观众影子被转成图形角色参与画面。"),
        ]),
        ("投影映射与空间", "Projection Mapping and Space", "TouchDesigner / MadMapper / Blender / WebGL", [
            ("小型雕塑投影映射", "Small Sculpture Projection Mapping", "训练 UV、遮罩和空间对齐", "projection mapping sculpture tutorial", "把影像精准投到雕塑表面。"),
            ("建筑立面预演器", "Facade Mapping Previewer", "训练透视校正和预演", "facade projection mapping preview", "先在电脑上模拟建筑投影效果。"),
            ("展厅灯光模拟", "Gallery Lighting Simulator", "训练 3D 场景和灯光", "gallery lighting simulator three.js", "在布展前测试灯光角度和作品阴影。"),
            ("多投影融合模板", "Multi Projector Blend Template", "训练边缘融合和校准", "multi projector edge blending", "多台投影拼接成一整面墙。"),
            ("投影内容排期器", "Projection Content Scheduler", "训练定时播放和媒体管理", "media playback scheduler installation", "按时间自动切换投影内容。"),
            ("墙面作品定位工具", "Wall Artwork Placement Tool", "训练拖拽、比例和测量", "artwork wall placement web app", "在平面图上拖作品，检查尺寸和间距。"),
            ("沉浸式房间草图", "Immersive Room Sketch", "训练 360 视觉和空间叙事", "immersive room projection concept", "设计四面墙影像如何包围观众。"),
            ("投影遮罩编辑器", "Projection Mask Editor", "训练多边形编辑和蒙版", "projection mask editor canvas", "用鼠标画出哪里投、哪里不投。"),
            ("空间声光同步", "Spatial Sound Light Sync", "训练时间码和同步", "timecode sound light sync", "音乐、灯光、投影按同一时间线运行。"),
            ("展览动线 3D 预览", "3D Exhibition Walkthrough", "训练 Three.js 漫游和模型加载", "three.js gallery walkthrough", "在网页里预览观众走进展厅的感觉。"),
            ("投影色彩校正表", "Projection Color Calibration Chart", "训练色彩校正和测试图", "projection color calibration chart", "让投影颜色更接近设计稿。"),
            ("现场媒体服务器面板", "Media Server Control Panel", "训练控制界面和状态监控", "media server control panel touchdesigner", "集中控制播放、暂停、切换、音量和状态。"),
            ("投影文件命名助手", "Projection File Naming Helper", "训练文件规则和批量重命名", "batch rename projection files", "减少现场文件混乱导致播错内容。"),
            ("舞台视觉 cue 表", "Stage Visual Cue Sheet", "训练流程表和触发点", "stage visual cue sheet", "把每个视觉变化和音乐、灯光节点对齐。"),
            ("投影项目报价计算器", "Projection Quote Calculator", "训练表格公式和项目估算", "projection mapping quote calculator", "按设备、天数、人员、内容复杂度自动估价。"),
        ]),
        ("声音与视听", "Audio Visual", "Web Audio / Tone.js / MIDI / TouchDesigner", [
            ("音乐频谱海报", "Audio Spectrum Poster", "训练 FFT 和图像输出", "audio spectrum generative poster", "把一段音乐变成静态海报。"),
            ("声音反应粒子", "Sound Reactive Particles", "训练音频分析和粒子系统", "sound reactive particles p5.js", "声音越强，粒子越密或飞得越快。"),
            ("MIDI 控制视觉台", "MIDI Visual Console", "训练 MIDI 输入和参数映射", "midi controlled visuals", "用旋钮和推子现场控制视觉。"),
            ("环境声音地图", "Ambient Sound Map", "训练录音、地图和标签", "sound map web app", "记录城市声音并放到地图上。"),
            ("自动节拍剪辑器", "Beat Cut Video Tool", "训练节拍检测和视频切分", "beat detection video editing", "根据音乐节拍自动切换画面。"),
            ("语音生成抽象画", "Voice Generated Abstract Art", "训练麦克风输入和波形", "voice reactive art", "观众说话会生成专属抽象图。"),
            ("声音情绪灯光", "Audio Mood Lighting", "训练音频特征和颜色映射", "audio mood lighting installation", "根据音乐情绪改变空间灯光。"),
            ("音画同步短视频模板", "Audio Visual Short Template", "训练时间线和导出", "audio visualizer short video", "把作品图和音乐合成适合社媒的短视频。"),
            ("合成器视觉皮肤", "Synth Visual Skin", "训练 Web Audio 和界面", "web audio synthesizer visualizer", "做一个会发声也会动的网页乐器。"),
            ("现场 VJ 素材库", "Live VJ Clip Library", "训练媒体管理和标签", "vj clip library management", "把现场视觉素材按速度、颜色、主题整理。"),
            ("声音驱动字体", "Audio Driven Typography", "训练文字动画和音频参数", "audio reactive typography", "声音让标题放大、扭曲或抖动。"),
            ("作品声音导览", "Artwork Audio Guide", "训练音频播放和二维码", "museum audio guide web app", "扫码就能听作品讲解。"),
            ("节奏型灯带控制", "Rhythm LED Controller", "训练节拍和 LED 通信", "beat synced led controller", "灯带按音乐节拍闪烁。"),
            ("声音采样拼贴", "Sound Sample Collage", "训练采样、触发和混音", "sound collage web audio", "把不同声音片段拼成互动音景。"),
            ("视听演出计时器", "AV Performance Timer", "训练时间线和 cue 提醒", "av performance timer cue", "演出时提示下一段视觉和声音节点。"),
        ]),
        ("3D 与 WebGL", "3D and WebGL", "Three.js / Blender / WebGPU / GLSL", [
            ("网页 3D 作品展柜", "Web 3D Artwork Cabinet", "训练模型加载和相机控制", "three.js 3d art gallery", "观众在网页里旋转查看 3D 作品。"),
            ("GLTF 作品查看器", "GLTF Artwork Viewer", "训练 3D 文件格式和材质", "gltf viewer three.js", "把 Blender 模型放到网页里展示。"),
            ("材质球实验室", "Material Ball Lab", "训练 PBR 材质和灯光", "three.js pbr material lab", "快速测试金属、玻璃、陶瓷等材质效果。"),
            ("3D 展厅导览", "3D Gallery Tour", "训练碰撞、导航和加载优化", "3d gallery tour webgl", "让客户线上走进虚拟展厅。"),
            ("程序化雕塑生成器", "Procedural Sculpture Generator", "训练几何变形和导出", "procedural sculpture three.js", "用参数生成 3D 雕塑草案。"),
            ("粒子星云空间", "Particle Nebula Space", "训练 GPU 粒子和后期效果", "three.js particle nebula", "做沉浸式网页背景或展览视觉。"),
            ("3D 字体标题", "3D Typography Title", "训练字体转几何体", "three.js 3d text typography", "把展览标题变成立体字。"),
            ("互动材质贴图器", "Interactive Texture Mapper", "训练 UV 和纹理替换", "three.js texture mapping tool", "给 3D 物体快速换作品图或纹理。"),
            ("模型压缩流程", "Model Compression Pipeline", "训练 glTF 压缩和性能", "draco compression gltf three.js", "让 3D 网页加载更快。"),
            ("网页 AR 预览模型", "Web AR Model Preview", "训练 WebXR 和模型展示", "webxr ar model viewer", "用手机把 3D 作品放到真实空间看。"),
            ("3D 数据雕塑", "3D Data Sculpture", "训练数据映射到几何", "data sculpture three.js", "把销售、声音或城市数据变成 3D 形体。"),
            ("低多边形人物生成", "Low Poly Figure Generator", "训练几何简化和风格化", "low poly generator three.js", "生成适合视觉海报的低多边形角色。"),
            ("反射镜面空间", "Reflective Mirror Room", "训练环境贴图和反射", "three.js mirror reflection room", "模拟镜面、玻璃、金属空间效果。"),
            ("网页雕塑配置器", "Sculpture Configurator", "训练参数 UI 和 3D 更新", "3d configurator three.js", "客户可调尺寸、颜色、材质预览雕塑方案。"),
            ("WebGPU 视觉实验", "WebGPU Visual Experiment", "训练新一代浏览器图形管线", "webgpu generative art", "用更强的浏览器图形能力做高性能视觉。"),
        ]),
        ("AR / VR / XR", "AR VR XR", "WebXR / Unity / ARKit / ARCore", [
            ("手机 AR 看画", "Mobile AR Artwork Preview", "训练 AR 放置和比例", "ar artwork preview app", "客户用手机把作品挂到自己墙上预览。"),
            ("AR 展览导览", "AR Exhibition Guide", "训练定位、识别和信息层", "ar exhibition guide app", "手机对准作品出现文字、音频或动画讲解。"),
            ("VR 虚拟展厅", "VR Virtual Gallery", "训练 VR 场景和交互", "vr art gallery unity", "戴头显进入一个完整的虚拟展览。"),
            ("AR 作品说明卡", "AR Artwork Label", "训练图像识别和叠加内容", "ar image tracking artwork label", "扫作品或展签，显示更多内容。"),
            ("空间绘画工具", "Spatial Painting Tool", "训练 3D 绘制和手柄输入", "vr painting tool unity", "在三维空间里画线条和色块。"),
            ("AR 雕塑摆放器", "AR Sculpture Placement", "训练平面识别和模型加载", "ar sculpture placement app", "把雕塑模型放到真实房间里看尺寸。"),
            ("VR 策展预演", "VR Curatorial Preview", "训练展厅建模和路线", "vr curatorial preview", "布展前在虚拟空间检查作品关系。"),
            ("AR 明信片动画", "AR Animated Postcard", "训练图像跟踪和视频叠加", "ar animated postcard", "实体明信片被手机扫描后出现动画。"),
            ("WebXR 小展厅", "WebXR Mini Gallery", "训练浏览器 XR", "webxr art gallery", "不用安装 App，网页直接进入 AR/VR 展示。"),
            ("混合现实装置说明", "MR Installation Guide", "训练 3D 标注和空间说明", "mixed reality installation guide", "在现场用空间标注指导安装。"),
            ("AR 艺术课程卡", "AR Art Course Cards", "训练教育内容叠加", "ar education cards art", "课程卡片扫描后出现步骤动画。"),
            ("VR 声音空间", "VR Spatial Audio Room", "训练空间音频", "vr spatial audio art", "观众移动时声音方向和远近会变。"),
            ("AR 城市艺术路线", "AR Public Art Walk", "训练地图和 AR 导航", "ar public art walking tour", "把公共艺术路线做成手机导览。"),
            ("XR 作品发布会", "XR Artwork Launch", "训练多人空间和直播", "xr virtual event art", "线上观众一起进入虚拟发布现场。"),
            ("AR 尺寸校验工具", "AR Scale Check Tool", "训练现实尺度和单位", "ar scale measurement app", "检查作品尺寸在真实空间里是否合适。"),
        ]),
        ("计算机视觉", "Computer Vision", "OpenCV / MediaPipe / Python / TensorFlow", [
            ("作品自动裁边工具", "Artwork Auto Cropper", "训练边缘检测和透视校正", "opencv artwork auto crop", "拍照后自动裁掉背景，校正作品边缘。"),
            ("展览人流计数器", "Exhibition People Counter", "训练目标检测和计数", "people counter opencv exhibition", "统计展厅人流，不用人工数。"),
            ("姿态生成绘画", "Pose Driven Drawing", "训练人体姿态识别", "mediapipe pose drawing", "身体动作变成线条或图形。"),
            ("手势控制画布", "Gesture Controlled Canvas", "训练手势识别和交互", "hand tracking canvas mediapipe", "用手势切换颜色、放大、旋转。"),
            ("作品相似图检索", "Visual Similarity Search", "训练特征向量和检索", "image similarity search art", "从作品库里找视觉相近的图片。"),
            ("画面构图评分器", "Composition Score Assistant", "训练规则检测和辅助评估", "image composition analysis", "检查主体位置、留白、对称等构图特征。"),
            ("色彩偏差校正", "Color Correction Tool", "训练白平衡和颜色校正", "opencv color correction artwork", "让拍摄图更接近真实作品颜色。"),
            ("展签 OCR 录入", "Artwork Label OCR", "训练文字识别和数据清洗", "ocr artwork label", "拍展签自动识别标题、作者、年份。"),
            ("观众表情反馈", "Audience Expression Feedback", "训练表情识别和统计", "facial expression recognition installation", "把观众表情转成匿名统计反馈。"),
            ("作品瑕疵检测", "Artwork Defect Detection", "训练图像差异和异常检测", "image defect detection artwork", "检查印刷、装裱或拍摄瑕疵。"),
            ("动态轮廓投影", "Live Contour Projection", "训练轮廓提取和实时渲染", "opencv contour projection", "实时提取观众轮廓并投影成图形。"),
            ("二维码展览入口", "QR Exhibition Entry", "训练二维码识别和跳转", "qr code exhibition app", "扫码进入线上资料、语音导览或问卷。"),
            ("作品拍摄质量检查", "Artwork Photo Quality Checker", "训练模糊、曝光和倾斜检测", "photo quality check opencv", "自动提醒照片是否模糊、过曝、歪斜。"),
            ("动作触发音画", "Motion Triggered AV", "训练运动检测和触发", "motion detection audio visual", "画面中有人移动就触发声音或视觉。"),
            ("图像风格聚类器", "Image Style Clustering", "训练聚类和降维", "art image clustering", "把一批作品按视觉风格自动分组。"),
        ]),
        ("机器人与硬件", "Robotics and Hardware", "Arduino / Raspberry Pi / ESP32 / Serial", [
            ("自动绘图机", "Pen Plotter", "训练步进电机和矢量路径", "arduino pen plotter svg", "机器按 SVG 路径画线，适合生成艺术落地成纸上作品。"),
            ("灯光矩阵控制器", "LED Matrix Controller", "训练矩阵显示和动画", "esp32 led matrix art", "用很多 LED 点组成动态像素画。"),
            ("机械臂绘画实验", "Robot Arm Painting", "训练机械臂路径和控制", "robot arm painting art", "让机械臂按规划路径绘画。"),
            ("热敏打印诗歌机", "Thermal Printer Poetry Machine", "训练串口打印和文本生成", "thermal printer poetry arduino", "按钮一按打印一段生成诗或展览签语。"),
            ("电子墨水展签", "E Ink Artwork Label", "训练低功耗显示和内容更新", "e ink museum label esp32", "展签内容可远程更新，省去重新打印。"),
            ("旋转光影装置", "Rotating Light Sculpture", "训练电机、光源和速度控制", "motorized light sculpture arduino", "转速改变会形成不同光影图案。"),
            ("物联网作品状态监控", "IoT Artwork Monitor", "训练传感器联网和告警", "iot artwork environmental monitor", "监测温湿度、震动，保护作品。"),
            ("蓝牙灯光遥控器", "Bluetooth Light Remote", "训练 BLE 和移动控制", "bluetooth led controller esp32", "手机控制装置灯光模式。"),
            ("伺服机关节雕塑", "Servo Joint Sculpture", "训练伺服电机和动作编排", "servo kinetic sculpture", "雕塑部件可按程序运动。"),
            ("按钮互动展台", "Button Interactive Pedestal", "训练按钮输入和媒体触发", "arduino button media trigger", "观众按按钮切换屏幕内容或声音。"),
            ("RFID 作品识别柜", "RFID Artwork Cabinet", "训练 RFID 和库存记录", "rfid art inventory", "作品放入柜中自动识别并记录。"),
            ("温湿度保护面板", "Climate Protection Panel", "训练传感器和数据显示", "temperature humidity art conservation", "实时显示环境是否适合作品保存。"),
            ("机械翻页展册", "Mechanical Page Turner", "训练电机结构和控制", "mechanical page turner art", "自动翻动实体书页或展册。"),
            ("激光切割灯箱", "Laser Cut Lightbox", "训练矢量文件和灯光结构", "laser cut lightbox art", "把图案切割成层叠灯箱。"),
            ("硬件项目测试台", "Hardware Test Bench", "训练供电、串口和日志", "arduino hardware test bench", "集中测试传感器、灯带、电机是否稳定。"),
        ]),
        ("数字制造", "Digital Fabrication", "Blender / Rhino / Grasshopper / CNC / 3D Print", [
            ("参数化展架", "Parametric Display Stand", "训练尺寸参数和结构", "parametric display stand grasshopper", "输入作品尺寸自动生成展架结构。"),
            ("3D 打印首饰系列", "3D Printed Jewelry Series", "训练建模、切片和材料", "3d printed jewelry parametric", "把生成式图案变成可打印首饰。"),
            ("激光切割拼插装置", "Laser Cut Slot Installation", "训练卡榫结构和板材厚度", "laser cut slot structure art", "不用螺丝，用插接结构搭装置。"),
            ("CNC 浮雕生成", "CNC Relief Generator", "训练高度图和刀路", "cnc relief heightmap", "把图像明暗转成可雕刻的浮雕。"),
            ("纸艺展开图工具", "Paper Craft Unfold Tool", "训练 3D 展开和编号", "papercraft unfold tool", "把立体形态展开成可打印折纸图。"),
            ("展览模型切片器", "Exhibition Model Slicer", "训练模型切片和装配", "laser cut model slicer", "把 3D 模型切成板材层片。"),
            ("陶瓷纹理压印板", "Ceramic Texture Stamp", "训练图案转模型", "3d printed ceramic texture stamp", "把纹样变成陶瓷压印工具。"),
            ("参数化灯罩", "Parametric Lampshade", "训练曲面、孔洞和光影", "parametric lampshade grasshopper", "用参数控制灯罩孔洞和投影效果。"),
            ("模块化展墙系统", "Modular Exhibition Wall System", "训练模块结构和清单", "modular exhibition wall design", "设计可重复使用的展墙模块。"),
            ("可变尺寸画框", "Variable Size Frame Generator", "训练尺寸计算和加工图", "parametric picture frame generator", "输入作品尺寸自动生成画框加工尺寸。"),
            ("机器人刀路草图", "Robot Toolpath Sketch", "训练路径规划和制造", "robotic fabrication toolpath", "为机械臂切割、画线或涂胶准备路径。"),
            ("生成式浮雕墙砖", "Generative Relief Tiles", "训练重复模块和高度变化", "generative relief tiles", "做一组可拼接的立体墙砖。"),
            ("作品包装盒结构", "Artwork Packaging Box", "训练纸盒展开和保护结构", "custom artwork packaging box dieline", "按作品尺寸生成包装盒和缓冲结构。"),
            ("展览导视牌切割图", "Exhibition Signage Cut Files", "训练字体转曲和切割文件", "laser cut exhibition signage", "把导视文字做成可切割文件。"),
            ("材料损耗计算器", "Material Waste Calculator", "训练排版优化和成本", "sheet material nesting calculator", "计算板材、纸张、亚克力的用量和损耗。"),
        ]),
        ("艺术教育工具", "Art Education Tools", "React / p5.js / LMS / Quiz", [
            ("创意编程课程平台", "Creative Coding Course Platform", "训练课程结构和代码嵌入", "creative coding course platform", "学生能看教程、写代码、交作业。"),
            ("色彩理论互动课", "Interactive Color Theory Lesson", "训练颜色模型和交互", "interactive color theory web", "拖动色相、饱和度、明度理解配色。"),
            ("构图练习小游戏", "Composition Practice Game", "训练拖拽和评分", "composition practice web app", "学生拖动画面元素练习构图。"),
            ("艺术史时间轴", "Art History Timeline", "训练时间线和资料卡片", "art history timeline web app", "把艺术运动、人物、作品放到时间线上。"),
            ("课堂作品互评系统", "Peer Critique System", "训练上传、评论和权限", "art critique platform", "学生上传作品，互相评论，老师汇总。"),
            ("图像分析练习器", "Image Analysis Trainer", "训练标注和问答", "image analysis education tool", "围绕作品练习观察、描述、解释。"),
            ("提示词训练卡", "Prompt Training Cards", "训练卡片系统和随机题", "prompt engineering flashcards", "随机抽取风格、媒介、构图条件练提示词。"),
            ("艺术材料知识库", "Art Materials Knowledge Base", "训练搜索和分类", "art material database", "整理颜料、纸张、布面、工具的用途和价格。"),
            ("学生作品集生成器", "Student Portfolio Builder", "训练模板和导出", "student portfolio builder", "学生填资料后一键生成作品集页面或 PDF。"),
            ("在线展览作业墙", "Online Class Exhibition Wall", "训练上传和画廊布局", "online student exhibition wall", "全班作品像展览一样在线展示。"),
            ("艺术词汇双语表", "Bilingual Art Vocabulary", "训练词库和测验", "bilingual art vocabulary app", "中英术语对照，适合国际申请和作品集。"),
            ("创作过程日志", "Creative Process Journal", "训练日记、图片和版本", "creative process journal app", "记录草图、修改和反思，方便教学评价。"),
            ("数字绘画笔刷实验", "Digital Brush Lab", "训练笔刷参数和对比", "digital brush lab canvas", "学生理解笔刷大小、透明度、纹理如何影响画面。"),
            ("课堂签到二维码", "Class QR Attendance", "训练二维码和表格", "qr attendance classroom", "扫码签到，自动统计课程出勤。"),
            ("艺术项目评分表", "Art Rubric Builder", "训练表单和评分规则", "art rubric builder", "生成清楚的作品评分维度和反馈表。"),
        ]),
        ("展览管理", "Exhibition Management", "Spreadsheets / Airtable / Notion / Web App", [
            ("展览项目总看板", "Exhibition Master Board", "训练项目管理和状态流", "exhibition project management board", "把作品、人员、设备、时间统一管理。"),
            ("作品借展追踪", "Loan Artwork Tracker", "训练借出归还和提醒", "art loan tracker", "记录作品借给谁、何时归还、保险状态。"),
            ("布展任务分配表", "Installation Task Planner", "训练任务拆分和责任人", "installation task planner", "清楚谁负责运输、灯光、墙面、文字、现场。"),
            ("设备清单管理", "Equipment Inventory Manager", "训练库存和借用记录", "equipment inventory exhibition", "投影、音响、线材、电脑不再靠记忆管理。"),
            ("展览预算模板", "Exhibition Budget Template", "训练预算分类和公式", "exhibition budget spreadsheet", "自动计算制作、运输、宣传、人员成本。"),
            ("艺术品保险记录", "Artwork Insurance Register", "训练文档和附件管理", "artwork insurance register", "记录估值、保险、运输和照片证据。"),
            ("展签批量生成器", "Artwork Label Batch Generator", "训练数据转 Word/PDF", "batch generate museum labels", "从表格自动生成统一格式展签。"),
            ("开幕 RSVP 系统", "Opening RSVP System", "训练表单和名单管理", "rsvp system exhibition opening", "统计谁会来、带几个人、联系方式。"),
            ("媒体联系清单", "Press Contact CRM", "训练联系人分类和跟进", "press contact crm art", "管理媒体、博主、机构联系人。"),
            ("展览风险清单", "Exhibition Risk Checklist", "训练检查表和责任人", "exhibition risk checklist", "提前检查消防、用电、版权、运输、保险风险。"),
            ("撤展归档流程", "Deinstallation Archive Workflow", "训练归档和文件命名", "exhibition deinstallation archive", "撤展后把照片、合同、反馈、销售记录归档。"),
            ("展览复盘报告", "Exhibition Review Report", "训练数据汇总和文档生成", "exhibition post mortem report", "把观众、销售、传播、问题整理成报告。"),
            ("现场问题记录器", "Onsite Issue Logger", "训练移动表单和通知", "onsite issue tracker exhibition", "现场发现问题马上记录和分派。"),
            ("作品运输标签", "Artwork Shipping Labels", "训练标签打印和二维码", "artwork shipping label generator", "每个箱子贴标签，扫码看到作品和运输信息。"),
            ("策展资料索引", "Curatorial Research Index", "训练资料库和引用", "curatorial research database", "把参考文章、图片、访谈、出处整理成可检索索引。"),
        ]),
        ("客户与商业", "Client and Business", "CRM / Forms / Automation / Payments", [
            ("艺术客户 CRM", "Art Client CRM", "训练客户数据和跟进", "artist client crm", "记录客户偏好、预算、购买历史和跟进时间。"),
            ("报价生成器", "Quote Generator", "训练模板和自动计算", "artist quote generator", "按尺寸、授权、工期、材料自动生成报价。"),
            ("授权合同助手", "Licensing Contract Assistant", "训练合同字段和条款模板", "art licensing contract template", "把授权范围、期限、地区、价格写清楚。"),
            ("定金尾款提醒", "Deposit Balance Reminder", "训练日期提醒和邮件", "payment reminder automation", "自动提醒客户付定金或尾款。"),
            ("委托创作流程表", "Commission Workflow Tracker", "训练阶段管理", "art commission workflow", "从需求、草图、定稿、制作到交付全流程追踪。"),
            ("客户需求问卷", "Client Brief Form", "训练表单设计和数据整理", "creative brief form artist", "让客户先填清楚尺寸、风格、用途、预算。"),
            ("作品可售清单", "Available Works List", "训练库存和筛选", "available artworks list", "快速发给客户当前可购买作品。"),
            ("销售漏斗看板", "Art Sales Pipeline", "训练看板和概率", "art sales pipeline crm", "看哪些客户只是咨询，哪些快成交。"),
            ("发票资料收集器", "Invoice Info Collector", "训练表单和校验", "invoice info form automation", "提前收集开票信息，减少来回问。"),
            ("合作方资料库", "Partner Directory", "训练联系人和标签", "art partner directory", "整理画廊、品牌、工厂、摄影、运输资源。"),
            ("版画编号管理", "Edition Number Manager", "训练编号和状态", "limited edition print manager", "管理第几版卖给谁，避免编号混乱。"),
            ("授权使用追踪", "License Usage Tracker", "训练授权状态和到期提醒", "art license usage tracker", "记录客户能在哪些渠道使用作品，到期自动提醒。"),
            ("课程报名系统", "Workshop Enrollment System", "训练报名、支付和名单", "art workshop enrollment system", "报名、付款、名单、提醒统一管理。"),
            ("客户满意度回访", "Client Feedback Follow Up", "训练问卷和自动邮件", "client feedback automation", "交付后自动收集反馈和复购机会。"),
            ("艺术服务价格表", "Service Pricing Sheet", "训练价格模型和版本", "creative service pricing spreadsheet", "把插画、授权、课程、装置等服务价格标准化。"),
        ]),
        ("内容自动化", "Content Automation", "Python / APIs / Canva / Video", [
            ("社媒发布日历", "Social Posting Calendar", "训练日历和内容排期", "social media content calendar artist", "提前安排作品发布、过程分享、展览预告。"),
            ("作品文案生成器", "Artwork Caption Generator", "训练文案模板和语气", "artwork caption generator", "根据作品信息生成不同平台的发布文案。"),
            ("九宫格排版器", "Nine Grid Layout Maker", "训练图片裁切和布局", "instagram nine grid generator", "把一张大图切成社媒九宫格。"),
            ("短视频字幕工具", "Short Video Caption Tool", "训练语音转文字和字幕", "auto subtitle short video", "给作品过程视频自动加字幕。"),
            ("海报尺寸批量适配", "Poster Size Batch Resizer", "训练图片尺寸和模板", "batch resize posters social media", "一张图自动适配小红书、抖音、Instagram、微信。"),
            ("作品水印批处理", "Artwork Watermark Batch Tool", "训练批量图片处理", "batch watermark artwork images", "给作品图统一加水印和版权信息。"),
            ("发布素材命名器", "Content Asset Naming Tool", "训练命名规则和文件整理", "batch rename content assets", "素材按日期、平台、主题命名，不再乱。"),
            ("展览倒计时模板", "Exhibition Countdown Templates", "训练模板和批量生成", "exhibition countdown social templates", "开幕前自动生成倒计时图文。"),
            ("作品过程视频合成", "Process Video Assembler", "训练图片序列和视频导出", "image sequence to video tool", "把过程照片合成短视频。"),
            ("多平台文案改写器", "Multi Platform Caption Rewriter", "训练文本改写和格式", "rewrite captions for social platforms", "同一内容改成小红书、微博、英文 Instagram 版本。"),
            ("标签关键词建议器", "Hashtag Keyword Suggester", "训练关键词和分类", "artist hashtag generator", "根据作品主题生成更容易搜索的标签。"),
            ("内容表现复盘表", "Content Performance Review", "训练指标和复盘", "social media performance spreadsheet", "统计哪些内容带来咨询和成交。"),
            ("邮件简报生成器", "Newsletter Generator", "训练邮件模板和列表", "artist newsletter generator", "把新作品、展览、课程整理成邮件。"),
            ("作品图压缩发布器", "Image Compression Publisher", "训练图片压缩和格式", "web image compression pipeline", "让网站和社媒图片清晰但体积小。"),
            ("自动封面挑选器", "Auto Cover Selector", "训练图片评分和裁切", "auto select video thumbnail", "从视频里挑最适合当封面的画面。"),
        ]),
        ("移动应用", "Mobile Apps", "React Native / Flutter / Swift / Kotlin", [
            ("艺术家口袋库存", "Pocket Art Inventory", "训练移动端表单和同步", "mobile art inventory app", "手机随时查作品状态和价格。"),
            ("展览扫码导览", "QR Exhibition Guide App", "训练二维码和内容页面", "qr museum guide mobile app", "观众扫码看图文、音频、视频导览。"),
            ("作品拍摄助手", "Artwork Photo Assistant", "训练相机、网格和校正", "artwork photography assistant app", "拍作品时辅助对齐、曝光、裁切。"),
            ("客户看图确认 App", "Client Approval App", "训练评论和版本确认", "client approval app design", "客户在手机上确认草图、留言和选择版本。"),
            ("艺术课程打卡", "Art Course Check In", "训练签到和学习记录", "course check in mobile app", "学生记录练习、提交作业、查看反馈。"),
            ("城市艺术地图 App", "Public Art Map App", "训练定位和地图", "public art map mobile app", "按位置发现附近艺术空间和公共艺术。"),
            ("展览志愿者 App", "Exhibition Volunteer App", "训练任务和通知", "volunteer management app exhibition", "志愿者查看排班、任务和现场通知。"),
            ("收藏家私享预览", "Collector Private Preview App", "训练登录和私密内容", "collector preview app art", "给重要客户看专属作品清单。"),
            ("艺术灵感速记", "Art Inspiration Capture", "训练图片、语音和标签", "inspiration capture mobile app", "看到灵感马上拍照、录音、加标签。"),
            ("创作计时器", "Studio Work Timer", "训练计时和统计", "studio work timer app", "记录每个项目花了多少时间。"),
            ("材料采购清单", "Art Supply Shopping List", "训练清单和预算", "art supply list app", "记录材料库存和采购计划。"),
            ("展览票务小程序", "Exhibition Ticket Mini App", "训练票务和核销", "exhibition ticket app", "线上预约，现场扫码核销。"),
            ("艺术活动日历", "Art Event Calendar App", "训练日历和提醒", "art event calendar app", "记录展览、讲座、申请截止日期。"),
            ("作品授权查看器", "License Viewer App", "训练合同摘要和搜索", "license tracker mobile app", "手机快速查看某作品授权状态。"),
            ("艺术家名片 App", "Artist Digital Card", "训练电子名片和链接聚合", "digital business card artist", "一页展示作品集、社媒、邮箱和二维码。"),
        ]),
        ("开源与工程能力", "Open Source and Engineering", "GitHub / CI / Testing / Docs", [
            ("作品集代码模板库", "Portfolio Starter Template", "训练模板化和文档", "artist portfolio starter github", "做一个可反复复制的新网站起点。"),
            ("创意编程组件库", "Creative Coding Component Library", "训练组件封装", "creative coding component library", "把常用画布、控制面板、导出按钮封装起来。"),
            ("项目 README 生成器", "Project README Generator", "训练文档结构", "readme generator github project", "每个技术作品都有清楚说明、截图、运行方式。"),
            ("自动部署作品集", "Auto Deploy Portfolio", "训练 CI/CD 和托管", "github actions deploy portfolio", "代码推到 GitHub 后网站自动更新。"),
            ("视觉回归测试", "Visual Regression Tests", "训练截图比较", "visual regression testing playwright", "改代码后自动检查页面视觉有没有坏。"),
            ("素材许可证检查", "Asset License Checker", "训练文件扫描和规则", "asset license checker", "检查项目里图片、字体、音乐有没有记录授权。"),
            ("作品数据 JSON 标准", "Artwork JSON Schema", "训练数据结构和校验", "artwork metadata json schema", "统一作品标题、尺寸、媒介、年份、价格字段。"),
            ("创意代码范例库", "Creative Code Examples Repo", "训练案例整理", "creative coding examples repository", "把小实验整理成可复用学习仓库。"),
            ("Issue 模板和需求表", "Issue Templates for Art Projects", "训练协作流程", "github issue templates creative project", "团队协作时需求、Bug、设计反馈更清楚。"),
            ("版本化作品说明", "Versioned Artist Statements", "训练 Git 版本管理", "version control artist statement", "作品文案每次修改都有记录，可回溯。"),
            ("项目截图自动生成", "Auto Screenshot Generator", "训练浏览器自动化", "playwright screenshot generator", "自动给网页作品生成展示截图。"),
            ("技术作品归档站", "Technical Work Archive", "训练静态站和元数据", "technical portfolio archive", "把每个技术实验按时间和主题归档。"),
            ("代码片段搜索库", "Code Snippet Search Library", "训练搜索和标签", "code snippet search app", "快速找之前写过的动画、图表、交互代码。"),
            ("插件使用手册", "Plugin Usage Manual", "训练技术写作和流程图", "plugin documentation manual", "把常用插件的用法整理成自己的 SOP。"),
            ("开源贡献路线图", "Open Source Contribution Roadmap", "训练项目拆解和 PR 流程", "open source contribution roadmap", "从修文档、小 Bug 到提交功能，逐步建立技术信用。"),
        ]),
        ("AI 助手与文本", "AI Assistants and Text", "LLM / RAG / Vector Search / Automation", [
            ("艺术家知识库助手", "Artist Knowledge Base Assistant", "训练 RAG 和资料检索", "artist knowledge base rag", "把简历、作品说明、采访整理成可问答资料库。"),
            ("客户邮件回复助手", "Client Email Reply Assistant", "训练邮件语气和模板", "client email assistant artist", "根据客户问题生成专业但不生硬的回复。"),
            ("展览文案润色器", "Exhibition Text Polisher", "训练文案改写", "exhibition text editing ai", "把展览介绍改得更清楚、更适合发布。"),
            ("作品说明多版本生成", "Artwork Statement Variants", "训练风格化写作", "artwork statement generator", "同一作品生成学术版、客户版、社媒版。"),
            ("中英翻译术语库", "Chinese English Art Glossary", "训练术语对照和检索", "bilingual art glossary", "常用艺术、技术、商业词汇中英统一。"),
            ("策展主题研究助手", "Curatorial Research Assistant", "训练资料摘要和引用管理", "curatorial research assistant ai", "围绕主题整理关键词、案例和论点。"),
            ("合同风险提示器", "Contract Risk Highlighter", "训练条款识别", "contract risk highlighter", "提醒授权、付款、交付、版权条款是否缺失。"),
            ("作品命名助手", "Artwork Title Assistant", "训练标题生成和筛选", "artwork title generator", "根据主题生成中英文作品名。"),
            ("课程大纲生成器", "Workshop Syllabus Generator", "训练教学结构", "art workshop syllabus generator", "把课程目标、课时、材料、作业整理出来。"),
            ("艺术申请材料助手", "Art Application Assistant", "训练申请文书结构", "artist grant application assistant", "辅助整理驻留、基金、展览申请材料。"),
            ("采访问题生成器", "Interview Question Generator", "训练问题设计", "artist interview questions generator", "为艺术家访谈或媒体采访准备问题。"),
            ("项目复盘摘要器", "Project Debrief Summarizer", "训练摘要和行动项", "project retrospective summarizer", "把会议记录、展览反馈整理成复盘。"),
            ("评论情绪分析", "Comment Sentiment Analyzer", "训练文本分类", "social comment sentiment analysis", "分析观众评论是喜欢、困惑还是询价。"),
            ("灵感笔记分类器", "Inspiration Note Classifier", "训练文本标签和聚类", "note classification ai", "把零散灵感按主题自动归类。"),
            ("双语发布检查器", "Bilingual Publishing Checker", "训练翻译校对和语气", "bilingual copy checker", "检查中文英文发布内容是否意思一致。"),
        ]),
        ("区块链与数字藏品", "Blockchain and Digital Collectibles", "Ethereum / IPFS / Smart Contracts", [
            ("数字作品证书", "Digital Artwork Certificate", "训练哈希和元数据", "digital artwork certificate blockchain", "给数字作品生成可验证的证书记录。"),
            ("限量版铸造流程", "Limited Edition Minting Flow", "训练智能合约和版数", "limited edition nft smart contract", "控制数字版画的数量、编号和归属。"),
            ("IPFS 作品归档", "IPFS Artwork Archive", "训练去中心化存储", "ipfs artwork metadata", "把作品文件和元数据做长期保存。"),
            ("藏品元数据编辑器", "Collectible Metadata Editor", "训练 JSON 和属性字段", "nft metadata editor", "编辑标题、描述、属性、图片链接。"),
            ("钱包登录展厅", "Wallet Login Gallery", "训练钱包连接和权限", "wallet gated gallery", "持有某作品的人才能进入专属页面。"),
            ("链上版税说明页", "On Chain Royalty Explainer", "训练合约字段和可视化", "nft royalty smart contract", "解释版税如何设置、在哪里生效。"),
            ("藏品持有人地图", "Collector Holder Map", "训练链上数据和地图", "nft holder map visualization", "看数字藏品持有人分布。"),
            ("数字作品出处链", "Provenance Chain Viewer", "训练交易记录可视化", "art provenance blockchain viewer", "展示作品从发行到转手的记录。"),
            ("智能合约测试模板", "Smart Contract Test Template", "训练自动测试", "smart contract tests nft", "在发布前测试版数、转让、权限是否正确。"),
            ("链上展览门票", "On Chain Exhibition Ticket", "训练票据和核销", "blockchain exhibition ticket", "用数字票做入场或会员权益。"),
            ("数字藏品落地页", "Digital Collectible Landing Page", "训练作品介绍和购买流程", "nft landing page artist", "给一组数字作品做清楚的介绍页。"),
            ("合约风险清单", "Contract Risk Checklist", "训练安全检查", "smart contract audit checklist nft", "检查合约权限、费用、元数据是否有风险。"),
            ("版税收入表", "Royalty Revenue Tracker", "训练交易数据和表格", "nft royalty revenue tracker", "记录数字作品版税收入。"),
            ("链下权益管理", "Off Chain Benefits Manager", "训练会员权益和核验", "token gated benefits manager", "数字藏品对应实体活动、课程、折扣等权益。"),
            ("数字作品法律说明", "Digital Art Rights Explainer", "训练版权说明和条款", "digital art rights nft license", "清楚写明买家买到的是作品、授权还是收藏凭证。"),
        ]),
        ("研究与档案", "Research and Archive", "Databases / OCR / Metadata / Search", [
            ("个人作品档案库", "Personal Artwork Archive", "训练数据库和元数据", "personal artwork archive database", "把所有作品、照片、说明、展览记录统一归档。"),
            ("纸质资料 OCR 库", "Paper Archive OCR Library", "训练扫描和文字识别", "ocr archive documents", "把旧展册、合同、手稿扫描成可搜索文本。"),
            ("艺术家访谈索引", "Artist Interview Index", "训练音频转写和标签", "artist interview transcript database", "访谈录音转文字并按主题检索。"),
            ("展览图片档案", "Exhibition Photo Archive", "训练图片管理和标签", "exhibition photo archive", "把布展、开幕、现场照片整理清楚。"),
            ("作品出处数据库", "Artwork Provenance Database", "训练历史记录和关系", "artwork provenance database", "记录作品何时创作、展出、售出、收藏。"),
            ("文献引用管理", "Reference Citation Manager", "训练引用格式和资料卡", "art research citation manager", "研究文章、书籍、网页有统一引用。"),
            ("视觉参考检索器", "Visual Reference Searcher", "训练图片标签和相似检索", "visual reference search database", "快速找以前收藏的参考图。"),
            ("展览年表生成器", "Exhibition Chronology Generator", "训练时间线和导出", "exhibition chronology generator", "自动生成艺术家展览年表。"),
            ("收藏证明整理器", "Collection Proof Organizer", "训练文件归档和清单", "art collection certificate organizer", "整理收藏证书、发票、运输单、合同。"),
            ("艺术项目档案包", "Art Project Archive Package", "训练文件结构和压缩", "art project archive package", "项目结束后一键生成归档文件夹结构。"),
            ("跨语言资料摘要", "Cross Language Research Summary", "训练翻译和摘要", "cross language research summary", "英文资料摘要成中文，并保留关键词。"),
            ("图片元数据清理", "Image Metadata Cleaner", "训练 EXIF 和隐私", "remove exif metadata images", "发布前清理图片里的位置、设备等隐私信息。"),
            ("展览资料检索站", "Exhibition Search Portal", "训练全文搜索", "exhibition archive search portal", "输入关键词就能找历史展览资料。"),
            ("作品编号规则系统", "Artwork Numbering System", "训练编号规则和校验", "artwork inventory numbering system", "统一作品编号，避免重复和混乱。"),
            ("档案备份监控", "Archive Backup Monitor", "训练备份状态和提醒", "archive backup monitor", "检查云盘和硬盘备份是否完整。"),
        ]),
        ("安全、版权与发布", "Security Copyright Publishing", "Watermark / Metadata / Access Control", [
            ("作品图片防盗链", "Image Hotlink Protection", "训练网站安全和资源权限", "image hotlink protection website", "防止别人直接盗用你网站图片链接。"),
            ("高清图访问权限", "High Res Access Control", "训练登录和文件权限", "high resolution image access control", "只有客户或合作方能看高清图。"),
            ("图片水印策略工具", "Watermark Strategy Tool", "训练批处理和透明度", "watermark strategy artwork", "按用途生成不同强度的水印版本。"),
            ("版权声明生成器", "Copyright Notice Generator", "训练文案模板和元数据", "copyright notice generator artist", "为网站、PDF、图片生成统一版权声明。"),
            ("作品授权检查表", "Artwork License Checklist", "训练授权字段和风险", "art licensing checklist", "检查授权范围、期限、地区、媒体、费用是否完整。"),
            ("网站隐私页模板", "Website Privacy Template", "训练网页政策和表单", "artist website privacy policy", "有表单和邮件订阅时需要说明数据怎么用。"),
            ("客户文件下载页", "Client File Download Portal", "训练临时链接和过期时间", "secure client file download", "给客户下载文件，同时控制有效期。"),
            ("合同版本对比", "Contract Version Diff", "训练文本差异比较", "contract version comparison", "看客户改了合同哪里。"),
            ("图片发布尺寸规范", "Image Publishing Presets", "训练导出规格", "image export presets social web print", "统一网站、社媒、印刷预览的图片尺寸。"),
            ("作品署名监测", "Artwork Attribution Monitor", "训练网页搜索和记录", "artwork attribution monitoring", "定期搜索作品是否被错误转载。"),
            ("敏感信息打码器", "Sensitive Info Redactor", "训练图像和文档脱敏", "redact sensitive information documents", "发布前遮住地址、电话、身份证、价格等信息。"),
            ("客户资料加密表", "Encrypted Client Records", "训练加密和访问权限", "encrypted client records", "保护客户联系方式和购买记录。"),
            ("数字签名交付单", "Digital Delivery Receipt", "训练签名和 PDF", "digital delivery receipt artwork", "交付后让客户确认收到文件或作品。"),
            ("版权证据时间戳", "Copyright Evidence Timestamp", "训练哈希和时间记录", "copyright timestamp proof", "给创作文件做时间证明和校验记录。"),
            ("发布前检查清单", "Pre Publish Checklist", "训练自动检查和流程", "pre publish checklist website art", "检查图片、链接、版权、错别字、尺寸后再发布。"),
        ]),
    ]

    projects = []
    for category_zh, category_en, stack, rows in categories:
        for zh, en, value, search, explain in rows:
            projects.append({
                "category": f"{category_zh} / {category_en}",
                "name": f"{zh}\n{en}",
                "value": value,
                "stack": stack,
                "search": search,
                "explain": explain,
            })
    assert len(projects) >= 300
    return projects


def add_source_table(doc):
    add_heading(doc, "七、补充调研来源", 1)
    add_para(doc, "我用 GitHub 和 Chrome 先看了艺术科技常见方向，再把方向拆成后面的 300+ 个可做项目。下面是这次扩展时用来定方向的来源，不是让你照抄，而是让你知道后面项目库从哪里延展出来。")
    rows = [
        ("GitHub Topics: generative-art", "生成艺术主题页显示大量公共仓库，典型方向包括 AI 图像、creative coding、WebGL、shader、地图和数据可视化。", "https://github.com/topics/generative-art"),
        ("openFrameworks README", "openFrameworks 是 C++ 创意编程工具包，适合装置、实时视觉、硬件和跨平台作品。", "https://github.com/openframeworks/openFrameworks"),
        ("p5.js 官网", "p5.js 面向艺术家、设计师、教育者和初学者，是适合入门创意编程的 JavaScript 开源库。", "https://p5js.org/"),
        ("three.js manual", "three.js 把 WebGL 里的场景、相机、材质、灯光等复杂部分封装起来，适合网页 3D。", "https://threejs.org/manual/en/fundamentals.html"),
        ("Processing 官网", "Processing 提供免费开源的开发环境、示例和参考资料，适合视觉艺术编程学习。", "https://processing.org/"),
        ("TouchDesigner / Derivative", "TouchDesigner 常用于高性能媒体系统、互动装置、投影、音画和实时演出。", "https://derivative.ca/"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    set_table_borders(table)
    widths = [5.0, 12.0, 9.2]
    for idx, header in enumerate(("来源 / Source", "搜索内容摘要 / Search Summary", "链接 / Link")):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, BLUE)
        cell_text(cell, header, size=8.2, bold=True, color="FFFFFF")
        set_cell_width(cell, widths[idx])
    repeat_table_header(table.rows[0])
    for row_idx, row in enumerate(rows, start=1):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_shading(cells[i], "FFFFFF" if row_idx % 2 else LIGHT_GRAY)
            cell_text(cells[i], value, size=7.6, bold=(i == 0))
            set_cell_width(cells[i], widths[i])


def add_project_library(doc):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.1)
    section.right_margin = Cm(1.1)

    add_heading(doc, "八、300+ 个艺术行业技术项目库", 1)
    add_note(
        doc,
        "怎么用这 300+ 条",
        "先不要全部做。建议按你的当前目标选 10 个：作品包装选 Canva/网站/内容自动化，互动装置选 TouchDesigner/Arduino/视觉识别，商业效率选 CRM/报价/展览管理。每条都给了中英项目名和搜索关键词，方便你继续用 GitHub、Google、YouTube、B站搜索教程或开源代码。",
        fill=AMBER,
    )

    widths = [0.9, 3.3, 3.7, 4.0, 3.6, 4.6, 6.2]
    headers = ["序号", "方向", "项目 / Project", "适合练什么", "技术栈", "搜索关键词", "技术白话解释"]
    projects = build_projects()
    grouped = []
    for item in projects:
        if not grouped or grouped[-1][0] != item["category"]:
            grouped.append((item["category"], []))
        grouped[-1][1].append(item)

    current_index = 1
    for group_index, (category, rows) in enumerate(grouped):
        if group_index:
            doc.add_page_break()
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(category)
        set_east_asian_font(r, size=11, bold=True, color=BLUE)

        table = doc.add_table(rows=1, cols=7)
        table.autofit = False
        set_table_borders(table)
        for idx, header in enumerate(headers):
            cell = table.rows[0].cells[idx]
            set_cell_shading(cell, BLUE)
            cell_text(cell, header, size=7.4, bold=True, color="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_width(cell, widths[idx])
        repeat_table_header(table.rows[0])

        for offset, item in enumerate(rows):
            cells = table.add_row().cells
            values = [
                str(current_index),
                item["category"],
                item["name"],
                item["value"],
                item["stack"],
                item["search"],
                item["explain"],
            ]
            for col, value in enumerate(values):
                set_cell_shading(cells[col], "FFFFFF" if offset % 2 == 0 else LIGHT_GRAY)
                cell_text(cells[col], value, size=6.8 if col not in (0, 1) else 6.6, bold=(col == 2))
                set_cell_width(cells[col], widths[col])
            current_index += 1


def rebuild_footer(doc):
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.text = ""
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = footer.add_run("艺术行业效率插件建议 | 白话版 + 300+ 个技术项目库")
        set_east_asian_font(r, size=8, color="78808C")


def main():
    doc = Document(DOCX_PATH)
    add_source_table(doc)
    add_project_library(doc)
    rebuild_footer(doc)
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
