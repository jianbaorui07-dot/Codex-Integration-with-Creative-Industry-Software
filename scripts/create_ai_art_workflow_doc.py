from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUT_PATH = r"C:\Users\jian\Documents\New project\output\AI艺术出图项目体系白话总结_5人工作组.docx"


ACCENT = "1F4E5F"
ACCENT_DARK = "173A46"
ACCENT_LIGHT = "EAF3F5"
MUTED = "6B7280"
GRID = "CBD5E1"
TEXT = "111827"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=GRID, size="6"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=120, start=120, bottom=120, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, width_pct=100):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_pct * 50))
    tbl_w.set(qn("w:type"), "pct")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_cell_text(cell, bold=False, color=TEXT, size=9.5, align=None):
    for paragraph in cell.paragraphs:
        if align is not None:
            paragraph.alignment = align
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.15
        for run in paragraph.runs:
            run.font.name = "Microsoft YaHei"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            run.font.size = Pt(size)
            run.font.color.rgb = RGBColor.from_string(color)
            run.bold = bold


def add_hyperlink(paragraph, text, url, color="0563C1", underline=True):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    if color:
        c = OxmlElement("w:color")
        c.set(qn("w:val"), color)
        r_pr.append(c)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        r_pr.append(u)
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Microsoft YaHei")
    r_fonts.set(qn("w:hAnsi"), "Microsoft YaHei")
    r_fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r_pr.append(r_fonts)
    new_run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def set_doc_defaults(doc):
    section = doc.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.75)
    section.right_margin = Cm(1.75)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.line_spacing = 1.18
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size, color, before, after in [
        ("Title", 24, ACCENT_DARK, 0, 8),
        ("Subtitle", 11, MUTED, 0, 18),
        ("Heading 1", 17, ACCENT_DARK, 14, 6),
        ("Heading 2", 13, ACCENT, 10, 4),
        ("Heading 3", 11.5, ACCENT_DARK, 8, 3),
    ]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = style_name.startswith("Heading") or style_name == "Title"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(34)
    run = p.add_run("AI 艺术出图项目体系白话总结")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(ACCENT_DARK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ComfyUI / Stable Diffusion / WebUI / LoRA / ControlNet 等工具总览与 5 人工作组方案")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(MUTED)

    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, 88)
    cell = table.cell(0, 0)
    set_cell_shading(cell, ACCENT_LIGHT)
    set_cell_border(cell, color="B7CDD4")
    set_cell_margins(cell, top=220, bottom=220, start=240, end=240)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("一句话：这不是单个软件清单，而是一套从“想法”到“成图交付”的 AI 出图生产线。")
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(ACCENT_DARK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("适用对象：想快速理解 AI 出图工具、搭建小团队流程、制定内部培训和项目交付规范的人。")
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor.from_string(TEXT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("整理日期：2026-05-16")
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor.from_string(MUTED)

    doc.add_page_break()


def add_callout(doc, title, body, fill=ACCENT_LIGHT):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, 100)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, color="B7CDD4")
    set_cell_margins(cell, top=150, bottom=150, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(ACCENT_DARK)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(body)
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(TEXT)


def add_bullets(doc, items, style=None):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.55)
        p.paragraph_format.first_line_indent = Cm(-0.22)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run("• " + item)
        r.font.name = "Microsoft YaHei"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor.from_string(TEXT)


def add_numbered(doc, items):
    for idx, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.first_line_indent = Cm(-0.25)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f"{idx}. {item}")
        r.font.name = "Microsoft YaHei"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor.from_string(TEXT)


def add_table(doc, headers, rows, widths=None, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    set_table_width(table, 100)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, ACCENT)
        set_cell_border(cell, color=ACCENT_DARK)
        set_cell_margins(cell, top=120, bottom=120, start=110, end=110)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.text = text
        style_cell_text(cell, bold=True, color="FFFFFF", size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
        if widths:
            cell.width = widths[i]
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cell = cells[i]
            set_cell_border(cell, color=GRID)
            set_cell_margins(cell, top=115, bottom=115, start=110, end=110)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = text
            align = WD_ALIGN_PARAGRAPH.CENTER if i == 0 and len(headers) > 3 else WD_ALIGN_PARAGRAPH.LEFT
            style_cell_text(cell, size=font_size, align=align)
            if widths:
                cell.width = widths[i]
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_link_table(doc, rows):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.autofit = False
    set_table_width(table, 100)
    headers = ["类别", "项目与用途", "可复制链接"]
    widths = [1350, 3550, 5200]
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, ACCENT)
        set_cell_border(cell, color=ACCENT_DARK)
        set_cell_margins(cell, top=115, bottom=115, start=105, end=105)
        set_cell_width(cell, widths[i])
        cell.text = text
        style_cell_text(cell, bold=True, color="FFFFFF", size=9.2, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_repeat_table_header(table.rows[0])

    for category, name, usage, url in rows:
        cells = table.add_row().cells
        values = [category, f"{name}\n{usage}", ""]
        for i, value in enumerate(values):
            cell = cells[i]
            set_cell_border(cell, color=GRID)
            set_cell_margins(cell, top=110, bottom=110, start=90, end=90)
            set_cell_width(cell, widths[i])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if i != 2:
                cell.text = value
                style_cell_text(cell, size=8.7, align=WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT)
        p = cells[2].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        add_hyperlink(p, url, url)
        for run in p.runs:
            run.font.name = "Microsoft YaHei"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            run.font.size = Pt(7.4)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("AI 艺术出图项目体系白话总结 | 5 人工作组方案")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)


def build_doc():
    doc = Document()
    set_doc_defaults(doc)
    add_footer(doc.sections[0])
    add_cover(doc)

    doc.add_heading("1. 整体说明：这套体系到底是什么", level=1)
    p = doc.add_paragraph()
    p.add_run("AI 艺术出图项目").bold = True
    p.add_run("不是只有一个软件，也不是只会输入提示词。它更像一条小型创意生产线：先有创意方向，再选模型和工具，再控制构图、风格、人物、动作，最后做精修、放大、归档和交付。")
    for run in p.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(10.5)

    add_callout(
        doc,
        "小白版理解",
        "Stable Diffusion 是“会画画的大脑”；WebUI、ComfyUI、Fooocus 是“操作台”；LoRA、ControlNet、IP-Adapter 是“控制工具”；Kohya_ss 是“训练工具”；Hugging Face、Civitai 是“模型和素材市场”。",
    )

    doc.add_heading("核心目标", level=2)
    add_bullets(
        doc,
        [
            "把一个人的试错，变成团队可复用的流程。",
            "把随机出图，变成有风格、有版本、有验收标准的生产。",
            "把“会玩软件”，升级成“能稳定交付图像资产”。",
            "让新手知道先学什么、进阶人员知道怎么搭工作流、管理者知道怎么安排分工。",
        ],
    )

    doc.add_heading("2. 整体架构：从想法到成图的 6 层", level=1)
    rows = [
        ("1", "需求层", "确定要画什么、给谁看、什么风格、什么尺寸。", "需求单、参考图、风格板、交付规格"),
        ("2", "模型层", "决定 AI 的基本画风和能力，像选择画师。", "SD/SDXL/Flux 等大模型、Checkpoint、VAE"),
        ("3", "操作层", "人和模型交互的界面，决定上手难度和流程复杂度。", "WebUI、ComfyUI、Fooocus、InvokeAI、SD.Next"),
        ("4", "控制层", "让画面不再完全随机，控制姿势、构图、人物一致性和参考图。", "ControlNet、IP-Adapter、Img2Img、Inpainting"),
        ("5", "风格层", "把某种人物、产品、品牌风格或画风固定下来。", "LoRA、Embedding、风格提示词、训练数据集"),
        ("6", "交付层", "把图修好、放大、命名、记录参数、交给业务或客户。", "Upscaler、Photoshop/修图、版本管理、素材库"),
    ]
    add_table(doc, ["层级", "名称", "白话解释", "常见工具/资产"], rows, font_size=8.9)

    doc.add_heading("3. 常见项目与工具白话表", level=1)
    tool_rows = [
        ("Stable Diffusion / SD", "AI 出图的底层模型", "像发动机。真正负责把文字或图片变成新图片。", "所有人都要理解，但不一定直接操作底层代码。"),
        ("Checkpoint / 大模型", "主画师", "不同模型像不同画师，有的擅长真人，有的擅长二次元，有的擅长产品图。", "提示词人员、流程工程师"),
        ("VAE", "颜色和质感调节器", "有时换 VAE 会让颜色更舒服、细节更顺。", "进阶人员"),
        ("WebUI / A1111", "最常见的新手操作台", "打开网页，输入提示词，调参数，点生成。插件多，上手资料多。", "新手、提示词人员"),
        ("ComfyUI", "节点式高级工作台", "像搭积木，把模型、提示词、控制、放大、保存一步步连起来。适合团队复用流程。", "流程工程师、技术负责人"),
        ("Fooocus", "傻瓜相机版出图", "隐藏很多复杂参数，适合快速试风格和给新人练手。", "完全新手、创意试图"),
        ("InvokeAI", "偏专业的创作平台", "界面更完整，适合图像生成、编辑和管理。", "设计师、团队创作"),
        ("SD.Next", "增强型 WebUI", "功能多，支持范围广，适合喜欢折腾和扩展的人。", "进阶玩家、技术人员"),
        ("ControlNet", "构图和姿势控制器", "给它线稿、姿势、深度图，AI 就更容易按你的结构画。", "需要精准构图的项目"),
        ("IP-Adapter", "参考图控制器", "让 AI 参考某张图的人脸、风格或构图，再生成新图。", "人物一致性、风格参考"),
        ("LoRA", "小型风格包", "让模型记住某个人物、角色、产品或画风。文件小，加载方便。", "品牌/角色长期项目"),
        ("Kohya_ss", "LoRA 训练工具", "把准备好的图片喂给工具，训练出自己的 LoRA。", "训练和素材管理员"),
        ("AnimateDiff", "静态图变动画", "让文生图模型生成简单动图或短视频。", "短视频、动态素材"),
        ("Real-ESRGAN / Upscaler", "图片放大修复", "把小图变大、变清晰，适合最终交付。", "后期和质检"),
        ("Diffusers", "开发者工具箱", "给程序员用的扩散模型库，适合做产品化、自动化和 API。", "开发者、平台团队"),
        ("Hugging Face / Civitai", "模型素材库", "找模型、LoRA、示例和社区资源的地方。", "全员都会用，但要注意版权和许可"),
    ]
    add_table(doc, ["工具/概念", "它是什么", "小白解释", "最适合谁"], tool_rows, font_size=8.3)

    doc.add_heading("4. 可直接复制的官方/常用链接", level=1)
    doc.add_paragraph("下面保留完整 URL，方便复制到浏览器、群聊、项目文档或 README。链接会随开源项目迁移而变化，建议团队每季度核对一次。")
    link_rows = [
        ("核心模型", "Stability AI Generative Models", "Stable Diffusion 相关官方模型代码与说明", "https://github.com/Stability-AI/generative-models"),
        ("操作台", "ComfyUI", "节点式工作流，适合搭团队生产线", "https://github.com/Comfy-Org/ComfyUI"),
        ("操作台", "Stable Diffusion WebUI / A1111", "经典网页操作台，新手资料多、插件多", "https://github.com/AUTOMATIC1111/stable-diffusion-webui"),
        ("操作台", "Fooocus", "更简单的出图工具，适合快速试图", "https://github.com/lllyasviel/Fooocus"),
        ("操作台", "InvokeAI", "偏专业的创作平台和 WebUI", "https://github.com/invoke-ai/InvokeAI"),
        ("操作台", "SD.Next", "增强型 WebUI，覆盖图片、视频、处理功能", "https://github.com/vladmandic/sdnext"),
        ("控制插件", "ControlNet 原始项目", "控制姿势、边缘、深度、线稿等", "https://github.com/lllyasviel/ControlNet"),
        ("控制插件", "WebUI ControlNet 扩展", "A1111 WebUI 常用 ControlNet 插件", "https://github.com/Mikubill/sd-webui-controlnet"),
        ("参考图", "IP-Adapter", "用参考图控制风格、人物、构图", "https://github.com/tencent-ailab/IP-Adapter"),
        ("训练", "Kohya_ss", "训练 LoRA、微调模型的 GUI/CLI 工具", "https://github.com/bmaltais/kohya_ss"),
        ("动画", "AnimateDiff", "把文生图模型扩展成动画生成", "https://github.com/guoyww/AnimateDiff"),
        ("放大修复", "Real-ESRGAN", "图片/视频放大和修复", "https://github.com/xinntao/Real-ESRGAN"),
        ("开发库", "Hugging Face Diffusers", "开发者使用的扩散模型工具库", "https://github.com/huggingface/diffusers"),
        ("模型平台", "Hugging Face Models", "模型、数据和开源资源平台", "https://huggingface.co/models"),
        ("模型平台", "Civitai", "社区模型、LoRA、样图平台", "https://civitai.com/"),
    ]
    add_link_table(doc, link_rows)

    doc.add_heading("5. 5 人工作组怎么配合", level=1)
    add_callout(
        doc,
        "推荐定位",
        "5 人小组不要每个人都做全部事情。最稳的方式是：1 人管方向，1 人管提示词和风格，1 人管工作流，1 人管模型/训练/素材，1 人管后期质检和交付。",
        fill="F4F7F7",
    )
    role_rows = [
        ("1", "项目负责人 / 美术总监", "接需求、定风格、看最终结果", "需求说明、参考图、验收标准、版本拍板", "不一定亲自跑模型，但必须能判断好坏"),
        ("2", "提示词与风格设计师", "写提示词、整理风格词、做参考图", "Prompt 模板、负面词、风格板、试图记录", "负责把“想要的感觉”翻译成 AI 能懂的话"),
        ("3", "ComfyUI / WebUI 流程工程师", "搭工作流、装插件、做批量生成", "ComfyUI 工作流、节点配置、参数模板", "负责把一次成功变成可重复成功"),
        ("4", "模型与素材管理员", "找模型、管 LoRA、训练 LoRA、管版权", "模型库、LoRA 库、训练数据集、许可记录", "防止模型乱用、版本混乱、侵权风险"),
        ("5", "后期质检与交付", "修图、放大、检查手脸文字、导出交付", "高清成图、修复版、交付包、问题清单", "负责最后一公里，保证能交给客户/业务"),
    ]
    add_table(doc, ["编号", "角色", "主要职责", "产出物", "白话说明"], role_rows, font_size=8.15)

    doc.add_heading("6. 标准工作流：一张图如何从需求走到交付", level=1)
    workflow = [
        "需求确认：负责人明确用途、风格、尺寸、数量、禁区和交付时间。",
        "风格探索：提示词人员用 Fooocus 或 WebUI 快速试 20-50 张方向图。",
        "技术路线：流程工程师判断用 WebUI 快速出，还是用 ComfyUI 搭可复用工作流。",
        "控制生成：需要固定姿势、构图、人物时加入 ControlNet、IP-Adapter、LoRA。",
        "批量产出：用固定参数批量生成，保留 seed、模型、LoRA、提示词和工作流文件。",
        "后期精修：用局部重绘、放大、修脸、修手、去瑕疵，得到可交付版本。",
        "验收归档：负责人验收；交付人员按命名规范打包；素材管理员记录模型来源和授权。",
    ]
    add_numbered(doc, workflow)

    doc.add_heading("7. 工具配合建议", level=1)
    pairing_rows = [
        ("新人练习", "Fooocus + WebUI", "先学提示词、尺寸、种子、负面词，不急着学复杂节点。"),
        ("正式项目出图", "ComfyUI + Checkpoint + LoRA + ControlNet", "把固定流程保存成 workflow，后续同类需求直接复用。"),
        ("人物/品牌一致性", "LoRA + IP-Adapter + Inpainting", "先用 LoRA 锁定风格，再用参考图和局部重绘修细节。"),
        ("构图要求高", "ControlNet + Img2Img", "用线稿、姿势图、草图控制画面，减少随机性。"),
        ("大量素材生产", "ComfyUI 批处理 + 命名规则 + 质检表", "不要只看单张效果，要看批量稳定性和可追溯性。"),
        ("需要训练风格", "Kohya_ss + 数据集规范 + LoRA 版本库", "训练前先清洗数据，训练后做固定测试集对比。"),
        ("最终交付", "Inpainting + Upscaler + 后期软件", "成图不是终点，最终图要经过修复、放大和验收。"),
    ]
    add_table(doc, ["场景", "推荐组合", "为什么这样配合"], pairing_rows, font_size=8.8)

    doc.add_page_break()
    doc.add_heading("8. 团队资料和文件夹规范", level=1)
    add_bullets(
        doc,
        [
            "01_需求：需求单、参考图、尺寸、交付标准。",
            "02_提示词：正向词、负面词、风格词、测试记录。",
            "03_工作流：ComfyUI workflow JSON、WebUI 参数截图、插件版本。",
            "04_模型资产：Checkpoint、LoRA、VAE、Embedding、许可说明。",
            "05_过程图：草稿、试错图、不同 seed 的候选图。",
            "06_精修图：局部重绘、放大、修复过程文件。",
            "07_交付图：最终 JPG/PNG/WebP、客户版、源文件、打包说明。",
            "08_复盘：本次好用参数、失败原因、下次可复用模板。",
        ],
    )

    doc.add_heading("命名建议", level=2)
    add_callout(
        doc,
        "文件命名模板",
        "项目名_用途_风格_模型_版本_日期_序号，例如：BrandA_海报_赛博写实_SDXL_v03_20260516_001.png",
        fill="F8FAFC",
    )

    doc.add_heading("9. 新人培训路线", level=1)
    training_rows = [
        ("第 1 天", "理解概念", "知道 SD、模型、LoRA、WebUI、ComfyUI 分别是什么。", "能说清楚“模型、操作台、插件、训练”的区别。"),
        ("第 2 天", "会用简单工具", "用 Fooocus/WebUI 出图，学会提示词、负面词、尺寸、seed。", "能稳定出 10 张同主题候选图。"),
        ("第 3 天", "会看参数", "理解采样器、步数、CFG、模型、LoRA 权重。", "能复现别人一张图的大致风格。"),
        ("第 4 天", "会控制画面", "学习 ControlNet、参考图、局部重绘。", "能按指定姿势或草图生成画面。"),
        ("第 5 天", "会交付", "放大、修脸修手、命名归档、提交验收。", "能独立完成一套小型交付包。"),
    ]
    add_table(doc, ["阶段", "主题", "训练内容", "验收标准"], training_rows, font_size=8.9)

    doc.add_heading("10. 管理者最该盯的 6 件事", level=1)
    add_bullets(
        doc,
        [
            "风格是否统一：同一项目不要一会儿像摄影、一会儿像二次元、一会儿像油画。",
            "版本是否可追溯：每张重要图都要知道用的模型、LoRA、提示词和 seed。",
            "素材是否合规：商业项目要检查模型和训练素材许可，不要乱用来路不明素材。",
            "流程是否可复用：能沉淀成 workflow 和 prompt 模板的，就不要只靠个人手感。",
            "交付是否稳定：检查手、脸、文字、边缘、分辨率、压缩质量和色彩。",
            "知识是否共享：每周沉淀好用模型、失败案例和参数模板，不让经验只留在个人电脑里。",
        ],
    )

    doc.add_page_break()
    doc.add_heading("11. 推荐落地方案", level=1)
    plan_rows = [
        ("第一阶段：先跑起来", "1-2 周", "统一工具、装好 WebUI/ComfyUI、建立模型库、训练新人。", "能完成基础出图和交付。"),
        ("第二阶段：稳定风格", "2-4 周", "建立 prompt 模板、风格板、LoRA 选择规范、质检表。", "同类图能保持风格一致。"),
        ("第三阶段：流程复用", "1-2 月", "把常用任务做成 ComfyUI 工作流，支持批量生成。", "从“会做一张”变成“稳定做一批”。"),
        ("第四阶段：自有资产", "持续", "训练自有 LoRA，沉淀品牌角色、产品风格和素材库。", "形成团队自己的出图能力。"),
    ]
    add_table(doc, ["阶段", "时间", "重点动作", "目标"], plan_rows, font_size=8.8)

    doc.add_heading("12. 最后总结", level=1)
    p = doc.add_paragraph()
    p.add_run("这套 AI 艺术出图体系的重点，不是追求工具越多越好，而是让工具各司其职。").bold = True
    p.add_run("WebUI 和 Fooocus 适合入门和快速探索；ComfyUI 适合沉淀成团队工作流；ControlNet、IP-Adapter、LoRA 负责把随机性降下来；Kohya_ss 负责训练团队自己的风格资产；后期和质检负责把图真正变成可交付结果。")
    for run in p.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(10.5)

    add_callout(
        doc,
        "给 5 人小组的最终建议",
        "先统一流程，再追求高级效果。先让每个人知道自己负责哪一段，再让团队沉淀模型库、提示词库、工作流库和交付标准。这样 AI 出图才不会停留在“谁手感好谁能做”，而是变成整个小组都能稳定产出的能力。",
    )

    doc.save(OUT_PATH)


if __name__ == "__main__":
    build_doc()
    print(OUT_PATH)
