from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = r"C:\Users\jian\Documents\New project\output\艺术行业效率插件建议.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None, size=10.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table, color="D8DEE8"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_width(cell, cm):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.color.rgb = RGBColor(32, 67, 112) if level == 1 else RGBColor(67, 91, 123)
    run.bold = True
    return p


def add_body(doc, text, bold_start=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.22
    if bold_start and text.startswith(bold_start):
        r1 = p.add_run(bold_start)
        r1.bold = True
        r2 = p.add_run(text[len(bold_start):])
        runs = [r1, r2]
    else:
        runs = [p.add_run(text)]
    for run in runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(10.5)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.18
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10.2)
    return p


def add_note_box(doc, title, body, fill="EEF5FF"):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    set_table_borders(table, "C7D5EA")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(11)
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.color.rgb = RGBColor(32, 67, 112)
    p2 = cell.add_paragraph()
    p2.paragraph_format.line_spacing = 1.2
    p2.paragraph_format.space_after = Pt(2)
    r2 = p2.add_run(body)
    r2.font.size = Pt(10.2)
    r2.font.name = "Microsoft YaHei"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    for style_name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11.5)]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def build_doc():
    doc = Document()
    style_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(5)
    r = title.add_run("艺术行业效率插件建议")
    r.bold = True
    r.font.size = Pt(24)
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.color.rgb = RGBColor(20, 57, 97)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(14)
    sr = subtitle.add_run("白话版：哪些插件最适合艺术工作者，怎么用来省时间")
    sr.font.size = Pt(11.5)
    sr.font.name = "Microsoft YaHei"
    sr._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    sr.font.color.rgb = RGBColor(90, 98, 110)

    add_note_box(
        doc,
        "一句话结论",
        "最值得优先用的是 Canva、Imagegen、PPT、Word 文档、视频制作和表格管理这几类工具。它们能分别帮你做视觉物料、找创意参考、做提案、写项目说明、做短视频和管报价排期。浏览器、人脉搜索、代码管理属于辅助型工具，看你的业务形态再用。",
        fill="EAF2FF",
    )

    add_heading(doc, "一、最值得优先用的插件", 1)
    data = [
        ("Canva", "海报、展览邀请函、作品集封面、社媒图、品牌模板", "把重复设计做成模板，日常发图不必每次从零排版。"),
        ("Imagegen", "概念草图、氛围图、材质参考、海报背景、视觉方向探索", "前期找灵感更快，可以先看几十个方向，再决定真正创作。"),
        ("Presentations", "客户提案、策展方案、品牌合作 PPT、课程课件", "把想法整理成能拿出去讲的页，适合谈合作、报价、汇报。"),
        ("Documents", "报价说明、合同草案、项目说明书、艺术家简介、展览文案", "把零散想法写成正式文件，客户看起来更清楚，也更专业。"),
        ("HyperFrames", "作品短视频、展览预告、动态海报、视频封面动画", "把静态作品变成短视频内容，适合小红书、抖音、Instagram、B站。"),
        ("Spreadsheets", "报价表、成本表、作品库存、客户跟进、展览排期", "把钱、时间、作品数量管清楚，减少漏报、漏发、漏交付。"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.autofit = False
    set_table_borders(table)
    headers = ["插件", "适合做什么", "能帮你省什么"]
    widths = [3.0, 6.4, 6.2]
    for idx, h in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_text(cell, h, bold=True, color="FFFFFF", size=10.2)
        set_cell_shading(cell, "204370")
        set_width(cell, widths[idx])
    for row_idx, row in enumerate(data, start=1):
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, bold=(idx == 0), size=9.6)
            set_width(cells[idx], widths[idx])
            set_cell_shading(cells[idx], "F7FAFD" if row_idx % 2 == 0 else "FFFFFF")

    add_heading(doc, "二、按艺术工作流怎么用", 1)
    add_heading(doc, "1. 灵感和创作前期", 2)
    add_body(doc, "推荐：Imagegen + Browser/Chrome。", bold_start="推荐：")
    add_body(doc, "你可以先用 Imagegen 做不同风格、色彩、构图的方向图，再用 Browser 或 Chrome 查真实案例、展览信息、艺术家资料和市场参考。这样不是让 AI 替你创作，而是让它帮你更快筛选方向。")
    add_bullet(doc, "适合：概念艺术、插画、装置方案、空间陈列、品牌视觉初稿。")
    add_bullet(doc, "省时间点：少刷很多无关参考，先把方向看出来。")

    add_heading(doc, "2. 作品包装和社媒发布", 2)
    add_body(doc, "推荐：Canva + HyperFrames。", bold_start="推荐：")
    add_body(doc, "Canva 适合做稳定的视觉模板，比如作品发布图、展览海报、课程招生图、艺术市集摊位宣传。HyperFrames 适合把图片、文字、音乐、转场做成短视频，让静态作品更容易被看到。")
    add_bullet(doc, "适合：小红书九宫格、Instagram Reels、抖音预告、展览开幕倒计时。")
    add_bullet(doc, "省时间点：一套模板反复用，发内容不会每次都卡在排版上。")

    add_heading(doc, "3. 客户提案和商业合作", 2)
    add_body(doc, "推荐：Presentations + Documents。", bold_start="推荐：")
    add_body(doc, "客户通常不只看作品，还看你能不能把项目讲清楚。Presentations 能做合作方案和提案 PPT，Documents 能做项目说明、报价说明、交付范围、艺术家简介。")
    add_bullet(doc, "适合：品牌联名、空间艺术、商业插画、展览策划、艺术课程。")
    add_bullet(doc, "省时间点：把口头沟通变成正式文件，减少反复解释。")

    add_heading(doc, "4. 报价、排期和库存管理", 2)
    add_body(doc, "推荐：Spreadsheets。", bold_start="推荐：")
    add_body(doc, "艺术工作很容易卡在杂事上：哪件作品卖了、哪张图授权给谁、成本多少、客户尾款有没有收、展览什么时候布撤展。Spreadsheets 可以把这些做成表格和公式，后面直接更新，不用每次重新算。")
    add_bullet(doc, "适合：作品库存表、授权记录、材料成本、项目报价、月度收入支出。")
    add_bullet(doc, "省时间点：少靠记忆，多靠表格提醒和自动计算。")

    add_heading(doc, "5. 人脉和长期机会", 2)
    add_body(doc, "推荐：Happenstance。", bold_start="推荐：")
    add_body(doc, "如果你需要找策展人、品牌市场负责人、画廊、收藏顾问、艺术机构联系人，Happenstance 可以帮你从职业网络里找可能认识的人。它适合找机会，不适合替代真正的人际经营。")
    add_bullet(doc, "适合：找合作入口、约访谈、找展览资源、找品牌对接人。")
    add_bullet(doc, "省时间点：先找到可能的人，再决定怎么联系。")

    add_heading(doc, "三、不同类型艺术从业者的插件组合", 1)
    combos = [
        ("自由插画师 / 视觉设计师", "Imagegen + Canva + PPT + Word + 表格", "先快速出视觉方向，再把报价和交付范围写清楚。"),
        ("艺术家 / 画家 / 摄影师", "Canva + 短视频 + Word + 表格", "重点是作品包装、展览资料、库存和销售记录。"),
        ("策展人 / 艺术机构", "PPT + Word + 表格 + 浏览器", "重点是策展方案、展览排期、预算表和资料调研。"),
        ("艺术教育 / 工作坊", "PPT + Canva + Word + 短视频", "重点是课程课件、招生海报、课程介绍和短视频宣传。"),
        ("数字艺术 / 交互作品", "GitHub + 浏览器 + 短视频 + PPT", "如果作品涉及网页、代码、交互展示等内容，可以管版本，也能测试展示效果。"),
    ]
    combo_table = doc.add_table(rows=1, cols=3)
    combo_table.style = "Table Grid"
    combo_table.autofit = False
    set_table_borders(combo_table)
    for idx, h in enumerate(["你的类型", "建议组合", "使用重点"]):
        cell = combo_table.rows[0].cells[idx]
        set_cell_text(cell, h, bold=True, color="FFFFFF", size=10.2)
        set_cell_shading(cell, "204370")
    combo_widths = [4.1, 6.0, 5.5]
    for row_idx, row in enumerate(combos, start=1):
        cells = combo_table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, bold=(idx == 0), size=9.5)
            set_width(cells[idx], combo_widths[idx])
            set_cell_shading(cells[idx], "FFFFFF" if row_idx % 2 else "F7FAFD")

    add_heading(doc, "四、建议你先从这三件事开始", 1)
    add_body(doc, "不要一口气全用。先做能马上省时间的三件事：")
    add_bullet(doc, "第一，用 Canva 做一套固定模板：作品发布图、展览通知、报价封面、课程宣传。")
    add_bullet(doc, "第二，用 Spreadsheets 做一个“作品 + 客户 + 报价”总表，至少记录作品名、尺寸、价格、状态、客户、收款情况。")
    add_bullet(doc, "第三，用 Documents 和 Presentations 各做一份固定底稿：一个项目说明 Word，一个客户提案 PPT。以后每个项目只改内容，不重做结构。")

    add_note_box(
        doc,
        "使用原则",
        "这些插件最适合处理重复、整理、排版、转格式、做初稿这些事。真正的审美判断、创作取舍、客户关系和最终作品质量，还是要由你来把关。",
        fill="FFF6E6",
    )

    doc.add_page_break()
    add_heading(doc, "五、插件优先级建议", 1)
    priorities = [
        ("第一优先级", "Canva、Word 文档、表格", "每天都可能用，直接解决宣传、文件、报价和管理。"),
        ("第二优先级", "Imagegen、PPT、短视频", "能提高创意探索、商业提案和视频内容效率。"),
        ("第三优先级", "浏览器、人脉搜索、GitHub", "按需要使用。调研、人脉、代码作品或网站项目时很有用。"),
    ]
    pr_table = doc.add_table(rows=1, cols=3)
    pr_table.style = "Table Grid"
    pr_table.autofit = False
    set_table_borders(pr_table)
    for idx, h in enumerate(["优先级", "插件", "理由"]):
        cell = pr_table.rows[0].cells[idx]
        set_cell_text(cell, h, bold=True, color="FFFFFF", size=10.2)
        set_cell_shading(cell, "204370")
    for row_idx, row in enumerate(priorities, start=1):
        cells = pr_table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, bold=(idx == 0), size=9.7)
            set_width(cells[idx], [3.3, 5.0, 7.3][idx])
            set_cell_shading(cells[idx], "FFFFFF" if row_idx % 2 else "F7FAFD")

    add_heading(doc, "六、一个简单的每周用法", 1)
    add_bullet(doc, "周一：用 Spreadsheets 看项目进度、收款、排期。")
    add_bullet(doc, "周二到周三：用 Imagegen 和 Browser 找方向，整理灵感。")
    add_bullet(doc, "周四：用 Canva 做社媒图，用 HyperFrames 做短视频版本。")
    add_bullet(doc, "周五：用 Documents 写项目说明，用 Presentations 准备客户提案。")
    add_bullet(doc, "月底：复盘哪些内容带来咨询，哪些作品更容易成交，再调整模板和报价。")

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("艺术行业效率插件建议 | 白话版")
    fr.font.name = "Microsoft YaHei"
    fr._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(120, 128, 140)

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT)
