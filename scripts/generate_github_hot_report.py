from __future__ import annotations

import json
import math
import re
import statistics
import unicodedata
import urllib.parse
import urllib.request
from collections import Counter
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "output"
SCRATCH_DIR = ROOT / "scratch"
REPORT_DATE = "2026-05-15"
WINDOW_START = "2026-02-15"
WINDOW_END = "2026-05-15"
QUERY = f"created:>={WINDOW_START} fork:false archived:false"
SEARCH_URL = (
    "https://api.github.com/search/repositories?"
    + urllib.parse.urlencode(
        {
            "q": QUERY,
            "sort": "stars",
            "order": "desc",
            "per_page": "100",
        }
    )
)


def fetch_projects() -> list[dict]:
    req = urllib.request.Request(
        SEARCH_URL,
        headers={
            "Accept": "application/vnd.github+json",
            "User-Agent": "Codex-GitHub-Hot-Projects-Report",
            "X-GitHub-Api-Version": "2022-11-28",
        },
    )
    with urllib.request.urlopen(req, timeout=80) as resp:
        payload = json.loads(resp.read().decode("utf-8"))
    items = payload["items"][:100]
    if len(items) < 100:
        raise RuntimeError(f"GitHub returned only {len(items)} repositories")
    return items


def short_date(iso_value: str) -> str:
    return datetime.fromisoformat(iso_value.replace("Z", "+00:00")).strftime("%Y-%m-%d")


def compact_number(value: int) -> str:
    if value >= 100000:
        return f"{value / 10000:.1f}w"
    if value >= 10000:
        return f"{value / 10000:.2f}w"
    return f"{value:,}"


def classify_project(item: dict) -> str:
    text = " ".join(
        [
            item.get("name") or "",
            item.get("full_name") or "",
            item.get("description") or "",
            " ".join(item.get("topics") or []),
        ]
    ).lower()
    checks = [
        ("AI / Agent", ["agent", "ai", "llm", "gpt", "claude", "model", "rag", "prompt", "mcp"]),
        ("开发工具", ["code", "cli", "sdk", "dev", "tool", "debug", "editor", "workflow"]),
        ("前端 / 设计", ["ui", "design", "react", "vue", "css", "component", "figma", "frontend"]),
        ("数据 / 研究", ["data", "research", "benchmark", "analytics", "dataset", "training"]),
        ("基础设施", ["infra", "deploy", "cloud", "kubernetes", "docker", "server", "database"]),
        ("安全", ["security", "auth", "privacy", "encrypt", "vulnerability"]),
        ("效率 / 个人工具", ["productivity", "notes", "career", "job", "email", "calendar", "todo"]),
        ("教程 / 资源", ["awesome", "course", "guide", "learn", "tutorial", "template", "list"]),
        ("创作 / 多媒体", ["video", "image", "audio", "game", "music", "creative"]),
    ]
    for label, keywords in checks:
        if any(keyword in text for keyword in keywords):
            return label
    return "其他"


def sanitize_text(value: str) -> str:
    value = re.sub(r"https?://\S+", "链接", value or "")
    value = " ".join(value.split())
    cleaned = []
    for ch in value:
        category = unicodedata.category(ch)
        if category.startswith("C") or category == "So":
            continue
        cleaned.append(ch)
    return "".join(cleaned).strip()


def make_takeaway(item: dict, category: str) -> str:
    desc = (item.get("description") or "暂无公开简介").strip()
    desc = sanitize_text(desc)
    if len(desc) > 78:
        desc = desc[:75].rstrip() + "..."
    language = item.get("language") or "多语言/未标注"
    return f"{desc} 主要看点：{category}方向，{language}生态，短期关注度上升快。"


def make_short_summary(item: dict) -> str:
    desc = sanitize_text(item.get("description") or "暂无公开简介")
    if len(desc) > 64:
        return desc[:61].rstrip() + "..."
    return desc


def set_run_font(run, size: int | float | None = None, bold: bool | None = None, color: str | None = None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_font(paragraph, size: int | float = 9, color: str | None = None):
    for run in paragraph.runs:
        set_run_font(run, size=size, color=color)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=80, bottom=80, end=80):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text: str, size: int | float = 8.2, bold: bool = False, align=None, color: str | None = None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_column_widths(table, widths_inches: list[float]):
    for row in table.rows:
        for idx, width in enumerate(widths_inches):
            row.cells[idx].width = Inches(width)


def add_hyperlink(paragraph, url: str, text: str, size: int | float = 8.2):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1D4ED8")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    font = OxmlElement("w:rFonts")
    font.set(qn("w:ascii"), "Microsoft YaHei")
    font.set(qn("w:hAnsi"), "Microsoft YaHei")
    font.set(qn("w:eastAsia"), "Microsoft YaHei")
    r_pr.append(font)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    r_pr.append(sz)
    new_run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_small_meta(doc: Document, label: str, value: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(label)
    set_run_font(r1, size=9, bold=True, color="334155")
    r2 = p.add_run(value)
    set_run_font(r2, size=9, color="334155")


def build_document(projects: list[dict]) -> Path:
    OUTPUT_DIR.mkdir(exist_ok=True)
    SCRATCH_DIR.mkdir(exist_ok=True)
    json_path = SCRATCH_DIR / f"github-hot-projects-top100-{REPORT_DATE}.json"
    json_path.write_text(
        json.dumps(
            {
                "generated_at": datetime.now(timezone.utc).isoformat(),
                "window_start": WINDOW_START,
                "window_end": WINDOW_END,
                "methodology": "GitHub Search API, created within the last three months, fork:false, archived:false, sorted by stars desc.",
                "query_url": SEARCH_URL,
                "items": projects,
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )

    enriched = []
    for idx, item in enumerate(projects, start=1):
        category = classify_project(item)
        enriched.append(
            {
                "rank": idx,
                "full_name": item["full_name"],
                "url": item["html_url"],
                "stars": item["stargazers_count"],
                "forks": item["forks_count"],
                "language": item.get("language") or "未标注",
                "created": short_date(item["created_at"]),
                "category": category,
                "summary": make_takeaway(item, category),
                "short_summary": make_short_summary(item),
            }
        )

    stars = [p["stars"] for p in enriched]
    forks = [p["forks"] for p in enriched]
    language_counts = Counter(p["language"] for p in enriched)
    category_counts = Counter(p["category"] for p in enriched)

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.1)
    section.bottom_margin = Cm(1.1)
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(9)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("GitHub 最近三个月爆火项目观察报告")
    set_run_font(run, size=21, bold=True, color="0F172A")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(14)
    r = subtitle.add_run(f"Top 100 新兴仓库 | {WINDOW_START} 至 {WINDOW_END} | 按 star 数降序")
    set_run_font(r, size=11, color="475569")

    meta = doc.add_table(rows=2, cols=4)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.autofit = False
    set_column_widths(meta, [2.3, 4.0, 2.3, 4.0])
    meta_cells = [
        ("数据源", "GitHub Search API"),
        ("抓取日期", REPORT_DATE),
        ("筛选条件", QUERY),
        ("结果数量", "100 个仓库"),
    ]
    for i, (label, value) in enumerate(meta_cells):
        row = i // 2
        col = (i % 2) * 2
        set_cell_text(meta.cell(row, col), label, 8.5, True, WD_ALIGN_PARAGRAPH.CENTER, "334155")
        set_cell_text(meta.cell(row, col + 1), value, 8.5, False, WD_ALIGN_PARAGRAPH.LEFT, "334155")
        set_cell_shading(meta.cell(row, col), "E2E8F0")
        set_cell_shading(meta.cell(row, col + 1), "F8FAFC")

    doc.add_paragraph()
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(4)
    h.paragraph_format.space_after = Pt(5)
    hr = h.add_run("一、阅读口径")
    set_run_font(hr, size=13, bold=True, color="0F172A")

    bullets = [
        "本报告把“最近三个月爆火”定义为：在窗口期内新创建、非 fork、非 archived，并按当前 star 数排序的仓库。",
        "这个口径更偏向发现新兴项目，不等同于 GitHub 官方 Trending，也不代表过去三个月的净增 star 排名。",
        "描述、语言、stars、forks、创建时间等字段来自 GitHub API；分类和中文看点为基于仓库公开信息的辅助归纳。",
    ]
    for text in bullets:
        p = doc.add_paragraph(style=None)
        p.style = doc.styles["List Bullet"]
        p.paragraph_format.space_after = Pt(2)
        rr = p.add_run(text)
        set_run_font(rr, size=9.2, color="334155")

    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(8)
    h.paragraph_format.space_after = Pt(5)
    hr = h.add_run("二、总体观察")
    set_run_font(hr, size=13, bold=True, color="0F172A")

    insight_table = doc.add_table(rows=2, cols=4)
    insight_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    insight_table.autofit = False
    set_column_widths(insight_table, [2.5, 3.4, 2.5, 3.4])
    top = enriched[0]
    insights = [
        ("最高 star", f"{top['full_name']}：{top['stars']:,}"),
        ("Top100 中位 star", f"{int(statistics.median(stars)):,}"),
        ("Top100 总 forks", f"{sum(forks):,}"),
        ("AI/Agent 占比", f"{category_counts.get('AI / Agent', 0)} / 100"),
    ]
    for i, (label, value) in enumerate(insights):
        row = i // 2
        col = (i % 2) * 2
        set_cell_text(insight_table.cell(row, col), label, 8.5, True, WD_ALIGN_PARAGRAPH.CENTER, "0F172A")
        set_cell_text(insight_table.cell(row, col + 1), value, 8.5, False, WD_ALIGN_PARAGRAPH.LEFT, "334155")
        set_cell_shading(insight_table.cell(row, col), "DBEAFE")
        set_cell_shading(insight_table.cell(row, col + 1), "F8FAFC")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    lang_text = "、".join([f"{lang} {count}" for lang, count in language_counts.most_common(6)])
    cat_text = "、".join([f"{cat} {count}" for cat, count in category_counts.most_common(6)])
    r = p.add_run(f"语言分布靠前：{lang_text}。领域分布靠前：{cat_text}。")
    set_run_font(r, size=9.2, color="334155")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(
        "整体上，AI Agent、开发者工作流、设计系统、研究自动化和效率工具仍是最明显的新增爆点；"
        "高 forks 项目通常意味着模板、配置、资源集合或可直接部署工具更容易形成二次传播。"
    )
    set_run_font(r, size=9.2, color="334155")

    doc.add_page_break()
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(8)
    h.paragraph_format.space_after = Pt(5)
    hr = h.add_run("三、Top 10 快览")
    set_run_font(hr, size=13, bold=True, color="0F172A")

    top10 = doc.add_table(rows=1, cols=5)
    top10.alignment = WD_TABLE_ALIGNMENT.CENTER
    top10.autofit = False
    set_column_widths(top10, [0.45, 2.85, 0.95, 1.05, 5.95])
    for cell, text in zip(top10.rows[0].cells, ["#", "项目", "Stars", "领域", "一句话看点"]):
        set_cell_text(cell, text, 8.5, True, WD_ALIGN_PARAGRAPH.CENTER, "FFFFFF")
        set_cell_shading(cell, "1E293B")
    set_repeat_table_header(top10.rows[0])
    for pinfo in enriched[:10]:
        row = top10.add_row()
        values = [
            str(pinfo["rank"]),
            pinfo["full_name"],
            compact_number(pinfo["stars"]),
            pinfo["category"],
            pinfo["short_summary"],
        ]
        for i, value in enumerate(values):
            align = WD_ALIGN_PARAGRAPH.CENTER if i in (0, 2, 3) else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(row.cells[i], value, 7.6 if i == 4 else 8.0, False, align, "334155")
            set_cell_shading(row.cells[i], "F8FAFC" if pinfo["rank"] % 2 else "F1F5F9")

    doc.add_page_break()
    h = doc.add_paragraph()
    h.paragraph_format.space_after = Pt(5)
    hr = h.add_run("四、完整 Top 100 清单")
    set_run_font(hr, size=13, bold=True, color="0F172A")

    headers = ["#", "项目", "Stars", "Forks", "语言", "创建", "领域", "简介"]
    widths = [0.36, 2.18, 0.65, 0.62, 0.76, 0.74, 0.92, 4.25]
    for start in range(0, len(enriched), 10):
        if start:
            doc.add_page_break()
        segment = enriched[start : start + 10]
        segment_title = doc.add_paragraph()
        segment_title.paragraph_format.space_after = Pt(4)
        segment_title.paragraph_format.space_before = Pt(2)
        sr = segment_title.add_run(f"第 {segment[0]['rank']} - {segment[-1]['rank']} 名")
        set_run_font(sr, size=10.5, bold=True, color="334155")

        table = doc.add_table(rows=1, cols=8)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        set_column_widths(table, widths)
        for cell, text in zip(table.rows[0].cells, headers):
            set_cell_text(cell, text, 7.8, True, WD_ALIGN_PARAGRAPH.CENTER, "FFFFFF")
            set_cell_shading(cell, "0F172A")
        set_repeat_table_header(table.rows[0])

        for pinfo in segment:
            row = table.add_row()
            shade = "FFFFFF" if pinfo["rank"] % 2 else "F8FAFC"
            row_values = [
                str(pinfo["rank"]),
                None,
                compact_number(pinfo["stars"]),
                compact_number(pinfo["forks"]),
                pinfo["language"],
                pinfo["created"],
                pinfo["category"],
                pinfo["short_summary"],
            ]
            for i, value in enumerate(row_values):
                cell = row.cells[i]
                if value is None:
                    cell.text = ""
                    paragraph = cell.paragraphs[0]
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    paragraph.paragraph_format.space_after = Pt(0)
                    add_hyperlink(paragraph, pinfo["url"], pinfo["full_name"], 7.3)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    set_cell_margins(cell, top=56, start=64, bottom=56, end=64)
                else:
                    align = WD_ALIGN_PARAGRAPH.CENTER if i in (0, 2, 3, 4, 5, 6) else WD_ALIGN_PARAGRAPH.LEFT
                    set_cell_text(cell, value, 7.1 if i == 7 else 7.3, False, align, "334155")
                    set_cell_margins(cell, top=56, start=64, bottom=56, end=64)
                set_cell_shading(cell, shade)

    doc.add_paragraph()
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(8)
    h.paragraph_format.space_after = Pt(4)
    hr = h.add_run("五、数据来源")
    set_run_font(hr, size=13, bold=True, color="0F172A")

    add_small_meta(doc, "GitHub API 查询：", SEARCH_URL)
    add_small_meta(doc, "原始数据文件：", str(json_path))
    add_small_meta(doc, "说明：", "仓库 star 和 fork 数会持续变化，本报告反映抓取时点的公开数据。")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = footer.add_run("GitHub 热门新兴项目 Top 100 | 数据抓取日期 2026-05-15")
    set_run_font(rr, size=8, color="64748B")

    docx_path = OUTPUT_DIR / f"github-hot-projects-top100-{REPORT_DATE}.docx"
    doc.save(docx_path)
    return docx_path


def main() -> None:
    projects = fetch_projects()
    path = build_document(projects)
    print(path)


if __name__ == "__main__":
    main()
