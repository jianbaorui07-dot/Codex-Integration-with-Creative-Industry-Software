from __future__ import annotations

import json
import re
import statistics
import unicodedata
import urllib.parse
import urllib.request
from collections import Counter
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DESKTOP = Path.home() / "Desktop"
SCRATCH_DIR = ROOT / "scratch"
REPORT_DATE = "2026-05-15"
WINDOW_START = "2026-02-15"
WINDOW_END = "2026-05-15"
BASE_DATA_PATH = SCRATCH_DIR / f"github-hot-projects-top100-{REPORT_DATE}.json"
QUERY = f"created:>={WINDOW_START} fork:false archived:false"
SEARCH_URL = (
    "https://api.github.com/search/repositories?"
    + urllib.parse.urlencode(
        {
            "q": QUERY,
            "sort": "stars",
            "order": "desc",
            "per_page": "100",
        }
    )
)


def fetch_projects() -> list[dict]:
    if BASE_DATA_PATH.exists():
        payload = json.loads(BASE_DATA_PATH.read_text(encoding="utf-8"))
        items = payload.get("items", [])[:100]
        if len(items) == 100:
            return items

    req = urllib.request.Request(
        SEARCH_URL,
        headers={
            "Accept": "application/vnd.github+json",
            "User-Agent": "Codex-GitHub-Hot-Projects-CN-Report",
            "X-GitHub-Api-Version": "2022-11-28",
        },
    )
    with urllib.request.urlopen(req, timeout=80) as resp:
        payload = json.loads(resp.read().decode("utf-8"))
    items = payload["items"][:100]
    if len(items) < 100:
        raise RuntimeError(f"GitHub returned only {len(items)} repositories")
    return items


def sanitize_text(value: str | None) -> str:
    value = re.sub(r"https?://\S+", "链接", value or "")
    value = " ".join(value.split())
    cleaned = []
    for ch in value:
        category = unicodedata.category(ch)
        if category.startswith("C") or category == "So":
            continue
        cleaned.append(ch)
    return "".join(cleaned).strip()


def short_date(iso_value: str) -> str:
    return datetime.fromisoformat(iso_value.replace("Z", "+00:00")).strftime("%Y-%m-%d")


def compact_number(value: int) -> str:
    if value >= 100000:
        return f"{value / 10000:.1f}w"
    if value >= 10000:
        return f"{value / 10000:.2f}w"
    return f"{value:,}"


def project_blob(item: dict) -> str:
    return " ".join(
        [
            item.get("name") or "",
            item.get("full_name") or "",
            item.get("description") or "",
            " ".join(item.get("topics") or []),
        ]
    ).lower()


def classify_project(item: dict) -> str:
    text = project_blob(item)
    checks = [
        ("AI / Agent", ["agent", "ai", "llm", "gpt", "claude", "model", "rag", "prompt", "mcp", "skill"]),
        ("开发工具", ["code", "cli", "sdk", "dev", "debug", "editor", "workflow"]),
        ("前端 / 设计", ["ui", "design", "react", "vue", "css", "component", "figma", "frontend"]),
        ("数据 / 研究", ["data", "research", "benchmark", "analytics", "dataset", "training", "experiment"]),
        ("基础设施", ["infra", "deploy", "cloud", "kubernetes", "docker", "server", "database", "vite"]),
        ("安全", ["security", "auth", "privacy", "encrypt", "vulnerability", "cybersecurity"]),
        ("效率 / 个人工具", ["productivity", "notes", "career", "job", "email", "calendar", "todo"]),
        ("教程 / 资源", ["awesome", "course", "guide", "learn", "tutorial", "template", "list"]),
        ("创作 / 多媒体", ["video", "image", "audio", "game", "music", "creative", "render"]),
    ]
    for label, keywords in checks:
        if any(keyword in text for keyword in keywords):
            return label
    return "其他"


def trim_cn(text: str, limit: int = 82) -> str:
    text = sanitize_text(text)
    if len(text) <= limit:
        return text
    return text[: limit - 3].rstrip() + "..."


SPECIAL_WHAT = {
    "ultraworkers/claw-code": "一个热度极高的开发工具仓库，公开简介偏传播，实际能力建议以 README 和示例为准。",
    "garrytan/gstack": "一套 Claude Code/AI 编程助手工作流配置，把设计、工程、发布、文档、QA 等角色做成工具。",
    "karpathy/autoresearch": "用 AI Agent 自动跑研究实验，围绕单卡训练和 nanochat 训练流程做自动化。",
    "VoltAgent/awesome-design-md": "收集 DESIGN.md 设计规范，让 AI 编码助手按品牌和设计系统生成一致界面。",
    "paperclipai/paperclip": "用于管理工作中多个 AI Agent 的开源应用，帮助集中查看和调度代理任务。",
    "JuliusBrussee/caveman": "把 Claude Code 的复杂说明压成极简、低 token 的技能式提示，减少沟通成本。",
    "MemPalace/mempalace": "开源 AI 记忆系统，用于给 Agent 保存长期上下文、偏好和任务线索。",
    "safishamsi/graphify": "给 Claude、Codex、Cursor 等准备的技能，把代码、SQL、文档转成可理解图谱。",
    "chenglou/pretext": "文本测量与布局工具，帮助前端精确计算文字宽高和排版结果。",
    "santifer/career-ops": "基于 Claude Code 的求职自动化系统，用来组织岗位搜索、材料生成和投递流程。",
    "addyosmani/agent-skills": "面向 AI 编程助手的工程技能库，覆盖实现、测试、发布、审查等生产级任务。",
    "nexu-io/open-design": "本地优先的设计助手/设计系统替代方案，目标是让 AI 生成更一致的产品界面。",
    "HKUDS/CLI-Anything": "把软件能力包装成 Agent 可调用的命令行入口，让不同工具更容易被代理操作。",
    "Gitlawb/openclaude": "Claude/Agent 类工具的开放实现或运行壳，重点是让代理能在更多环境中运行。",
    "googleworkspace/cli": "Google Workspace 命令行工具，可在终端里操作 Drive、Gmail 等办公服务。",
    "openai/symphony": "把项目任务拆成隔离的自治执行单元，让多个 Agent 并行推进工作。",
    "openai/codex-plugin-cc": "让 Claude Code 调用 Codex 做代码审查或任务委派的插件。",
    "heygen-com/hyperframes": "用 HTML 写视频和动效，再渲染成可交付视频，面向 Agent 自动生产内容。",
}


SPECIAL_USE = {
    "ultraworkers/claw-code": "先在小仓库试核心命令或 API，确认真实能力后再考虑放入团队工具链。",
    "garrytan/gstack": "把其中的角色模板复制到自己的 AI 助手配置里，按团队流程改成专属工作流。",
    "karpathy/autoresearch": "给它一个明确研究问题，先复现示例实验，再把结果当研究草稿人工校验。",
    "VoltAgent/awesome-design-md": "挑一个接近你产品的 DESIGN.md，放进项目根目录，让 Codex 生成 UI 时遵守。",
    "paperclipai/paperclip": "适合先用来集中管理几个日常 Agent 任务，比如研发、文档、运营自动化。",
    "JuliusBrussee/caveman": "把你的长提示改写成更短的规则，测试是否能降低 token 和上下文干扰。",
    "MemPalace/mempalace": "接入自己的 Agent 流程，保存用户偏好、任务记录和长期项目背景。",
    "safishamsi/graphify": "在复杂仓库里跑一遍，让 AI 先理解结构图，再让它改代码或写文档。",
    "chenglou/pretext": "用于前端组件或编辑器场景，先拿它验证文字测量和布局是否更稳定。",
    "santifer/career-ops": "拿来整理岗位和生成投递草稿，但简历、邮件和匹配结论一定人工复核。",
    "addyosmani/agent-skills": "把常用工程 skill 导入 Codex/Claude，作为写代码、测试和审查时的标准动作。",
    "nexu-io/open-design": "用它做本地设计工作台或规范参考，再把生成结果接入你的前端项目。",
    "HKUDS/CLI-Anything": "先把一个常用桌面或网页工具接成 CLI，再交给 Agent 做重复操作。",
    "Gitlawb/openclaude": "在隔离环境试跑，观察它是否适合替代或补充你的 Claude Code 工作流。",
    "googleworkspace/cli": "连接测试账号后做低风险办公自动化，比如查文件、整理邮件或批量操作。",
    "openai/symphony": "把一个大需求拆成多个子任务，验证它的隔离执行和结果汇总是否顺手。",
    "openai/codex-plugin-cc": "在 Claude Code 里把代码审查、排查或子任务交给 Codex 并回收结果。",
    "heygen-com/hyperframes": "用它把脚本和网页组件变成产品演示视频，适合做自动化内容生产。",
}


def chinese_only_snippet(desc: str, limit: int = 42) -> str:
    cjk_count = sum(1 for ch in desc if "\u4e00" <= ch <= "\u9fff")
    if cjk_count < 10:
        return ""
    cleaned = sanitize_text(desc)
    return trim_cn(cleaned, limit)


def explain_what(item: dict, category: str) -> str:
    text = project_blob(item)
    name = item["full_name"]

    if name in SPECIAL_WHAT:
        return SPECIAL_WHAT[name]

    desc_cn = chinese_only_snippet(item.get("description") or "")
    if desc_cn:
        return f"项目公开说明大意：{desc_cn}"

    rules = [
        (["manage agents", "agents at work"], "Agent 管理应用，用来集中查看、组织和调度工作里的多个 AI 代理。"),
        (["memory", "mempalace"], "AI 记忆层项目，用来保存长期上下文、用户偏好和任务历史。"),
        (["design.md", "design system", "design"], "设计规范/界面约束项目，让 AI 生成 UI 时更符合品牌和产品风格。"),
        (["skill", "skills"], "AI 编程助手技能库，整理提示词、命令、检查清单或工作流模板。"),
        (["cli", "command-line", "command line"], "命令行工具项目，把某类软件能力包装成终端命令或 Agent 可调用接口。"),
        (["career", "job search"], "求职自动化项目，帮助管理岗位搜索、材料生成、投递和进度追踪。"),
        (["research", "experiment", "benchmark"], "研究/实验自动化项目，帮助跑资料整理、实验流程、基准测试或结果记录。"),
        (["browser", "scrap", "web automation"], "浏览器自动化项目，用于网页操作、资料抓取、页面测试或在线任务代理。"),
        (["video", "render", "hyperframes"], "代码化视频/动效项目，把 HTML、脚本或组件渲染成演示内容。"),
        (["security", "cybersecurity"], "安全辅助项目，提供技能、知识库、检查清单或分析流程。"),
        (["workspace", "gmail", "drive", "google"], "办公自动化项目，围绕 Drive、Gmail 等协作工具做命令行或 Agent 操作。"),
        (["openclaw", "claude code", "opencode", "coding agent", "codex"], "AI 编程代理项目，围绕代码生成、审查、任务分派或运行环境搭建。"),
        (["training", "course", "learn"], "学习资料或训练材料项目，适合当教程、模板和最佳实践参考。"),
    ]
    for keywords, explanation in rules:
        if any(keyword in text for keyword in keywords):
            return explanation

    category_hint = {
        "AI / Agent": "AI/Agent 项目，用来增强智能体的编码、研究、记忆、浏览或自动化能力。",
        "开发工具": "开发者工具项目，主要帮助写代码、调试、生成配置或优化工程流程。",
        "前端 / 设计": "前端或设计项目，主要帮助生成界面、规范 UI 或复用视觉系统。",
        "数据 / 研究": "数据或研究项目，主要用于实验、基准、资料整理或模型探索。",
        "基础设施": "基础设施项目，主要帮助部署、服务编排、本地运行或工程环境管理。",
        "安全": "安全项目，主要用于检查、知识整理、权限或风险分析。",
        "效率 / 个人工具": "效率工具，主要帮助个人或团队减少重复操作、整理信息或推进工作流。",
        "教程 / 资源": "教程或资源集合，把经验、模板和参考资料集中整理。",
        "创作 / 多媒体": "创作或多媒体项目，主要帮助生成图像、视频、音频或交互内容。",
        "其他": "近期关注度较高的新项目，具体功能需要结合 README 和示例判断。",
    }[category]
    return category_hint


def explain_use(item: dict, category: str) -> str:
    text = project_blob(item)
    name = item["full_name"]

    if name in SPECIAL_USE:
        return SPECIAL_USE[name]

    if "design" in text:
        return "把它的规范或示例放进你的项目，作为 Codex/Claude 生成页面时的设计约束。"
    if "skill" in text:
        return "挑与你工作流相近的 skill/提示模板，复制到自己的 AI 助手配置里再按项目改。"
    if "memory" in text:
        return "接到自己的 Agent 流程里，先存任务记录和偏好，再测试它能否稳定召回。"
    if "cli" in text or "command-line" in text or "command line" in text:
        return "先按 README 本地安装 CLI，在测试仓库试跑常用命令，再接入日常脚本。"
    if "career" in text or "job" in text:
        return "用它整理岗位、生成投递材料或追踪进度，但简历和邮件建议人工复核。"
    if "research" in text or "experiment" in text or "benchmark" in text:
        return "给它一个明确研究问题，先复现示例流程，再把结论当草稿做人工校验。"
    if "browser" in text or "scrap" in text:
        return "拿来做网页自动化原型，比如自动浏览、抓资料、填表或测试页面流程。"
    if "video" in text or "render" in text:
        return "用它把脚本、网页或组件转成视频/动效内容，适合做演示和自动化素材生产。"
    if "security" in text or "cybersecurity" in text:
        return "作为安全检查清单或知识库参考，落地前要结合你的系统做人工审计。"
    if "google" in text or "workspace" in text or "gmail" in text or "drive" in text:
        return "连接办公账号后做低风险自动化，比如查资料、整理文件或批量处理日常任务。"

    defaults = {
        "AI / Agent": "把它当作 Agent 能力样板，先跑 demo，再挑可复用的提示、插件或流程接进自己的项目。",
        "开发工具": "在小项目里试用它的核心命令或 API，确认稳定后再纳入团队工具链。",
        "前端 / 设计": "用它快速生成界面原型、组件规范或设计参考，再按你的产品风格微调。",
        "数据 / 研究": "先复现实验或数据处理流程，再把代码/指标迁移到自己的研究问题上。",
        "基础设施": "先用本地或测试环境验证部署方式，再评估是否能进入生产工具链。",
        "安全": "把它当作辅助检查工具，不直接替代人工安全评审和权限控制。",
        "效率 / 个人工具": "用它处理一个重复任务小场景，跑顺后再扩展到完整个人工作流。",
        "教程 / 资源": "把它当资料库阅读，摘出模板、清单和最佳实践放到你的项目文档里。",
        "创作 / 多媒体": "先用示例素材跑通生成流程，再替换成你的品牌、脚本或内容素材。",
        "其他": "先看 README、issues 和示例，判断它是否解决你当前的具体痛点。",
    }
    return defaults[category]


def set_run_font(run, size: int | float | None = None, bold: bool | None = None, color: str | None = None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=70, start=70, bottom=70, end=70):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text: str, size: int | float = 7.0, bold: bool = False, align=None, color: str = "334155"):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.02
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)


def set_column_widths(table, widths_inches: list[float]):
    for row in table.rows:
        for index, width in enumerate(widths_inches):
            row.cells[index].width = Inches(width)


def add_hyperlink(paragraph, url: str, text: str, size: int | float = 7.0):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1D4ED8")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    font = OxmlElement("w:rFonts")
    font.set(qn("w:ascii"), "Microsoft YaHei")
    font.set(qn("w:hAnsi"), "Microsoft YaHei")
    font.set(qn("w:eastAsia"), "Microsoft YaHei")
    r_pr.append(font)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    r_pr.append(sz)
    new_run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def build_doc(projects: list[dict]) -> Path:
    SCRATCH_DIR.mkdir(exist_ok=True)
    data_path = SCRATCH_DIR / f"github-hot-projects-top100-cn-{REPORT_DATE}.json"
    enriched = []
    for index, item in enumerate(projects, start=1):
        category = classify_project(item)
        enriched.append(
            {
                "rank": index,
                "full_name": item["full_name"],
                "url": item["html_url"],
                "stars": item["stargazers_count"],
                "forks": item["forks_count"],
                "language": item.get("language") or "未标注",
                "created": short_date(item["created_at"]),
                "category": category,
                "what": explain_what(item, category),
                "use": explain_use(item, category),
            }
        )

    data_path.write_text(
        json.dumps(
            {
                "generated_at": datetime.now(timezone.utc).isoformat(),
                "query_url": SEARCH_URL,
                "items": enriched,
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.05)
    section.bottom_margin = Cm(1.05)
    section.left_margin = Cm(0.9)
    section.right_margin = Cm(0.9)

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("GitHub 最近三个月爆火项目 Top 100：中文说明版")
    set_run_font(run, size=20, bold=True, color="0F172A")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(12)
    run = subtitle.add_run(f"{WINDOW_START} 至 {WINDOW_END} | 新建仓库 | 按 star 排序 | 补充“干啥的 / 咋用”")
    set_run_font(run, size=10.5, color="475569")

    summary = doc.add_table(rows=2, cols=4)
    summary.alignment = WD_TABLE_ALIGNMENT.CENTER
    summary.autofit = False
    set_column_widths(summary, [2.0, 3.8, 2.0, 3.8])
    metadata = [
        ("数据源", "GitHub Search API"),
        ("抓取日期", REPORT_DATE),
        ("筛选条件", QUERY),
        ("项目数量", "100 个"),
    ]
    for index, (label, value) in enumerate(metadata):
        row = index // 2
        col = (index % 2) * 2
        set_cell_text(summary.cell(row, col), label, 8.4, True, WD_ALIGN_PARAGRAPH.CENTER, "0F172A")
        set_cell_text(summary.cell(row, col + 1), value, 8.4, False, WD_ALIGN_PARAGRAPH.LEFT, "334155")
        set_cell_shading(summary.cell(row, col), "DBEAFE")
        set_cell_shading(summary.cell(row, col + 1), "F8FAFC")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("怎么读")
    set_run_font(run, size=12.5, bold=True, color="0F172A")

    notes = [
        "“是干啥的”基于仓库名、公开简介、语言和 topics 做中文归纳；不是官方中文文档。",
        "“你可以咋用”给的是上手场景：先跑 README/demo，再把可复用的提示、CLI、模板或流程迁入自己的项目。",
        "Star 和 fork 会持续变化，本版反映生成时点的公开数据。",
    ]
    for note in notes:
        para = doc.add_paragraph(style="List Bullet")
        para.paragraph_format.space_after = Pt(1)
        run = para.add_run(note)
        set_run_font(run, size=8.6, color="334155")

    language_counts = Counter(item["language"] for item in enriched)
    category_counts = Counter(item["category"] for item in enriched)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(
        "快速观察："
        f"Top100 中位 star {int(statistics.median([p['stars'] for p in enriched])):,}；"
        f"语言靠前为 {', '.join(f'{k} {v}' for k, v in language_counts.most_common(5))}；"
        f"领域靠前为 {', '.join(f'{k} {v}' for k, v in category_counts.most_common(5))}。"
    )
    set_run_font(run, size=8.6, color="334155")

    doc.add_page_break()
    widths = [0.35, 1.6, 0.65, 0.68, 0.82, 3.42, 3.42]
    headers = ["#", "项目", "Stars", "语言", "领域", "是干啥的", "你可以咋用"]

    for start in range(0, len(enriched), 6):
        if start:
            doc.add_page_break()
        segment = enriched[start : start + 6]
        h = doc.add_paragraph()
        h.paragraph_format.space_after = Pt(4)
        run = h.add_run(f"第 {segment[0]['rank']} - {segment[-1]['rank']} 名")
        set_run_font(run, size=12, bold=True, color="0F172A")

        table = doc.add_table(rows=1, cols=7)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        set_column_widths(table, widths)
        for cell, text in zip(table.rows[0].cells, headers):
            set_cell_text(cell, text, 7.6, True, WD_ALIGN_PARAGRAPH.CENTER, "FFFFFF")
            set_cell_shading(cell, "0F172A")

        for item in segment:
            row = table.add_row()
            shade = "FFFFFF" if item["rank"] % 2 else "F8FAFC"
            values = [
                str(item["rank"]),
                None,
                compact_number(item["stars"]),
                item["language"],
                item["category"],
                item["what"],
                item["use"],
            ]
            for index, value in enumerate(values):
                cell = row.cells[index]
                if value is None:
                    cell.text = ""
                    paragraph = cell.paragraphs[0]
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    paragraph.paragraph_format.space_after = Pt(0)
                    add_hyperlink(paragraph, item["url"], item["full_name"], 6.9)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    set_cell_margins(cell)
                else:
                    align = WD_ALIGN_PARAGRAPH.CENTER if index in (0, 2, 3, 4) else WD_ALIGN_PARAGRAPH.LEFT
                    set_cell_text(cell, value, 7.2 if index in (5, 6) else 7.2, False, align, "334155")
                set_cell_shading(cell, shade)

    doc.add_page_break()
    h = doc.add_paragraph()
    run = h.add_run("数据来源")
    set_run_font(run, size=12.5, bold=True, color="0F172A")
    for label, value in [
        ("GitHub API 查询：", SEARCH_URL),
        ("原始数据文件：", str(data_path)),
        ("原始 Word：", str(DESKTOP / "github-hot-projects-top100-2026-05-15.docx")),
    ]:
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(1)
        r1 = para.add_run(label)
        set_run_font(r1, size=8.6, bold=True, color="334155")
        r2 = para.add_run(value)
        set_run_font(r2, size=8.6, color="334155")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("GitHub 爆火项目中文说明版 | 数据抓取日期 2026-05-15")
    set_run_font(run, size=7.5, color="64748B")

    out_path = DESKTOP / "github-hot-projects-top100-2026-05-15-中文说明版.docx"
    doc.save(out_path)
    return out_path


def main() -> None:
    projects = fetch_projects()
    out = build_doc(projects)
    print(out)


if __name__ == "__main__":
    main()
