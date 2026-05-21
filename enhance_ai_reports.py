from __future__ import annotations

import shutil
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT_ARG = Path(sys.argv[1]) if len(sys.argv) > 1 else Path.cwd()
ACCESS_DATE = "2026年5月20日"


ACCENT = RGBColor(31, 78, 121)
LIGHT_BLUE = "D9EAF7"
LIGHT_GREEN = "E2F0D9"
LIGHT_ORANGE = "FCE4D6"
LIGHT_GRAY = "F2F2F2"
TEXT = RGBColor(40, 40, 40)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = "B7B7B7", size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 10, color: RGBColor | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color or TEXT
    set_run_font(run)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_run_font(run, name: str = "微软雅黑") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def set_paragraph_font(paragraph, size: int = 10, bold: bool = False, color: RGBColor | None = None) -> None:
    for run in paragraph.runs:
        set_run_font(run)
        run.font.size = Pt(size)
        run.bold = bold
        if color:
            run.font.color.rgb = color


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = ACCENT if level <= 2 else TEXT
    run.font.size = Pt(15 if level == 1 else 12)
    set_run_font(run)
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_body(doc: Document, text: str, after: int = 5):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    set_run_font(run)
    run.font.size = Pt(10.5)
    run.font.color.rgb = TEXT
    return p


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(item)
        set_run_font(run)
        run.font.size = Pt(10.5)
        run.font.color.rgb = TEXT


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Cm(0.55)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(item)
        set_run_font(run)
        run.font.size = Pt(10.5)
        run.font.color.rgb = TEXT


def add_note_box(doc: Document, title: str, items: list[str], fill: str = LIGHT_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, "9ECAE1")
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = ACCENT
    set_run_font(r)
    for item in items:
        p2 = cell.add_paragraph()
        p2.paragraph_format.left_indent = Cm(0.2)
        p2.paragraph_format.space_after = Pt(2)
        r2 = p2.add_run(f"· {item}")
        r2.font.size = Pt(9.5)
        r2.font.color.rgb = TEXT
        set_run_font(r2)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, size=9, color=RGBColor(255, 255, 255))
        set_cell_shading(hdr[i], "1F4E79")
        set_cell_border(hdr[i], "1F4E79")
        if widths:
            hdr[i].width = Cm(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val, size=8.8)
            set_cell_border(cells[i])
            if i == 0:
                set_cell_shading(cells[i], LIGHT_GRAY)
            if widths:
                cells[i].width = Cm(widths[i])
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)


def add_code_block(doc: Document, code: str) -> None:
    for line in code.strip("\n").splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.35)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        r.font.name = "Consolas"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(30, 30, 30)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def add_refs(doc: Document, refs: list[tuple[str, str]]) -> None:
    add_heading(doc, "参考资料", 1)
    for i, (name, url) in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{i}. {name}：{url}，访问日期：{ACCESS_DATE}")
        set_run_font(r)
        r.font.size = Pt(9)
        r.font.color.rgb = TEXT


def setup_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.65)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)
    styles = doc.styles
    styles["Normal"].font.name = "微软雅黑"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    styles["Normal"].font.size = Pt(10.5)
    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = styles[style_name]
        style.font.name = "微软雅黑"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def add_cover(doc: Document, title: str, subtitle: str, direction: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = ACCENT
    set_run_font(r)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(12)
    r2 = p2.add_run(subtitle)
    r2.font.size = Pt(12)
    r2.font.color.rgb = RGBColor(89, 89, 89)
    set_run_font(r2)

    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows = [
        ("作业类型", "个人AI应用作品 / 技术报告"),
        ("报告日期", ACCESS_DATE),
        ("适用方向", direction),
        ("修订重点", "事实核对、技术参数、实施流程、风险控制、展示证据链"),
    ]
    for row_idx, (left, right) in enumerate(rows):
        cells = table.rows[row_idx].cells
        set_cell_text(cells[0], left, bold=True, size=9.5, color=ACCENT)
        set_cell_text(cells[1], right, size=9.5)
        set_cell_shading(cells[0], LIGHT_GRAY)
        set_cell_border(cells[0])
        set_cell_border(cells[1])
    doc.add_paragraph().paragraph_format.space_after = Pt(8)


REPORTS: list[dict] = [
    {
        "file": "01_LoRA训练技术报告.docx",
        "title": "LoRA训练技术报告",
        "subtitle": "面向角色一致性与风格迁移的轻量化模型训练方案",
        "direction": "AI视觉设计、模型微调、风格资产生产",
        "summary": [
            "原报告对LoRA定位基本正确：LoRA属于参数高效微调，不重训完整基础模型，而是在目标层加入低秩增量权重。",
            "已补充训练数据治理、caption策略、rank/alpha/学习率、验证样张和FLUX/SDXL差异，避免只停留在“能训练出风格”的描述。",
            "已强化风险控制：版权授权、过拟合、触发词污染、人物同质化、底模许可和课程展示证据链。",
        ],
        "sections": [
            ("一、技术定位与设计思路", [
                "LoRA（Low-Rank Adaptation）是一种参数高效微调方法：冻结基础模型主体权重，只训练少量低秩适配矩阵。用于扩散模型时，它把“角色身份、服装结构、纹样语言、材质质感或视觉风格”压缩成一个可加载的适配器文件，通常以 .safetensors 交付。",
                "本项目不把LoRA理解为单次出图技巧，而把它作为AI视觉资产库的生产方式。对课程作品而言，训练目标要先被拆清楚：人物LoRA解决角色一致性，风格LoRA解决系列视觉统一，纹样/物件LoRA解决特定图案和材质复用。三类目标不建议混在同一个训练集中，否则模型容易把脸、服装、背景、颜色和构图一起记死。",
                "技术路线应以“可复现”为核心：固定底模版本、训练脚本版本、数据集目录、caption规则、随机种子和样张提示词。最终交付不只给出成品图，还要展示数据整理、训练日志、阶段样张、推理加载和应用结果。",
            ]),
            ("二、环境与工具配置", [
                "入门可选择Kohya GUI、AI Toolkit或云端训练平台；需要可控性时使用kohya-ss/sd-scripts、Diffusers训练脚本或Ostris AI Toolkit。SDXL人物/风格LoRA仍然适合教学演示，因为资源需求较可控、生态成熟、验证方式清楚。",
                "FLUX.1、FLUX.2、Qwen-Image等新模型的LoRA训练已经进入常用工具链，但它们往往需要额外的文本编码器、AE文件、更多显存或CPU/GPU交换策略。课程作业如果电脑显存有限，建议先用SDXL完成一版稳定成果，再把新模型训练作为进阶方案说明。",
            ]),
            ("三、具体实施过程", [
                "先定义训练目标和验收标准。例如人物LoRA要能在正面、侧面、半身、全身和不同光照下保持身份稳定；风格LoRA要能迁移色彩、线条和构图语言，但不能把训练集中的固定背景硬记下来。",
                "整理数据集时要去除水印、错字、低清图、重复图和版权不明素材。人物数据建议覆盖角度、表情、距离和服装变化；风格数据要保持风格一致但题材多样；纹样/物件数据要加入局部特写和应用场景图。",
                "caption应采用“唯一触发词 + 主要对象 + 可变属性 + 风格/材质词”的结构。人物示例：ylq_character, young woman, red opera costume, frontal portrait。风格示例：yangliuqing_style, flat color, folk print, symmetrical composition。需要避免把不想学习的背景、摄影水印和错误文字写入caption。",
                "训练时先做小规模试跑：低步数输出样张，确认触发词有效、loss正常下降、样张没有严重污染后再延长训练。每隔固定epoch保存checkpoint和样张，保留失败版本用于说明参数如何影响结果。",
                "部署时把LoRA文件放入ComfyUI/models/loras，通过LoRA Loader加载，并在同一提示词、同一种子和同一采样器下对比“未加载LoRA/加载LoRA”的效果。",
            ]),
            ("四、关键参数与验证标准", []),
            ("五、作品应用效果", [
                "训练完成后，LoRA可以支撑系列化视觉生产：同一IP角色的设定板、表情包、海报主视觉、视频分镜、电商包装和文创延展都可以围绕同一个触发词生成。其应用价值不是“出图更炫”，而是让作品组之间拥有可复用的身份和风格控制能力。",
                "课程展示建议准备五段证据链：训练数据截图、caption样例、训练参数截图、阶段样张对比、ComfyUI加载LoRA后的应用结果。若能加入失败样张和修正原因，技术表达会更可信。",
            ]),
            ("六、创新点与风险控制", [
                "创新点在于把LoRA从个人化模型训练转化为设计流程资产。它可以把传统文化纹样、非遗风格、校园活动角色和商品视觉语言沉淀为可复用适配器，并与ControlNet、IP-Adapter、高清放大和电商打光组合成完整工作流。",
                "风险控制重点包括：使用授权数据；记录底模和工具版本；不要用单一构图数据训练风格；设置保留验证集；用固定提示词做阶段对比；避免训练真实人物身份或敏感肖像；最终图像必须人工复核文字、纹样和人物细节。",
            ]),
        ],
        "tables": [
            ("训练类型对比", ["训练类型", "数据建议", "关键技术点", "主要风险"], [
                ["人物LoRA", "20-80张，同一人物多角度、多表情、多景别", "唯一触发词、较低步数频繁验证、必要时加入正则图", "过拟合、变脸、服饰和背景被一起记忆"],
                ["风格LoRA", "50-200张，同一视觉语言但题材丰富", "caption保留风格词，弱化固定对象词，检查泛化题材", "只学颜色不学结构，或把固定构图硬记下来"],
                ["纹样/物件LoRA", "30-120张，局部特写+应用图", "保持纹样清晰，分辨主体/材质/背景", "纹样变形、伪文字、重复纹理"],
            ]),
            ("核心参数建议", ["参数", "技术含义", "建议范围", "判断方法"], [
                ["network_dim / rank", "LoRA容量，rank越高可记忆细节越多", "人物16-32；复杂风格32-64；FLUX可按层微调", "rank过低不像，过高易过拟合"],
                ["network_alpha", "缩放LoRA有效强度，影响训练稳定性", "常设为rank或rank的一半", "推理强度在0.6-1.0仍应稳定"],
                ["learning_rate", "控制权重更新速度", "SDXL常从1e-4试起；文本编码器更保守", "loss剧烈震荡或样张崩坏则降低"],
                ["resolution/buckets", "训练分辨率与多比例分桶", "SDXL常用1024；显存不足可768", "不同画幅样张不应明显拉伸"],
                ["sample_prompts", "训练中固定验证提示词", "每个目标至少3-5条", "同一seed下阶段对比身份/风格变化"],
            ]),
            ("FLUX训练补充", ["项目", "说明", "课程写法"], [
                ["模型结构", "FLUX使用Transformer结构，并依赖CLIP-L、T5-XXL和AE等组件", "不要套用SDXL的clip_skip等旧参数"],
                ["显存策略", "可使用cache_latents、cache_text_encoder_outputs、blocks_to_swap等方式降显存", "说明速度与显存的取舍"],
                ["采样/预测", "常见配置会涉及timestep_sampling、guidance_scale、model_prediction_type", "把新模型训练作为进阶路线，不强行承诺低显存稳定"],
            ]),
        ],
        "video": [
            "第1分钟说明LoRA解决“系列作品不统一”的问题。",
            "第2-3分钟展示数据集、caption和触发词设计。",
            "第4-6分钟展示训练参数、loss趋势和阶段样张。",
            "第7-8分钟在ComfyUI加载LoRA，对比未加载/已加载效果。",
            "第9-10分钟展示最终应用和风险控制表。",
        ],
        "refs": [
            ("Hu et al., LoRA: Low-Rank Adaptation of Large Language Models", "https://arxiv.org/abs/2106.09685"),
            ("Hugging Face Diffusers, LoRA training guide", "https://huggingface.co/docs/diffusers/training/lora"),
            ("kohya-ss, sd-scripts", "https://github.com/kohya-ss/sd-scripts"),
            ("ostris, AI Toolkit", "https://github.com/ostris/ai-toolkit"),
            ("Comfy-Org, ComfyUI", "https://github.com/Comfy-Org/ComfyUI"),
        ],
    },
    {
        "file": "02_ComfyUI搭建技术报告.docx",
        "title": "ComfyUI搭建技术报告",
        "subtitle": "面向AI视觉生产的节点式工作流环境建设",
        "direction": "AI视觉设计、节点工作流、智能体应用搭建",
        "summary": [
            "原报告对ComfyUI节点式工作流定位正确，但安装方式和插件管理表述偏旧。",
            "已补充Desktop、Windows Portable、源码/manual、comfy-cli等路线，并加入模型目录、Manager启用、版本锁定和工作流可复现要求。",
            "已强化节点链路、排错方法、JSON/PNG工作流复用、API化和安全风险。",
        ],
        "sections": [
            ("一、技术定位与设计思路", [
                "ComfyUI是面向生成式AI内容生产的节点式界面和推理引擎。它把模型加载、文本编码、采样、VAE解码、控制网络、参考图、局部重绘、高清放大和保存输出拆成可视化节点，使生成过程可观察、可复用、可调试。",
                "与只输入提示词的绘图工具相比，ComfyUI更适合做AI应用技术报告，因为它能展示完整的计算链路：输入是什么、模型如何加载、哪些节点控制结构、哪些节点处理细节、结果如何保存。对课程作业而言，这种可解释性比单张成品图更重要。",
                "本项目把ComfyUI设定为AI视觉设计的底层工作台，用于LoRA调用、ControlNet构图、IP-Adapter参考控制、局部重绘、高清放大、电商打光、多比例输出和后续批量处理。",
            ]),
            ("二、硬件与软件环境", [
                "当前ComfyUI可通过Desktop应用、Windows便携包、源码/manual安装或comfy-cli安装。Windows课堂展示优先选择Desktop或Portable，依赖少、启动快；需要自定义节点开发、API集成或版本控制时选择源码安装。",
                "模型目录必须清晰：checkpoints存基础模型，loras存LoRA，vae存VAE/AE，controlnet存ControlNet模型，clip/text_encoders存新模型所需文本编码器。大型模型的目录要求不同，不能把所有文件都塞进checkpoints。",
                "显卡方面，8GB显存可完成基础SDXL推理和小规模放大；复杂ControlNet、视频、FLUX或多节点并行建议使用更高显存或云端GPU。低显存时优先使用分块、低批量、FP16/BF16、显存释放和模型卸载策略。",
            ]),
            ("三、具体搭建过程", [
                "第一步确认安装路线。便携包适合“下载、解压、运行”的课程演示；源码路线建议建立独立虚拟环境，并记录Python、PyTorch、CUDA、ComfyUI commit和自定义节点版本。",
                "第二步整理模型目录。每次新增模型后重启或刷新模型列表，并在报告中写明模型来源、文件类型和放置目录。对课程项目，应避免使用来源不明的ckpt/pt文件。",
                "第三步启用或安装ComfyUI-Manager。当前版本中Manager多通过内置或官方方式启用；旧环境也可通过custom_nodes安装。安装自定义节点时要检查README、依赖和安全性，不要批量安装未知节点。",
                "第四步先跑最小工作流：Checkpoint Loader、CLIP Text Encode、KSampler、VAE Decode、Save Image。只有基础链路稳定后，再接入LoRA、ControlNet、IP-Adapter、Inpaint和Upscale。",
                "第五步保存工作流JSON，并把seed、模型名、LoRA强度、采样器、步数、CFG、分辨率写入报告。ComfyUI生成的PNG/WebP可携带工作流元数据，适合课堂复现演示。",
            ]),
            ("四、工作流设计方法", [
                "工作流应遵循“先小后大、先稳定后复杂”的原则。先建立四个基础模板：基础文生图、LoRA人物、ControlNet构图、高清放大。每个模板都要能独立运行，再把它们组合成完整生产链。",
                "复杂工作流要分区命名：输入区、模型区、控制区、采样区、后处理区、输出区。节点过多时使用Group和颜色标签，避免视频展示时观众看不懂。",
                "可复现性要求包括：固定随机种子或说明随机策略；记录模型文件和哈希；保存workflow JSON；标注自定义节点版本；在新增节点后做一次基础回归测试。",
            ]),
            ("五、应用效果", [
                "搭建完成后，ComfyUI可以成为完整的AI视觉生产系统。用户输入主题、参考图和参数后，可以生成主视觉底图，再通过LoRA保持角色一致，通过ControlNet控制姿势和构图，通过局部重绘修复细节，通过高清放大输出海报或展板级图像。",
                "报告中的应用效果应突出四点：流程可复用、结果可追溯、参数可调整、节点可扩展。团队协作时，JSON工作流、模型目录说明和节点依赖清单比单张图片更有价值。",
            ]),
            ("六、创新点与风险控制", [
                "创新点在于用节点图把AI视觉生产流程显性化。它不仅能生成图像，还能把“参考图控制、结构控制、风格控制、细节修复、批量输出”变成可演示的应用系统。",
                "风险包括自定义节点版本冲突、依赖污染、模型路径混乱、显存不足、未知节点安全问题和工作流不可复现。解决方法是使用独立环境、一次只新增一个节点、保留requirements/版本记录、备份custom_nodes和workflow，并给每个模板保存成功样张。",
            ]),
        ],
        "tables": [
            ("安装路线对比", ["路线", "适合场景", "优点", "注意事项"], [
                ["Desktop", "课堂展示、快速上手", "安装简单，适合非开发用户", "自定义深度低于源码环境"],
                ["Windows Portable", "Windows本地演示", "可解压即用，依赖隔离较好", "需按显卡选择版本并更新驱动"],
                ["源码/manual", "开发、自定义节点、API集成", "可控性强，便于版本管理", "Python/PyTorch/CUDA依赖需严格匹配"],
                ["comfy-cli", "自动化安装和脚本化管理", "命令统一，适合批量部署", "仍需管理模型文件和自定义节点"],
            ]),
            ("基础工作流模板", ["模板", "核心节点", "关键参数", "适用输出"], [
                ["基础文生图", "Checkpoint Loader、CLIP Text Encode、KSampler、VAE Decode", "seed、steps、CFG、sampler、scheduler", "海报底图、概念图"],
                ["LoRA人物", "LoRA Loader、Reference/IP-Adapter、KSampler", "LoRA strength、trigger word、seed", "角色设定、表情动作延展"],
                ["构图控制", "ControlNet Loader、Preprocessor、Apply ControlNet", "control strength、start/end percent", "姿势、线稿、深度、边缘控制"],
                ["高清放大", "Upscale Model、Image/Latent Upscale、Tiled VAE、Inpaint", "scale、denoise、tile size、overlap", "展板、门型架、商品主图"],
            ]),
            ("常见问题排查", ["问题", "可能原因", "处理方法"], [
                ["Torch not compiled with CUDA enabled", "安装了CPU版PyTorch或CUDA版本不匹配", "重装匹配显卡驱动的PyTorch版本"],
                ["缺失节点/Missing nodes", "工作流依赖自定义节点未安装", "通过Manager安装，重启并检查启动日志"],
                ["模型不显示", "文件放错目录或格式不匹配", "按模型类型放入对应models子目录"],
                ["显存溢出", "分辨率、batch、ControlNet或视频节点过重", "降低分辨率，启用分块/VAE tile，减少并行节点"],
            ]),
        ],
        "video": [
            "展示安装路线和模型目录结构。",
            "运行最小文生图工作流，证明环境可用。",
            "加入LoRA、ControlNet或IP-Adapter，展示节点扩展逻辑。",
            "展示JSON工作流保存、导入和复现。",
            "说明版本记录和排错表，体现技术可靠性。",
        ],
        "refs": [
            ("Comfy-Org, ComfyUI GitHub", "https://github.com/Comfy-Org/ComfyUI"),
            ("ComfyUI Official Documentation", "https://docs.comfy.org/"),
            ("ComfyUI Docs, Install Custom Nodes", "https://docs.comfy.org/installation/install_custom_node"),
            ("ComfyUI Docs, Custom Nodes Overview", "https://docs.comfy.org/custom-nodes/overview"),
        ],
    },
    {
        "file": "03_超清放大技术报告.docx",
        "title": "超清放大技术报告",
        "subtitle": "面向海报、展板与AI图像交付的高清增强流程",
        "direction": "AI图像后期、超分辨率、打印输出",
        "summary": [
            "原报告的三层分类基本正确，但已纠正一个关键点：Tiled VAE主要降低高分辨率解码显存压力，不等同于超分模型本身。",
            "已补充像素超分、潜空间重绘、瓦片重绘、打印尺寸换算、denoise/tile/overlap等可执行参数。",
            "已强化质量验收，特别是人物脸、手、纹样文字、商品Logo和打印视距。",
        ],
        "sections": [
            ("一、技术定位与设计思路", [
                "超清放大不是简单把图片尺寸拉大，而是在尽量保持构图和主体结构的前提下提升边缘、纹样、材质和局部清晰度。AI图像常见问题包括分辨率不足、手部变形、脸部糊、伪文字、纹样重复、边缘噪点和过度锐化。",
                "本报告将高清增强拆成三类技术：像素级超分、潜空间重绘放大、瓦片/分块式重绘。像素超分结构最稳定但新增细节有限；潜空间重绘能补细节但会改变图像；瓦片式适合大尺寸海报，但需要控制接缝、局部风格漂移和重复纹理。",
                "课程项目应建立“先修结构、再放大、再局部检查、最后印刷换算”的流程。若原图人物比例、Logo、文字或纹样已经错误，放大只会让错误更明显。",
            ]),
            ("二、环境与工具配置", [
                "推荐在ComfyUI中建立高清放大模板，包含图像加载、超分模型、尺寸调整、图生图低强度重绘、瓦片处理、Tiled VAE解码和保存输出。后期可在Photoshop中完成锐化、降噪、色阶、局部蒙版和打印尺寸换算。",
                "需要明确区分模型功能：RealESRGAN、SwinIR、4x-UltraSharp等属于像素超分；KSampler/图生图负责潜空间细节重绘；Tiled VAE负责把高分辨率latent分块解码，主要解决显存压力，不直接决定细节质量。",
            ]),
            ("三、具体实施过程", [
                "第一步生成稳定底图。先把构图、主体、光影和色彩定好，不要在构图仍摇摆时直接追求8K。",
                "第二步做结构修复。用局部重绘修脸、手、服装破损、伪文字和异常纹样。商品图要单独保护Logo、包装文字和瓶身轮廓。",
                "第三步做第一次像素超分。建议先用2x或4x超分模型，观察边缘、噪点、材质和脸部是否被过度锐化。",
                "第四步做低强度图生图或瓦片重绘。角色一致性项目denoise可从0.15-0.25起步，复杂材质可试0.25-0.35；高于0.4时必须重点检查是否变脸、变Logo或改构图。",
                "第五步使用瓦片策略。常见tile size可从512或768开始，overlap保持64-128像素以减少接缝。显存不足时缩小tile，但要增加重叠并检查拼接边界。",
                "第六步导出交付文件。网络展示可用PNG；印刷前应换算实际尺寸、DPI和色彩模式，必要时输出TIFF或交给Photoshop做最终校色。",
            ]),
            ("四、关键参数与质量标准", [
                "放大质量不以“8K”字样为唯一标准，而以结构不变、细节提升、风格统一和可交付为标准。真正可用的高清图在100%查看时不应出现明显接缝、眼睛糊、伪文字、局部重复纹理或材质塑料化。",
                "打印尺寸换算要写入报告。例如A3横向在300dpi下约4961×3508像素；展板和门型架因观看距离较远，可按150-200dpi评估。课程展示至少要给出原图、第一次超分、瓦片重绘、最终修图四张对比。",
            ]),
            ("五、作品应用效果", [
                "该技术适合AI海报、门型架人物介绍、校园活动背景板、电商主图、文创包装和短视频封面。它的价值在于把AI初稿提升到可展示、可印刷、可交付的质量，而不是盲目增加像素。",
                "在角色项目中，高清放大要维护人物身份和服装语言；在非遗纹样项目中，要维护图案节奏和边缘清晰；在电商项目中，要维护Logo、包装比例、反光和接触阴影。",
            ]),
            ("六、创新点与风险控制", [
                "创新点在于把高清放大做成可控工作流，而非后期软件中的单次滤镜。通过分层放大、低强度重绘、瓦片处理、Tiled VAE和人工检查，可以在保持设计语言的同时提升细节。",
                "主要风险是AI在放大时重新创作：人物变脸、手指变多、纹样变乱、Logo变形、细节过度锐化。控制方法包括低denoise、保留原图参考、分区蒙版保护、tile overlap、固定seed和输出后人工复核。",
            ]),
        ],
        "tables": [
            ("放大方式对比", ["方式", "优点", "风险", "适用场景"], [
                ["像素级超分", "速度快，结构稳定", "新增细节有限，可能过锐化", "小图变清晰、初步放大"],
                ["图生图重绘放大", "能补材质和局部细节", "容易改变人物、Logo和构图", "角色海报、复杂材质"],
                ["瓦片式重绘", "适合大尺寸输出，显存压力低", "接缝、局部不一致、重复纹理", "展板、门型架、8K主视觉"],
                ["Tiled VAE", "降低高分辨率解码显存压力", "不是超分模型，不能单独提升清晰度", "高分辨率latent解码"],
            ]),
            ("参数建议", ["参数", "建议", "说明"], [
                ["初始图长边", "1024-1536起步", "先保证构图准确，不急于超大尺寸"],
                ["超分倍率", "2x或4x", "先观察结构是否稳定，再决定是否二次处理"],
                ["denoise", "0.15-0.35", "越高越会重画，角色和商品图应偏低"],
                ["tile size", "512-768", "显存低用小tile，注意overlap"],
                ["overlap", "64-128", "减少接缝和边缘不一致"],
            ]),
            ("质量检查清单", ["对象", "合格标准", "常见问题"], [
                ["人物脸部", "五官清楚，身份不变，左右脸稳定", "眼睛变形、磨皮过度、脸部重画"],
                ["手部结构", "手指数量正确，动作合理", "多指、断指、手掌粘连"],
                ["纹样文字", "纹样边缘清楚，文字区不生成伪字", "重复纹理、伪文字、符号乱跳"],
                ["商品Logo", "形状、比例、字体不变", "品牌字母错误、包装被改写"],
                ["打印输出", "按尺寸和DPI换算后仍清晰", "像素不足、锐化光晕、色偏"],
            ]),
        ],
        "video": [
            "展示低分辨率原图及局部问题。",
            "展示ComfyUI超清放大节点链路，并说明Tiled VAE的作用。",
            "对比不同denoise值对脸、纹样和材质的影响。",
            "展示tile size/overlap对接缝的影响。",
            "展示最终图在海报、展板或商品图中的应用。",
        ],
        "refs": [
            ("Comfy-Org, ComfyUI GitHub", "https://github.com/Comfy-Org/ComfyUI"),
            ("ComfyUI Examples, Upscale Models", "https://comfyanonymous.github.io/ComfyUI_examples/upscale_models/"),
            ("Real-ESRGAN GitHub", "https://github.com/xinntao/Real-ESRGAN"),
            ("SwinIR GitHub", "https://github.com/JingyunLiang/SwinIR"),
        ],
    },
    {
        "file": "04_电商打光工作流技术报告.docx",
        "title": "电商打光工作流技术报告",
        "subtitle": "面向商品主图、详情页与广告图的AI光影重构方案",
        "direction": "电商视觉、AI摄影后期、商品图生成",
        "summary": [
            "原报告方向正确，但已增加电商图最关键的主体保真、蒙版保护、Logo保护、阴影方向和批量输出标准。",
            "已补充技术链路：分割/抠图、mask feather、ControlNet轮廓控制、参考图控制、低强度inpaint、接触阴影和材质反射。",
            "已强化平台交付和法律风险：真实商品不得改变商标、包装文字、容量、外形比例和功能细节。",
        ],
        "sections": [
            ("一、技术定位与设计思路", [
                "电商打光工作流的核心目标不是“把商品图变漂亮”，而是在不改变商品真实信息的前提下重建商业摄影质感。电商图需要让消费者看清形状、材质、尺度、卖点和品牌信息，因此主体保真比艺术化更重要。",
                "本报告将工作流拆成六个环节：主体保真、背景重构、光源设计、阴影统一、材质增强、输出适配。AI可以帮助降低拍摄和后期成本，但不能随意改变Logo、包装文字、瓶身比例、容量标识和功能部件。",
                "课程展示应强调“可控商业视觉流程”：从普通商品图输入，到蒙版保护，再到背景、灯光、阴影和多比例输出，形成可重复使用的电商摄影模板。",
            ]),
            ("二、工作流环境", [
                "推荐使用ComfyUI搭建电商打光流程。输入可以是白底商品图、手机拍摄图或三维渲染图。关键节点包括背景移除/分割、蒙版处理、Inpaint、ControlNet Canny/Depth、IP-Adapter或参考图控制、超清放大和多尺寸保存。",
                "真实商品图应使用低重绘强度，并尽量只重绘背景、台面、阴影和反射区域。若必须处理商品表面，应单独制作保护蒙版，对Logo、包装字、二维码、容量和功能标识进行排除。",
            ]),
            ("三、具体实施过程", [
                "第一步准备商品原图。选择透视正常、主体清楚、边缘完整、没有严重反光和运动模糊的商品图。若输入质量过差，先做基础修图和边缘清理。",
                "第二步主体分离。通过抠图、分割或手工蒙版得到商品mask，并做适度dilation/feather，避免背景重绘时出现白边或吃掉商品边缘。",
                "第三步确定光源方案。白底平台主图通常使用高调柔光；品牌氛围图可使用侧逆光、轮廓光和暗调背景；节日广告图可增加暖光、背景道具和局部高光，但主光方向必须统一。",
                "第四步重建背景与阴影。使用局部重绘生成干净台面、背景纸、接触阴影和反射。阴影必须贴合商品底部，方向与主光一致，不能出现悬浮感。",
                "第五步强化材质。玻璃需要透亮和边缘高光；金属需要可控反射；纸盒需要纤维和哑光；塑料需要避免过亮油腻。提示词应描述摄影和材质，不要描述会改变商品本身的内容。",
                "第六步导出多尺寸结果。按平台需求输出1:1主图、4:5详情页、16:9横幅、9:16短视频封面和局部特写，并用统一色彩和锐化策略保持系列一致。",
            ]),
            ("四、提示词与参数规范", [
                "电商提示词应使用商业摄影语言，例如softbox lighting、realistic contact shadow、clean studio background、sharp product edges、premium material texture、controlled reflection。少用fantasy、magic、ultra surreal等容易改变商品属性的词。",
                "负向词应覆盖商品保真问题：logo distortion、wrong text、changed package shape、extra label、over reflection、dirty background、floating object、watermark、low resolution、plastic look。中文描述中也要明确“主体结构不变、包装文字不改、Logo不变形”。",
            ]),
            ("五、应用效果", [
                "最终可生成三类成果：标准白底主图、品牌氛围详情图、活动广告KV。一个商品能输出三套光影风格，体现工作流复用能力和商业落地价值。",
                "课程展示建议用同一商品进行三组对比：原始手机图、高调白底图、暗调质感图、节日广告图。同时展示局部放大检查：Logo、文字、边缘、阴影、反射和材质。",
            ]),
            ("六、创新点与风险控制", [
                "创新点在于把AI图像生成用于商业摄影流程，而不是只做艺术滤镜。通过主体保护、蒙版控制、参考图约束、低强度重绘和多比例输出，普通商品图可以转化为电商级视觉资产。",
                "风险包括商品形状被改、品牌字错误、包装容量变化、阴影不真实、材质失真和过度美化。解决方案是保留主体原像素、保护文字区、分层重绘、低denoise、输出前逐项放大检查，并保留原图和修改记录。",
            ]),
        ],
        "tables": [
            ("工作流模块", ["目标", "技术手段", "验收标准"], [
                ["主体保真", "抠图、蒙版、参考图控制、低denoise", "商品形状、Logo、包装结构不变"],
                ["背景重构", "Inpaint、纯色/场景背景、台面生成", "背景干净，商品边缘无白边"],
                ["光影统一", "主光、辅光、轮廓光、接触阴影", "阴影方向和强度与光源一致"],
                ["材质增强", "提示词、参考图、局部重绘", "玻璃、金属、纸、塑料质感可区分"],
                ["输出适配", "多比例裁切、高清放大、批量保存", "适配主图、详情页、广告KV"],
            ]),
            ("节点/工具建议", ["节点/工具", "作用", "注意事项"], [
                ["Remove Background / Segment", "分离商品主体", "边缘要检查，mask需适度羽化"],
                ["Mask / Inpaint", "重建背景、阴影和台面", "蒙版不得覆盖Logo和关键文字"],
                ["ControlNet Canny/Depth", "保持轮廓和空间关系", "强度过高会限制光影变化"],
                ["IP-Adapter / Reference", "保持商品视觉特征", "适合包装、鞋服、文创产品"],
                ["Upscale + PS", "交付前清晰化和人工收尾", "检查文字、噪点和锐化光晕"],
            ]),
            ("输出规格", ["输出类型", "常用比例", "用途", "检查重点"], [
                ["白底主图", "1:1", "平台列表页、SKU展示", "边缘、阴影、白底洁净"],
                ["详情页氛围图", "4:5 / 3:4", "材质、功能和品牌调性", "道具不抢主体"],
                ["广告KV", "16:9 / 9:16", "直播封面、短视频、活动海报", "视觉冲击与商品识别"],
                ["局部特写", "自由裁切", "工艺、纹样、卖点说明", "细节真实不过度锐化"],
            ]),
        ],
        "video": [
            "展示普通商品图作为输入，并说明需要保护的信息。",
            "展示抠图、mask扩张/羽化和保护区域。",
            "生成白底柔光、暗调侧光、节日暖光三种结果。",
            "放大检查Logo、文字、边缘、阴影和材质。",
            "展示多比例输出，说明如何用于真实电商页面。",
        ],
        "refs": [
            ("Comfy-Org, ComfyUI GitHub", "https://github.com/Comfy-Org/ComfyUI"),
            ("ComfyUI Official Documentation", "https://docs.comfy.org/"),
            ("ComfyUI Examples, Inpainting", "https://comfyanonymous.github.io/ComfyUI_examples/inpaint/"),
            ("ComfyUI Examples, ControlNet", "https://comfyanonymous.github.io/ComfyUI_examples/controlnet/"),
        ],
    },
    {
        "file": "05_Codex接入Blender5.0技术报告.docx",
        "title": "Codex接入Blender 5.0技术报告",
        "subtitle": "面向三维建模、材质与渲染自动化的AI编码工作流",
        "direction": "AI编码、Blender Python、三维自动化",
        "summary": [
            "原报告把“接入”解释为本地脚本工程协作，这一点正确；已进一步区分Codex App、IDE、CLI和Cloud，不再把Codex写成Blender内置插件。",
            "已补充Blender后台执行、bpy数据API、bpy.ops上下文限制、版本检测、5.0 API变化和自动化测试。",
            "已增强脚本工程结构、命令行示例、功能模块和风险控制。",
        ],
        "sections": [
            ("一、技术定位与设计思路", [
                "Codex接入Blender 5.0的本质，是让AI编码智能体参与Blender Python脚本、插件、批处理和资产生成流程。Codex负责理解需求、生成代码、修改脚本、解释报错和组织工程；Blender负责执行三维场景、几何体、材质、灯光、摄像机和渲染。",
                "这里的“接入”不是把Codex变成Blender内部按钮，而是在本地项目目录建立脚本工程。Codex通过读取README、修改Python脚本、运行命令和根据报错迭代，与Blender的Python API形成“需求—脚本—执行—渲染—复盘”的闭环。",
                "Blender 5.0对应的Python API存在版本差异，脚本不应只凭旧教程编写。需要在脚本中检查bpy.app.version，优先使用稳定的数据API，谨慎使用依赖界面上下文的bpy.ops。",
            ]),
            ("二、环境配置", [
                "Codex侧可选择App、IDE扩展、CLI或Cloud任务。课堂本地演示建议使用Codex App或CLI：在项目文件夹中运行，让它读取README、生成脚本、执行测试并修复报错。Cloud任务更适合连接GitHub仓库做长任务，不适合作为现场依赖。",
                "Blender侧安装官方5.0版本，并确认命令行可调用。Windows可把blender.exe加入PATH，也可以在命令中写完整路径。后台渲染使用--background或-b，脚本使用--python加载。",
            ]),
            ("三、具体实施过程", [
                "第一步建立项目目录，包含scripts、assets、textures、renders、tests和README。README要写清楚目标场景、单位、尺寸、材质、灯光、相机、渲染器和输出路径。",
                "第二步让Codex生成最小脚本create_scene.py：清空场景、设置单位、创建几何体、添加材质、布置灯光、设置相机、配置Cycles/EEVEE并输出PNG。",
                "第三步用Blender命令行执行脚本。若出现AttributeError、context is incorrect、路径错误或材质节点报错，把完整报错交给Codex修复。",
                "第四步把成功脚本拆成模块：scene_setup.py、materials.py、lighting.py、camera.py、render.py。这样后续只改参数即可批量生成不同商品摄影棚、展台草图或文创器物展示。",
                "第五步增加最小验收：脚本运行退出码为0，renders目录生成图片，场景中包含预期对象，摄像机和灯光存在，渲染分辨率正确。",
            ]),
            ("四、脚本功能设计", [
                "Blender自动化应优先使用bpy.data创建和管理对象、材质、灯光和相机；bpy.ops适合调用建模操作，但在后台模式下容易因为缺少界面上下文而失败。报告中要说明这一点，体现技术判断。",
                "材质系统可基于Principled BSDF设置base color、metallic、roughness、alpha、transmission等参数；灯光系统可使用Area Light模拟柔光箱，Point/Spot Light做补光或轮廓光；相机需设置焦距、位置、目标点和景深。",
                "Blender 5.0的Python API更新需要注意旧脚本兼容性，例如不要依赖运行时定义属性的旧式字典访问方式；导入导出、USD/Alembic等接口也要以5.0文档为准。",
            ]),
            ("五、应用效果", [
                "Codex接入Blender后，可以形成“文本需求生成三维场景”的半自动化流程。设计类学生可快速生成电商产品棚拍、文创包装展示、展台空间草图、国风器物三维示意和渲染测试。",
                "该流程的优势是可修改、可批量、可复现。同一脚本换参数即可生成多个尺寸、材质、灯光和相机方案，比手动重复建模更适合课程展示。",
            ]),
            ("六、创新点与风险控制", [
                "创新点在于把AI从二维图像生成拓展到三维软件自动化。Codex负责代码生成与排错，Blender负责真实三维空间和物理渲染，两者结合可以形成低成本数字摄影棚和三维设计智能体雏形。",
                "风险包括API版本变化、bpy.ops上下文错误、脚本生成模型过于简化、材质不真实、渲染耗时和路径兼容问题。解决方法是固定Blender版本、使用小脚本分步测试、保存Git版本、记录命令行日志，并对复杂造型保留人工建模环节。",
            ]),
        ],
        "tables": [
            ("项目目录建议", ["目录/文件", "作用", "技术要求"], [
                ["README.md", "写清需求和验收标准", "包含单位、尺寸、材质、镜头、输出路径"],
                ["scripts/create_scene.py", "生成基础场景", "可独立运行，避免硬编码绝对路径"],
                ["scripts/materials.py", "材质函数库", "按纸、金属、玻璃、塑料封装参数"],
                ["assets/models", "外部模型素材", "记录来源和许可"],
                ["renders", "渲染输出", "文件名包含方案、尺寸和时间"],
                ["tests", "脚本验收", "检查对象数量、相机、灯光和输出文件"],
            ]),
            ("功能模块", ["模块", "Blender Python操作", "应用价值"], [
                ["场景初始化", "删除默认物体、设置单位、世界背景", "保证每次运行一致"],
                ["几何生成", "创建立方体、圆柱、曲线、倒角", "快速生成展台、包装盒、道具"],
                ["材质系统", "Principled BSDF、粗糙度、金属度、透明度", "控制纸、金属、玻璃、塑料质感"],
                ["灯光系统", "Area/Point/Sun Light，三点布光", "模拟电商打光和海报光影"],
                ["相机渲染", "焦距、景深、分辨率、采样、输出路径", "形成可交付渲染图"],
            ]),
            ("技术注意事项", ["问题", "原因", "建议"], [
                ["bpy.ops在后台失败", "操作依赖界面上下文", "优先使用bpy.data，必要时构造context override"],
                ["旧脚本在5.0报错", "API变更或属性访问方式变化", "查5.0 Python API并用bpy.app.version分支"],
                ["材质看起来假", "只设置颜色，未控制粗糙度/金属度/透明度", "为不同材质建立参数模板"],
                ["渲染过慢", "采样过高、分辨率过大、材质复杂", "先低采样预览，再高质量输出"],
            ]),
        ],
        "code": """# 命令行执行示例
blender --background --python scripts/create_scene.py

# Codex提示词示例
请为Blender 5.0生成一个Python脚本：清空场景，设置毫米单位，
创建产品展示台、三点布光、相机和Cycles渲染设置，
将PNG输出到renders文件夹，并在脚本中打印bpy.app.version。""",
        "video": [
            "展示项目目录和README需求。",
            "展示Codex生成并修改create_scene.py。",
            "展示Blender命令行运行脚本和报错修复。",
            "展示最终渲染图，并说明材质/灯光参数如何批量替换。",
            "展示脚本版本记录和可复现命令。",
        ],
        "refs": [
            ("OpenAI Developers, Codex CLI", "https://developers.openai.com/codex/cli"),
            ("OpenAI Developers, Codex Quickstart", "https://developers.openai.com/codex/quickstart"),
            ("OpenAI, Codex GitHub", "https://github.com/openai/codex"),
            ("Blender, 5.0 Release Notes", "https://developer.blender.org/docs/release_notes/5.0/"),
            ("Blender, Python API 5.0", "https://docs.blender.org/api/5.0/"),
            ("Blender, Python API Current", "https://docs.blender.org/api/current/"),
        ],
    },
    {
        "file": "06_Codex接入CAD技术报告.docx",
        "title": "Codex接入CAD技术报告",
        "subtitle": "面向AutoCAD制图、参数化绘图与批处理的AI自动化方案",
        "direction": "AI编码、AutoCAD自动化、参数化制图",
        "summary": [
            "原报告总体正确：Codex适合作为脚本生成和排错助手，CAD负责精确制图。",
            "已补充AutoLISP、SCR、.NET/ObjectARX/ActiveX等开发接口，并明确课程项目优先做AutoLISP和脚本批处理。",
            "已加入SECURELOAD、TRUSTEDPATHS、APPLOAD/load、单位/图层/标注/系统变量恢复等关键技术点。",
        ],
        "sections": [
            ("一、技术定位与设计思路", [
                "Codex接入CAD的核心思路，是让AI编码智能体帮助生成AutoCAD脚本、AutoLISP程序、参数化绘图代码和批量处理命令。CAD软件负责精确几何和图纸规范，Codex负责把自然语言需求转化为可执行、可检查、可迭代的脚本。",
                "本报告以AutoCAD为主要对象。Autodesk官方生态提供AutoLISP/Visual LISP、ObjectARX、.NET API、ActiveX等自定义方式。课程作业不需要开发大型插件，最适合展示的是“需求说明—Codex生成LSP/SCR—AutoCAD加载运行—检查图纸—修复脚本”的闭环。",
                "与AI绘图不同，CAD强调尺寸、坐标、图层、线型、标注比例和输出规范。报告中必须体现精度意识，不能只说“看起来像图纸”。",
            ]),
            ("二、环境配置", [
                "AutoCAD侧需要合法授权版本或教育版，并确认可以加载AutoLISP。脚本文件建议放入受信任目录；AutoCAD安全模式下，SECURELOAD和TRUSTEDPATHS会影响LSP加载。",
                "Codex侧在项目目录中工作，读取README并生成lisp、scripts、dwg、output等文件。若使用Git，建议每次成功运行后提交脚本版本，便于回滚错误代码。",
            ]),
            ("三、具体实施过程", [
                "第一步明确CAD任务。需求必须包含单位、图纸尺寸、比例、坐标基准、图层、线型、文字高度、标注样式和输出格式。例如“生成A3横向图框，单位毫米，外框420×297，内边距10，右下角标题栏”。",
                "第二步让Codex生成AutoLISP。代码应包含命令函数、参数变量、图层创建、对象绘制、文字标注、错误处理和系统变量恢复。不要把关键尺寸写散，应集中定义常量。",
                "第三步在AutoCAD中加载脚本。可使用APPLOAD，也可在命令行输入(load \"路径\")。从AutoCAD 2021起LSP可使用ASCII或Unicode；旧版本要注意中文编码。",
                "第四步运行命令并核对结果。检查单位、坐标、线宽、图层、标注比例、标题栏位置和输出PDF。发现报错时，把AutoCAD命令行错误原文交给Codex修复。",
                "第五步固化模板。将图框生成、图层清理、批量导出、尺寸标注和标题栏填写脚本分开保存，形成CAD自动化工具包。",
            ]),
            ("四、技术模块设计", [
                "AutoLISP脚本应设置错误处理函数，在失败时恢复系统变量，避免污染用户CAD环境。涉及命令调用时要注意AutoCAD本地化命令名和系统变量；复杂对象可使用entmakex创建实体，减少交互式命令依赖。",
                "批处理可使用SCR脚本或AutoCAD命令行完成多图打开、清理、打印和导出。若要做云端自动化，可进一步研究Autodesk Platform Services的AutoCAD Automation API，但课程演示可以先聚焦本地AutoCAD。",
                "参数化绘图要把尺寸、比例、图层名、文字高度和标题栏字段写成变量。这样同一脚本可以生成A4/A3图框，或按输入参数生成不同尺寸构件。",
            ]),
            ("五、应用效果", [
                "Codex接入CAD后，可以显著减少重复制图和格式整理时间。最直观的课堂成果包括自动生成标准图框、自动建立图层、自动绘制简单平面构件、自动标注尺寸和批量导出PDF。",
                "最终展示应包含代码文件、AutoCAD加载画面、运行命令、生成图纸截图、输出PDF或DWG文件，并对关键尺寸做手工核对，证明结果不是视觉幻觉。",
            ]),
            ("六、创新点与风险控制", [
                "创新点在于把AI编码能力用于精确制图场景。Codex把自然语言需求转化为脚本，AutoCAD把脚本转化为可量测图纸，两者结合后形成参数化制图智能体雏形。",
                "风险包括AutoLISP语法错误、括号不匹配、单位错误、图层命名不规范、加载路径不受信任、系统变量未恢复和软件授权问题。解决方法是合法环境、受信任路径、逐步测试、错误处理、关键尺寸参数化和人工审核记录。",
            ]),
        ],
        "tables": [
            ("项目目录建议", ["目录/文件", "作用", "技术要求"], [
                ["README.md", "记录任务需求", "包含单位、比例、尺寸、图层、输出格式"],
                ["lisp/draw_frame.lsp", "图框和标题栏生成", "命令名清晰，尺寸参数集中"],
                ["lisp/layer_cleaner.lsp", "图层清理和规范化", "避免删除用户重要图层"],
                ["scripts/batch_plot.scr", "批量打印或导出", "先在测试DWG中验证"],
                ["dwg", "输入图纸", "原始文件和测试文件分开"],
                ["output", "PDF/DWG输出", "文件名包含日期和图纸编号"],
            ]),
            ("技术模块", ["模块", "实现方式", "应用价值"], [
                ["图框生成", "AutoLISP创建线段、矩形、文字", "减少重复绘图，统一版式"],
                ["图层规范", "自动创建/修改图层颜色线型", "符合制图规范"],
                ["标注系统", "设置文字高度、DIM变量、标注样式", "保证尺寸表达一致"],
                ["批量处理", "SCR脚本、命令行、批量打印", "多张DWG统一清理和输出"],
                ["参数化绘图", "变量控制尺寸、比例、位置", "同一脚本适配多尺寸图纸"],
            ]),
            ("加载与安全检查", ["检查项", "正确做法", "风险"], [
                ["APPLOAD/load", "先加载LSP，再运行自定义命令", "路径错误或文件未加载导致命令未知"],
                ["SECURELOAD", "把脚本放入受信任位置", "安全模式阻止未知路径代码"],
                ["TRUSTEDPATHS", "配置可信脚本目录", "临时目录脚本无法运行"],
                ["系统变量", "脚本结束恢复OSMODE/CMDECHO/CLAYER等", "污染用户绘图环境"],
                ["编码", "2021及以后可用Unicode，旧版注意ASCII", "中文字段乱码"],
            ]),
        ],
        "code": """AutoLISP提示词示例：
请生成一个AutoCAD AutoLISP脚本，命令名为DRAW_A3_FRAME。
功能：以毫米为单位绘制横向A3图框，外框420x297，内边距10，
右下角绘制标题栏，创建图层FRAME、TEXT、DIM，
保存并恢复系统变量，加入错误处理函数，并添加中文字段占位。

加载示例：
(load \"D:/cad_codex_project/lisp/draw_frame.lsp\")
DRAW_A3_FRAME""",
        "video": [
            "展示需求文档：图框尺寸、单位、标题栏和图层要求。",
            "展示Codex生成AutoLISP脚本，并说明参数区。",
            "展示AutoCAD通过APPLOAD或load加载脚本。",
            "运行命令并核对生成图框、图层和标注。",
            "修改参数后重新生成不同尺寸图纸，展示可复用性。",
        ],
        "refs": [
            ("OpenAI Developers, Codex CLI", "https://developers.openai.com/codex/cli"),
            ("OpenAI Developers, Codex Quickstart", "https://developers.openai.com/codex/quickstart"),
            ("Autodesk Help, AutoLISP Developer's Guide", "https://help.autodesk.com/cloudhelp/2022/ENU/AutoCAD-AutoLISP/files/GUID-265AADB3-FB89-4D34-AA9D-6ADF70FF7D4B.htm"),
            ("Autodesk Help, About Loading AutoLISP Applications", "https://help.autodesk.com/cloudhelp/2023/ENU/AutoCAD-Customization/files/GUID-4E633A41-30EA-4C6B-ABC5-11F58970E1EE.htm"),
            ("Autodesk Platform Services, AutoCAD API", "https://aps.autodesk.com/developer/overview/autocad-api"),
        ],
    },
]


def build_report(report: dict, out_path: Path) -> None:
    doc = Document()
    setup_doc(doc)
    add_cover(doc, report["title"], report["subtitle"], report["direction"])
    add_note_box(doc, "技术核对结论", report["summary"], LIGHT_BLUE)

    table_map = {title: (headers, rows) for title, headers, rows in report.get("tables", [])}

    for heading, paragraphs in report["sections"]:
        if report["file"].startswith(("05_", "06_")) and heading.startswith("四、"):
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        add_heading(doc, heading, 1)
        for text in paragraphs:
            add_body(doc, text)
        if heading == "四、关键参数与验证标准" and report["file"].startswith("01_"):
            for title in ("训练类型对比", "核心参数建议", "FLUX训练补充"):
                headers, rows = table_map[title]
                add_heading(doc, title, 2)
                add_table(doc, headers, rows)
        elif heading == "四、工作流设计方法" and report["file"].startswith("02_"):
            for title in ("安装路线对比", "基础工作流模板", "常见问题排查"):
                headers, rows = table_map[title]
                add_heading(doc, title, 2)
                add_table(doc, headers, rows)
        elif heading == "四、关键参数与质量标准" and report["file"].startswith("03_"):
            for title in ("放大方式对比", "参数建议", "质量检查清单"):
                headers, rows = table_map[title]
                add_heading(doc, title, 2)
                add_table(doc, headers, rows)
        elif heading == "四、提示词与参数规范" and report["file"].startswith("04_"):
            for title in ("工作流模块", "节点/工具建议", "输出规格"):
                headers, rows = table_map[title]
                add_heading(doc, title, 2)
                add_table(doc, headers, rows)
        elif heading == "四、脚本功能设计" and report["file"].startswith("05_"):
            for title in ("项目目录建议", "功能模块", "技术注意事项"):
                headers, rows = table_map[title]
                add_heading(doc, title, 2)
                add_table(doc, headers, rows)
            add_heading(doc, "命令行与提示词示例", 2)
            add_code_block(doc, report["code"])
        elif heading == "四、技术模块设计" and report["file"].startswith("06_"):
            for title in ("项目目录建议", "技术模块", "加载与安全检查"):
                headers, rows = table_map[title]
                add_heading(doc, title, 2)
                add_table(doc, headers, rows)
            add_heading(doc, "AutoLISP提示词与加载示例", 2)
            add_code_block(doc, report["code"])

    add_heading(doc, "七、视频展示建议", 1)
    add_numbered(doc, report["video"])
    add_refs(doc, report["refs"])

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = footer.add_run("AI应用技术报告 · 技术增强版")
        set_run_font(r)
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(127, 127, 127)

    tmp = out_path.with_suffix(".tmp.docx")
    doc.save(tmp)
    if out_path.exists():
        out_path.unlink()
    tmp.rename(out_path)


def main() -> None:
    root = ROOT_ARG
    if not root.exists():
        raise SystemExit(f"Root does not exist: {root}")

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup_dir = root / f"_original_backup_{timestamp}"
    backup_dir.mkdir(exist_ok=False)

    for report in REPORTS:
        path = root / report["file"]
        if not path.exists():
            raise SystemExit(f"Missing expected report: {path}")
        shutil.copy2(path, backup_dir / path.name)

    for report in REPORTS:
        build_report(report, root / report["file"])

    readme = root / "README_打开说明.txt"
    if readme.exists():
        shutil.copy2(readme, backup_dir / readme.name)
        readme.write_text(
            "本文件夹包含六份AI应用技术报告（技术增强版）。\n"
            f"修订日期：{ACCESS_DATE}\n\n"
            "修订内容：已逐份加强事实核对、技术参数、实施流程、风险控制、展示建议和参考资料。\n"
            f"原始文件已备份到：{backup_dir.name}\n\n"
            "建议使用 Microsoft Word 或 WPS 打开 .docx 文件；若从网盘或压缩包打开，请先完整解压后再编辑。\n",
            encoding="utf-8",
        )

    print(f"updated={len(REPORTS)}")
    print(f"backup={backup_dir}")
    for report in REPORTS:
        print(root / report["file"])


if __name__ == "__main__":
    main()
